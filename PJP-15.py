#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PJP – 256 Lossless Transforms + 2704 Transform‑Pair Sequences
+ Hybrid Dictionary Mode + Quantum Transforms + Base64 + 6‑bit Text
+ Transforms 28–30 + .docx transforms 31–32
+ Zaden Block Optimization + Algorithm 36 (powers‑of‑two + smart candidates)
  Option 9: Algorithm 37 + Absolute (Hybrid) – tries both, picks the smallest.
            Asks for input/output filenames.
  (Backend: zstandard if available, then paq, else raw)
============================================================================
"""

import math
import random
import decimal
import hashlib
import struct
import re
import os
import urllib.request
import sys
import subprocess
import importlib
import tempfile
import base64
import zipfile
import io
import xml.etree.ElementTree as ET
import time
from typing import Optional, List, Tuple, Dict, Callable
from collections import Counter, defaultdict

# ------------------------------------------------------------------
# Helper: install a single package via pip (silent, auto)
# ------------------------------------------------------------------
def install_package(pkg: str) -> bool:
    print(f"Installing {pkg}...")
    try:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', pkg])
        print(f"Successfully installed {pkg}")
        return True
    except Exception as e:
        print(f"Failed to install {pkg}: {e}")
        return False

# ------------------------------------------------------------------
# 1. Ask about quantum transforms – auto‑install if missing
# ------------------------------------------------------------------
USE_QUANTUM = False
HAS_QISKIT = False

quantum_choice = input("Enable quantum‑inspired transforms (requires Qiskit)? (y/n): ").strip().lower()
if quantum_choice == 'y':
    try:
        from qiskit import QuantumCircuit
        HAS_QISKIT = True
        USE_QUANTUM = True
        print("Quantum transforms ENABLED (Qiskit already installed).")
    except ImportError:
        print("Qiskit not found. Installing automatically...")
        if install_package('qiskit'):
            try:
                from qiskit import QuantumCircuit
                HAS_QISKIT = True
                USE_QUANTUM = True
                print("Quantum transforms ENABLED after automatic installation.")
            except ImportError:
                print("Qiskit installation succeeded but import failed – quantum transforms disabled.")
        else:
            print("Automatic installation failed – quantum transforms disabled.")
else:
    print("Quantum transforms disabled.")

# ------------------------------------------------------------------
# 2. Ask about other optional compression backends (zstandard, paq, etc.)
# ------------------------------------------------------------------
other_choice = input("Install other optional compression backends (zstandard, paq, mpmath, cython, python-docx)? (y/n): ").strip().lower()
if other_choice == 'y':
    for pkg in ['mpmath', 'zstandard', 'cython', 'paq', 'python-docx']:
        try:
            importlib.import_module(pkg)
        except ImportError:
            install_package(pkg)
else:
    print("Skipping other backends.")

# ---------- Optional compression backends ----------
try:
    import paq
except ImportError:
    paq = None

try:
    import zstandard as zstd
    zstd_cctx = zstd.ZstdCompressor(level=22)
    zstd_dctx = zstd.ZstdDecompressor()
    HAS_ZSTD = True
except ImportError:
    HAS_ZSTD = False

# ---------- (Re‑import Qiskit if it was just installed) ----------
if USE_QUANTUM and not HAS_QISKIT:
    try:
        from qiskit import QuantumCircuit
        HAS_QISKIT = True
    except ImportError:
        USE_QUANTUM = False
        print("Quantum transforms disabled because Qiskit could not be imported.")

PROGNAME = "PJP"

# ---------- Dictionary configuration ----------
DICT_DIR = "Dictionaries"
COMBINED_DICTIONARY_FILE = os.path.join(DICT_DIR, "dictionary_combined.txt")

DICTIONARY_FILES = [
    "generated.txt",
    "eng_news_2005_1M-sentences.txt",
    "eng_news_2005_1M-words.txt",
    "eng_news_2005_1M-sources.txt",
    "eng_news_2005_1M-co_n.txt",
    "eng_news_2005_1M-co_s.txt",
    "eng_news_2005_1M-inv_w_2.txt",
    "eng_news_2005_1M-inv_w_3.txt",
    "eng_news_2005_1M-inv_so.txt",
    "eng_news_2005_1M-meta.txt",
    "Dictionary.txt",
    "the-complete-reference-html-css-fifth-edition.txt",
]

DICTIONARY_URLS = [
    "https://drive.google.com/uc?export=download&id=1u_1dCEl8hhdEug6GwkOxHAuSx_6_Pme9",
    "https://drive.google.com/uc?export=download&id=1pVqNN5JZ2AeOCgRaHkv4Vv6Byr4zK20e",
    "https://drive.google.com/uc?export=download&id=1ZSC-Tn76x8itdN0rCp-Zw17hGudxbjxo",
    "https://drive.google.com/uc?export=download&id=1VB_7tzngs4GxjclSRyRDnxgS8znT2w2S",
    "https://drive.google.com/uc?export=download&id=1KVIRgiMrhCUCqQZJ3UT67ztls2GqGJzz",
    "https://drive.google.com/uc?export=download&id=1Z3Lx6SqL4HWsnmbJCez4kXWRQQhUXWKL",
    "https://drive.google.com/uc?export=download&id=1br2bdRMkZEVVRPKYmC4IIaZuAjxFJE4N",
    "https://drive.google.com/uc?export=download&id=1aE6ubPZiJ8rr3lEVk8fFJYjDQ1y1rU0X",
    "https://drive.google.com/uc?export=download&id=1uro3TZe-t5zPx2Qu2xrTL3lU8N0melk9",
    "https://drive.google.com/uc?export=download&id=1HqsTH1DqpWNpGbn9VtD7-SB6wVqA90R2",
    "https://drive.google.com/uc?export=download&id=1zZ8iMeBC3605NZhuc4UE9jx_w_lZFg5B",
    "https://drive.google.com/uc?export=download&id=1dDdqYDgm7f-smS7KF70Wf0KmyFo-ft1M",
]

MAX_LINE_ENTRIES = 1024

def download_and_merge_dictionaries():
    if not os.path.exists(DICT_DIR):
        os.makedirs(DICT_DIR)

    if os.path.exists(COMBINED_DICTIONARY_FILE):
        print(f"Combined dictionary '{COMBINED_DICTIONARY_FILE}' already exists. Skipping download.")
        return True

    all_words = set()
    success_count = 0

    for idx, (filename, url) in enumerate(zip(DICTIONARY_FILES, DICTIONARY_URLS)):
        local_path = os.path.join(DICT_DIR, filename)
        print(f"Downloading {filename} to {DICT_DIR}/ ...")
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response:
                content = response.read()

            if b'<html' in content[:200].lower():
                print(f"  WARNING: {filename} appears to be an HTML page (not a text file). Skipping.")
                continue

            with open(local_path, 'wb') as f:
                f.write(content)

            with open(local_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    w = line.strip()
                    if not w:
                        continue
                    try:
                        decoded = base64.b64decode(w, validate=True)
                        decoded_str = decoded.decode('utf-8')
                        all_words.add(decoded_str)
                    except Exception:
                        all_words.add(w)

            print(f"  Downloaded {filename} ({os.path.getsize(local_path)} bytes)")
            success_count += 1

        except Exception as e:
            print(f"  WARNING: Could not download {filename}: {e}")
            if os.path.exists(local_path):
                os.remove(local_path)

    if success_count == 0:
        print("ERROR: No dictionary files could be downloaded.")
        print("Proceeding without static word and line dictionaries.")
        return False

    try:
        with open(COMBINED_DICTIONARY_FILE, 'w', encoding='utf-8') as f:
            for word in sorted(all_words):
                f.write(word + '\n')
        print(f"Merged {len(all_words)} unique words into {COMBINED_DICTIONARY_FILE} "
              f"({os.path.getsize(COMBINED_DICTIONARY_FILE)} bytes)")
        return True
    except Exception as e:
        print(f"ERROR: Could not write combined dictionary: {e}")
        return False

# ---------- Constants ----------
PRIMES = [p for p in range(2, 256) if all(p % d != 0 for d in range(2, int(p ** 0.5) + 1))]
PI_DIGITS = [79, 17, 111]
BLOCK_SIZE = 1024

def find_nearest_prime_around(n: int) -> int:
    o = 0
    while True:
        c1, c2 = n - o, n + o
        if c1 >= 2 and all(c1 % d != 0 for d in range(2, int(c1 ** 0.5) + 1)):
            return c1
        if c2 >= 2 and all(c2 % d != 0 for d in range(2, int(c2 ** 0.5) + 1)):
            return c2
        o += 1

def sha256_8bytes(data: bytes) -> bytes:
    return hashlib.sha256(data).digest()[:8]

def xor_prime_hash(word: str) -> bytes:
    prime = 2147483647
    total = sum(ord(c) for c in word)
    transformed = total ^ prime
    return transformed.to_bytes(8, 'big')

# ---------- 6‑bit alphabet for transform 27 (exactly 64 chars) ----------
ALPHABET_6BIT = (
    "ABCDEFGHIJKLMNOPQRSTUVWXYZ"    # 26
    "abcdefghijklmnopqrstuvwxyz"    # 26
    "0123456789"                    # 10
    " \n"                           # 2  (space and newline)
)  # Total = 64
assert len(ALPHABET_6BIT) == 64
CHAR_TO_6BIT = {ch: i for i, ch in enumerate(ALPHABET_6BIT)}
SIXBIT_TO_CHAR = {i: ch for ch, i in CHAR_TO_6BIT.items()}

# ---------- Main Compressor Class ----------
class PJPCompressor:
    def __init__(self):
        download_and_merge_dictionaries()

        self.PI_DIGITS = PI_DIGITS.copy()
        self.seed_tables = self._gen_seed_tables(num=126, size=40, seed=42)
        self.fibonacci = self._gen_fib(100)
        self.PI_STR = "3.14159265358979323846264338327950288419716939937510"

        self._build_transform_maps()
        self.sequences = self._build_pair_sequences()
        self.pair_lookup = {idx: (t1, t2) for idx, (t1, t2) in enumerate(self.sequences)}

        self.static_dict, self.word_to_index = self._load_static_dictionary()
        self.line_dict, self.line_to_index = self._load_line_dictionary()

        # Precompute quantum permutations if enabled
        if USE_QUANTUM and HAS_QISKIT:
            self._precompute_quantum_transforms()

    # ------------------------------------------------------------------
    # Quantum transform generation (using Qiskit circuit as seed, no simulation)
    # ------------------------------------------------------------------
    def _generate_permutation_from_circuit(self, num_qubits: int, seed: int) -> List[int]:
        qc = QuantumCircuit(num_qubits)
        rng = random.Random(seed)
        for qubit in range(num_qubits):
            qc.h(qubit)
            qc.rz(rng.random() * 2 * math.pi, qubit)
            qc.rx(rng.random() * 2 * math.pi, qubit)
        for _ in range(num_qubits):
            for i in range(num_qubits - 1):
                qc.cx(i, i+1)
            qc.barrier()
            for i in range(num_qubits):
                qc.rz(rng.random() * 2 * math.pi, i)
                qc.rx(rng.random() * 2 * math.pi, i)

        try:
            qasm_str = qc.qasm()
        except AttributeError:
            qasm_str = qc.draw('text')

        final_seed = seed + hash(qasm_str) % 1000000
        rng2 = random.Random(final_seed)
        n = 1 << num_qubits
        perm = list(range(n))
        rng2.shuffle(perm)
        if num_qubits == 12:  # ultra: need 2704 permutation
            perm_2704 = list(range(2704))
            rng2 = random.Random(final_seed)
            rng2.shuffle(perm_2704)
            return perm_2704
        else:
            return perm

    def _precompute_quantum_transforms(self):
        self.quantum_fast_perms = []
        for i in range(9):
            seed = 1000 + i
            perm = self._generate_permutation_from_circuit(8, seed)
            self.quantum_fast_perms.append(perm)

        self.quantum_ultra_perms = []
        for i in range(17):
            seed = 2000 + i
            perm = self._generate_permutation_from_circuit(12, seed)
            self.quantum_ultra_perms.append(perm)

        self.quantum_fast_transforms = []
        for perm in self.quantum_fast_perms:
            fwd, rev = self._make_substitution_transform(perm, 256)
            self.quantum_fast_transforms.append((fwd, rev))

        self.quantum_ultra_transforms = []
        for perm in self.quantum_ultra_perms:
            fwd, rev = self._make_permutation_transform(perm, 2704)
            self.quantum_ultra_transforms.append((fwd, rev))

        for idx, (fwd, rev) in enumerate(self.quantum_fast_transforms, start=257):
            self.fwd_transforms[idx] = fwd
            self.rev_transforms[idx] = rev
        for idx, (fwd, rev) in enumerate(self.quantum_ultra_transforms, start=266):
            self.fwd_transforms[idx] = fwd
            self.rev_transforms[idx] = rev

    def _make_substitution_transform(self, perm: List[int], size: int):
        inv_perm = [0] * size
        for i, p in enumerate(perm):
            inv_perm[p] = i
        def forward(data: bytes) -> bytes:
            return bytes(perm[b] for b in data)
        def reverse(data: bytes) -> bytes:
            return bytes(inv_perm[b] for b in data)
        return forward, reverse

    def _make_permutation_transform(self, perm: List[int], block_size: int):
        inv_perm = [0] * block_size
        for i, p in enumerate(perm):
            inv_perm[p] = i
        def forward(data: bytes) -> bytes:
            out = bytearray()
            for offset in range(0, len(data), block_size):
                block = data[offset:offset+block_size]
                if len(block) < block_size:
                    out += block
                else:
                    new_block = bytearray(block_size)
                    for i in range(block_size):
                        new_block[perm[i]] = block[i]
                    out += new_block
            return bytes(out)
        def reverse(data: bytes) -> bytes:
            out = bytearray()
            for offset in range(0, len(data), block_size):
                block = data[offset:offset+block_size]
                if len(block) < block_size:
                    out += block
                else:
                    new_block = bytearray(block_size)
                    for i in range(block_size):
                        new_block[inv_perm[i]] = block[i]
                    out += new_block
            return bytes(out)
        return forward, reverse

    # ------------------------------------------------------------------
    # Dictionary loaders
    # ------------------------------------------------------------------
    def _load_static_dictionary(self):
        if not os.path.exists(COMBINED_DICTIONARY_FILE):
            print(f"ERROR: {COMBINED_DICTIONARY_FILE} not found. No dictionaries loaded.")
            return [], {}

        words_set = set()
        try:
            with open(COMBINED_DICTIONARY_FILE, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    w = line.strip()
                    if w:
                        words_set.add(w)
        except Exception as e:
            print(f"Warning: could not read {COMBINED_DICTIONARY_FILE}: {e}")
            return [], {}

        sorted_words = sorted(words_set)
        word_to_idx = {w: i for i, w in enumerate(sorted_words)}
        print(f"Loaded static word dictionary: {len(sorted_words)} unique words.")
        return sorted_words, word_to_idx

    def _load_line_dictionary(self):
        if not os.path.exists(COMBINED_DICTIONARY_FILE):
            return [], {}

        lines = []
        try:
            with open(COMBINED_DICTIONARY_FILE, 'r', encoding='utf-8', errors='ignore') as f:
                for raw_line in f:
                    phrase = raw_line.strip()
                    if phrase and phrase not in lines:
                        lines.append(phrase)
                        if len(lines) >= MAX_LINE_ENTRIES:
                            break
        except Exception as e:
            print(f"Warning: could not read {COMBINED_DICTIONARY_FILE}: {e}")
            return [], {}

        if not lines:
            return [], {}

        lines.sort(key=len, reverse=True)
        line_to_idx = {phrase: i for i, phrase in enumerate(lines)}
        print(f"Loaded line dictionary: {len(lines)} phrases.")
        return lines, line_to_idx

    # ------------------------------------------------------------------
    # pi / constant helpers
    # ------------------------------------------------------------------
    def get_pi_digits(self, n: int) -> str:
        if n < 1: return ""
        return self.PI_STR[2:2 + n]

    def find_lossless_k(self, n: int):
        if n < 1: return 0, True
        true_digits = self.get_pi_digits(n)
        true_scaled = int(self.PI_STR.replace('.', '')[:n + 1])
        DENOM = 16777216
        decimal.getcontext().prec = 50
        pi_dec = decimal.Decimal(self.PI_STR)
        k_float = (pi_dec - 3) * DENOM
        k_candidate = int(round(k_float))
        k_candidate = max(0, min(k_candidate, DENOM - 1))
        approx_scaled = (3 * 10 ** n * DENOM + k_candidate * 10 ** n) // DENOM
        return k_candidate, approx_scaled == true_scaled

    def to_bin(self, value: int, bits: int) -> str:
        return format(value, 'b').zfill(bits)

    def get_bit_size(self, k: int) -> int:
        return 23 if k <= 0x7FFFFF else 25

    def transform_17(self, data: bytes) -> bytes:
        if not data: return b''
        k, _ = self.find_lossless_k(7)
        bits_used = self.get_bit_size(k)
        bit_str = self.to_bin(k, bits_used)
        mask_bytes = []
        for i in range(0, len(bit_str), 8):
            byte_bits = bit_str[i:i + 8]
            if len(byte_bits) < 8:
                byte_bits = byte_bits.ljust(8, '0')
            mask_bytes.append(int(byte_bits, 2))
        mask = bytes(mask_bytes)
        t = bytearray(data)
        for i in range(len(t)):
            t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_17 = transform_17

    def get_basel_digits(self, n: int) -> str:
        decimal.getcontext().prec = n + 5
        pi = decimal.Decimal(self.PI_STR)
        basel = (pi * pi) / decimal.Decimal(6)
        s = str(basel).replace('.', '')
        return s[:n]

    def get_one_over_e_digits(self, n: int) -> str:
        decimal.getcontext().prec = n + 5
        e = decimal.Decimal(1).exp()
        inv_e = decimal.Decimal(1) / e
        s = str(inv_e).replace('.', '')
        return s[:n]

    def get_5e_digits(self, n: int) -> str:
        decimal.getcontext().prec = n + 5
        e = decimal.Decimal(1).exp()
        five_e = decimal.Decimal(5) * e
        s = str(five_e).replace('.', '')
        return s[:n]

    # ------------------------------------------------------------------
    # Seed tables, Fibonacci
    # ------------------------------------------------------------------
    def _gen_seed_tables(self, num=126, size=40, seed=42):
        random.seed(seed)
        return [[random.randint(5, 255) for _ in range(size)] for _ in range(num)]

    def _gen_fib(self, n):
        a, b = 0, 1
        res = [a, b]
        for _ in range(2, n):
            a, b = b, a + b
            res.append(b)
        return res

    def get_seed(self, idx: int, val: int) -> int:
        if 0 <= idx < len(self.seed_tables):
            return self.seed_tables[idx][val % 40]
        return 0

    # ------------------------------------------------------------------
    # Bit helpers (for RLE)
    # ------------------------------------------------------------------
    def _append_bits(self, bitlist: List[int], value: int, count: int):
        for i in range(count - 1, -1, -1):
            bitlist.append((value >> i) & 1)

    def _read_bits(self, bits: List[int], pos: int, count: int) -> int:
        val = 0
        for i in range(count):
            if pos + i >= len(bits): return 0
            val = (val << 1) | bits[pos + i]
        return val

    # ------------------------------------------------------------------
    # RLE transform 00
    # ------------------------------------------------------------------
    def transform_00(self, data: bytes) -> bytes:
        if not data: return b'\x00'
        best_result = None
        best_length = float('inf')
        best_shifts = []
        MAX_PASSES = 10
        current = bytearray(data)
        applied_shifts = []
        for _ in range(MAX_PASSES):
            best_shift = 0
            best_shifted = current
            best_score = float('-inf')
            for shift in range(256):
                tmp = bytearray(current)
                for j in range(len(tmp)):
                    tmp[j] = (tmp[j] + shift) % 256
                score = 0
                i = 0
                while i < len(tmp):
                    val = tmp[i]
                    run = 1
                    i += 1
                    while i < len(tmp) and tmp[i] == val:
                        run += 1
                        i += 1
                    score += run * run
                if score > best_score:
                    best_score = score
                    best_shifted = tmp
                    best_shift = shift
            applied_shifts.append(best_shift)
            rle_encoded = self._apply_rle_to_shifted(best_shifted, best_shift)
            if len(rle_encoded) < best_length:
                best_length = len(rle_encoded)
                best_result = rle_encoded
                best_shifts = applied_shifts.copy()
            current = best_shifted
            if len(rle_encoded) >= len(data):
                break
        if best_result is None or best_length >= len(data):
            return bytes([0]) + data
        header = bytearray([len(best_shifts)])
        header.extend(best_shifts)
        return header + best_result

    def _apply_rle_to_shifted(self, shifted_data: bytearray, shift: int) -> bytes:
        bits = []
        self._append_bits(bits, 0b010, 3)
        self._append_bits(bits, shift, 8)
        i = 0
        n = len(shifted_data)
        while i < n:
            val = shifted_data[i]
            run = 1
            i += 1
            while i < n and shifted_data[i] == val:
                run += 1
                i += 1
            while run >= 13:
                chunk = min(run, 268)
                self._append_bits(bits, 0b1111, 4)
                self._append_bits(bits, chunk - 13, 8)
                self._append_bits(bits, val, 8)
                run -= chunk
            if run == 1:
                self._append_bits(bits, 0b00, 2)
                self._append_bits(bits, val, 8)
            elif run <= 5:
                self._append_bits(bits, 0b01, 2)
                self._append_bits(bits, run - 2, 2)
                self._append_bits(bits, val, 8)
            elif run <= 12:
                self._append_bits(bits, 0b10, 2)
                self._append_bits(bits, run - 6, 3)
                self._append_bits(bits, val, 8)
        pad = (8 - len(bits) % 8) % 8
        self._append_bits(bits, 0, pad)
        out = bytearray()
        for j in range(0, len(bits), 8):
            byte = 0
            for k in range(8):
                if j + k < len(bits):
                    byte = (byte << 1) | bits[j + k]
            out.append(byte)
        return bytes(out)

    def reverse_transform_00(self, cdata: bytes) -> bytes:
        if not cdata or cdata == b'\x00': return b''
        if cdata[0] == 0: return cdata[1:]
        num_passes = cdata[0]
        if num_passes == 0 or len(cdata) < 1 + num_passes: return b''
        shifts = list(cdata[1:1 + num_passes])
        rle_data = cdata[1 + num_passes:]
        decoded = self._rle_decode(rle_data)
        if decoded is None: return b''
        current = bytearray(decoded)
        for shift in reversed(shifts):
            for i in range(len(current)):
                current[i] = (current[i] - shift) % 256
        return bytes(current)

    def _rle_decode(self, data: bytes) -> Optional[bytearray]:
        if not data: return None
        bits = []
        for b in data:
            for i in range(7, -1, -1):
                bits.append((b >> i) & 1)
        pos = 0
        nbits = len(bits)
        if nbits < 11: return None
        marker = self._read_bits(bits, pos, 3)
        pos += 3
        if marker != 0b010: return None
        pos += 8
        out = bytearray()
        while pos < nbits:
            if pos + 2 > nbits: break
            prefix = self._read_bits(bits, pos, 2)
            pos += 2
            if prefix == 0b00:
                if pos + 8 > nbits: break
                run = 1
            elif prefix == 0b01:
                if pos + 2 + 8 > nbits: break
                run = 2 + self._read_bits(bits, pos, 2)
                pos += 2
            elif prefix == 0b10:
                if pos + 3 + 8 > nbits: break
                run = 6 + self._read_bits(bits, pos, 3)
                pos += 3
            else:
                if pos + 2 + 8 + 8 > nbits: break
                if self._read_bits(bits, pos, 2) != 0b11: return None
                pos += 2
                run = 13 + self._read_bits(bits, pos, 8)
                pos += 8
            if pos + 8 > nbits: break
            val = self._read_bits(bits, pos, 8)
            pos += 8
            out.extend([val] * run)
        for i in range(pos, nbits):
            if bits[i] != 0: return None
        return out

    # ------------------------------------------------------------------
    # Transforms 01‑21 (all bijective on bytes except 1,14 which are handled separately)
    # ------------------------------------------------------------------
    def transform_01(self, d, r=100):
        t = bytearray(d)
        for prime in PRIMES:
            xor_val = prime if prime == 2 else max(1, math.ceil(prime * 4096 / 28672))
            for _ in range(r):
                for i in range(0, len(t), 3):
                    if i < len(t): t[i] ^= xor_val
        return bytes(t)
    reverse_transform_01 = transform_01

    def transform_02(self, d):
        if len(d) < 1: return b''
        t = bytearray(d)
        checksum = sum(d) % 256
        pattern_index = (len(d) + checksum) % 256
        pattern_values = self._get_pattern(4, pattern_index)
        for i in range(1, len(t), 4):
            if i < len(t): t[i] ^= pattern_values[i % len(pattern_values)]
        return bytes([pattern_index]) + bytes(t)
    def reverse_transform_02(self, d):
        if len(d) < 2: return b''
        pattern_index = d[0]
        t = bytearray(d[1:])
        pattern_values = self._get_pattern(4, pattern_index)
        for i in range(1, len(t), 4):
            if i < len(t): t[i] ^= pattern_values[i % len(pattern_values)]
        return bytes(t)

    def transform_03(self, d):
        if len(d) < 1: return b''
        t = bytearray(d)
        rotation = (len(d) * 13 + sum(d)) % 8
        if rotation == 0: rotation = 1
        for i in range(2, len(t), 5):
            if i < len(t): t[i] = ((t[i] << rotation) | (t[i] >> (8 - rotation))) & 0xFF
        return bytes([rotation]) + bytes(t)
    def reverse_transform_03(self, d):
        if len(d) < 2: return b''
        rotation = d[0]
        t = bytearray(d[1:])
        for i in range(2, len(t), 5):
            if i < len(t): t[i] = ((t[i] >> rotation) | (t[i] << (8 - rotation))) & 0xFF
        return bytes(t)

    def transform_04(self, d, r=100):
        t = bytearray(d)
        for _ in range(r):
            for i in range(len(t)): t[i] = (t[i] - (i % 256)) % 256
        return bytes(t)
    def reverse_transform_04(self, d, r=100):
        t = bytearray(d)
        for _ in range(r):
            for i in range(len(t)): t[i] = (t[i] + (i % 256)) % 256
        return bytes(t)

    def transform_05(self, d, s=3):
        t = bytearray(d)
        for i in range(len(t)): t[i] = ((t[i] << s) | (t[i] >> (8 - s))) & 0xFF
        return bytes(t)
    def reverse_transform_05(self, d, s=3):
        t = bytearray(d)
        for i in range(len(t)): t[i] = ((t[i] >> s) | (t[i] << (8 - s))) & 0xFF
        return bytes(t)

    def transform_06(self, d, sd=42):
        random.seed(sd)
        sub = list(range(256))
        random.shuffle(sub)
        t = bytearray(d)
        for i in range(len(t)): t[i] = sub[t[i]]
        return bytes(t)
    def reverse_transform_06(self, d, sd=42):
        random.seed(sd)
        sub = list(range(256))
        random.shuffle(sub)
        inv = [0]*256
        for i in range(256): inv[sub[i]] = i
        t = bytearray(d)
        for i in range(len(t)): t[i] = inv[t[i]]
        return bytes(t)

    def transform_07(self, d, r=100):
        t = bytearray(d)
        sh = len(d) % len(self.PI_DIGITS)
        pi_rot = self.PI_DIGITS[sh:] + self.PI_DIGITS[:sh]
        sz = len(d) % 256
        for i in range(len(t)): t[i] ^= sz
        for _ in range(r):
            for i in range(len(t)): t[i] ^= pi_rot[i % len(pi_rot)]
        return bytes(t)
    reverse_transform_07 = transform_07

    def transform_08(self, d, r=100):
        t = bytearray(d)
        sh = len(d) % len(self.PI_DIGITS)
        pi_rot = self.PI_DIGITS[sh:] + self.PI_DIGITS[:sh]
        p = find_nearest_prime_around(len(d) % 256)
        for i in range(len(t)): t[i] ^= p
        for _ in range(r):
            for i in range(len(t)): t[i] ^= pi_rot[i % len(pi_rot)]
        return bytes(t)
    reverse_transform_08 = transform_08

    def transform_09(self, d, r=100):
        t = bytearray(d)
        sh = len(d) % len(self.PI_DIGITS)
        pi_rot = self.PI_DIGITS[sh:] + self.PI_DIGITS[:sh]
        p = find_nearest_prime_around(len(d) % 256)
        seed = self.get_seed(len(d) % len(self.seed_tables), len(d))
        for i in range(len(t)): t[i] ^= p ^ seed
        for _ in range(r):
            for i in range(len(t)): t[i] ^= pi_rot[i % len(pi_rot)] ^ (i % 256)
        return bytes(t)
    reverse_transform_09 = transform_09

    def transform_10(self, data: bytes) -> bytes:
        if not data: return b'\x00'
        cnt = sum(1 for i in range(len(data)-1) if data[i:i+2] == b'X1')
        n = (((cnt * 2) + 1) // 3) * 3 % 256
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= n
        return bytes([n]) + bytes(t)
    def reverse_transform_10(self, data: bytes) -> bytes:
        if len(data) < 1: return b''
        n = data[0]
        t = bytearray(data[1:])
        for i in range(len(t)): t[i] ^= n
        return bytes(t)

    def transform_11(self, data: bytes) -> bytes:
        if not data: return b''
        t = bytearray(data)
        length = len(t)
        for i in range(length):
            fib_idx = (i + length) % len(self.fibonacci)
            fib_val = self.fibonacci[fib_idx] % 256
            pos_val = (i * 13 + length * 17) % 256
            key = (fib_val ^ pos_val) % 256
            t[i] ^= key
        return bytes(t)
    reverse_transform_11 = transform_11

    def transform_12(self, data: bytes) -> bytes:
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= self.fibonacci[i % len(self.fibonacci)] % 256
        return bytes(t)
    reverse_transform_12 = transform_12

    def transform_13(self, d):
        if not d: return b''
        repeats = self._calculate_repeats(d)
        current_value = len(d) % 256
        prime_values = []
        count = 0
        while count < repeats:
            current_value = find_nearest_prime_around(current_value)
            prime_values.append(current_value)
            count += 1
        t = bytearray(d)
        xor_value = prime_values[-1] if prime_values else 0
        for i in range(len(t)): t[i] ^= xor_value
        repeat_byte = (repeats - 1) % 256
        return bytes([repeat_byte]) + bytes(t)
    def reverse_transform_13(self, d):
        if len(d) < 2: return b''
        repeat_byte = d[0]
        repeats = (repeat_byte + 1) % 256
        if repeats == 0: repeats = 256
        t = bytearray(d[1:])
        current_value = len(t) % 256
        prime_values = []
        count = 0
        while count < repeats:
            current_value = find_nearest_prime_around(current_value)
            prime_values.append(current_value)
            count += 1
        xor_value = prime_values[-1] if prime_values else 0
        for i in range(len(t)): t[i] ^= xor_value
        return bytes(t)

    # Transform 14 is NOT bijective; skipped in pair base.

    def transform_15(self, d):
        if len(d) < 1: return b''
        t = bytearray(d)
        pattern_index = len(d) % 256
        pattern_values = self._get_pattern(3, pattern_index)
        for i in range(0, len(t), 3):
            if i < len(t): t[i] = (t[i] + pattern_values[i % len(pattern_values)]) % 256
        return bytes([pattern_index]) + bytes(t)
    def reverse_transform_15(self, d):
        if len(d) < 2: return b''
        pattern_index = d[0]
        t = bytearray(d[1:])
        pattern_values = self._get_pattern(3, pattern_index)
        for i in range(0, len(t), 3):
            if i < len(t): t[i] = (t[i] - pattern_values[i % len(pattern_values)]) % 256
        return bytes(t)

    def transform_16(self, data: bytes) -> bytes:
        if not data: return b''
        xor_byte = (len(data) * 7 + 13) % 256
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= xor_byte
        return bytes(t)
    reverse_transform_16 = transform_16

    # transform_17 defined earlier
    def transform_18(self, data: bytes) -> bytes:
        if not data: return b''
        digits = self.get_basel_digits(max(10, len(data)//2 + 5))
        mask = bytes(int(digits[i:i+2]) % 256 for i in range(0, len(digits), 2))
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_18 = transform_18

    def transform_19(self, data: bytes) -> bytes:
        if not data: return b''
        digits = self.get_one_over_e_digits(max(10, len(data)//2 + 5))
        mask = bytes(int(digits[i:i+2]) % 256 for i in range(0, len(digits), 2))
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_19 = transform_19

    def transform_20(self, data: bytes) -> bytes:
        if not data: return b''
        digits = self.get_5e_digits(max(10, len(data)//2 + 5))
        mask = bytes(int(digits[i:i+2]) % 256 for i in range(0, len(digits), 2))
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_20 = transform_20

    def transform_21(self, data: bytes) -> bytes:
        if not data: return b''
        shift = 255
        t = bytearray(data)
        for i in range(len(t)): t[i] = (t[i] + shift) % 256
        return bytes(t)
    def reverse_transform_21(self, data: bytes) -> bytes:
        if not data: return b''
        shift = 255
        t = bytearray(data)
        for i in range(len(t)): t[i] = (t[i] - shift) % 256
        return bytes(t)

    # ------------------------------------------------------------------
    # Transform 22 – Base64 encode/decode (NOT bijective; skipped in pair base)
    # ------------------------------------------------------------------
    def transform_22(self, data: bytes) -> bytes:
        return base64.b64encode(data)

    def reverse_transform_22(self, data: bytes) -> bytes:
        try:
            return base64.b64decode(data, validate=False)
        except Exception:
            return data

    # ------------------------------------------------------------------
    # Transform 23 – SHA‑256 word tokenizer (text‑only, NOT bijective)
    # ------------------------------------------------------------------
    def transform_23(self, data: bytes) -> bytes:
        if not data: return b'\x00\x00\x00\x00'
        try:
            text = data.decode('latin-1')
        except:
            text = data.decode('latin-1', errors='replace')
        pattern = r'([A-Za-z0-9_]+)'
        tokens = re.split(pattern, text)
        hash_to_word = {}
        token_list = []
        for i, tok in enumerate(tokens):
            if i % 2 == 1:
                word_bytes = tok.encode('latin-1')
                h = sha256_8bytes(word_bytes)
                if h in hash_to_word:
                    if hash_to_word[h] != word_bytes:
                        token_list.append((False, word_bytes))
                        continue
                else:
                    hash_to_word[h] = word_bytes
                token_list.append((True, h))
            else:
                if tok:
                    token_list.append((False, tok.encode('latin-1')))
        dict_entries = list(hash_to_word.items())
        num_entries = len(dict_entries)
        result = bytearray()
        result += struct.pack('>I', num_entries)
        for h, wb in dict_entries:
            result += h
            result += struct.pack('>H', len(wb))
            result += wb
        for is_word, payload in token_list:
            if is_word:
                result += b'\x01'
                result += payload
            else:
                result += b'\x00'
                result += struct.pack('>H', len(payload))
                result += payload
        return bytes(result)

    def reverse_transform_23(self, data: bytes) -> bytes:
        if not data: return b''
        if len(data) < 4: return data
        num_entries = struct.unpack('>I', data[:4])[0]
        pos = 4
        hash_to_word = {}
        for _ in range(num_entries):
            if pos + 10 > len(data): break
            h = data[pos:pos+8]
            pos += 8
            wlen = struct.unpack('>H', data[pos:pos+2])[0]
            pos += 2
            if pos + wlen > len(data): break
            wb = data[pos:pos+wlen]
            pos += wlen
            hash_to_word[h] = wb
        out = bytearray()
        while pos < len(data):
            if pos >= len(data): break
            typ = data[pos]
            pos += 1
            if typ == 1:
                if pos + 8 > len(data): break
                h = data[pos:pos+8]
                pos += 8
                wb = hash_to_word.get(h)
                out += wb if wb else h
            elif typ == 0:
                if pos + 2 > len(data): break
                rawlen = struct.unpack('>H', data[pos:pos+2])[0]
                pos += 2
                if pos + rawlen > len(data): break
                out += data[pos:pos+rawlen]
                pos += rawlen
            else:
                break
        return bytes(out)

    # ------------------------------------------------------------------
    # Transform 24 – XOR‑prime word tokenizer (text‑only, NOT bijective)
    # ------------------------------------------------------------------
    def transform_24(self, data: bytes) -> bytes:
        if not data: return b'\x00\x00\x00\x00'
        try:
            text = data.decode('latin-1')
        except:
            text = data.decode('latin-1', errors='replace')
        pattern = r'([A-Za-z0-9_]+)'
        tokens = re.split(pattern, text)
        hash_to_word = {}
        token_list = []
        for i, tok in enumerate(tokens):
            if i % 2 == 1:
                word_bytes = tok.encode('latin-1')
                h = xor_prime_hash(tok)
                if h in hash_to_word:
                    if hash_to_word[h] != word_bytes:
                        token_list.append((False, word_bytes))
                        continue
                else:
                    hash_to_word[h] = word_bytes
                token_list.append((True, h))
            else:
                if tok:
                    token_list.append((False, tok.encode('latin-1')))
        dict_entries = list(hash_to_word.items())
        num_entries = len(dict_entries)
        result = bytearray()
        result += struct.pack('>I', num_entries)
        for h, wb in dict_entries:
            result += h
            result += struct.pack('>H', len(wb))
            result += wb
        for is_word, payload in token_list:
            if is_word:
                result += b'\x01'
                result += payload
            else:
                result += b'\x00'
                result += struct.pack('>H', len(payload))
                result += payload
        return bytes(result)

    def reverse_transform_24(self, data: bytes) -> bytes:
        if not data: return b''
        if len(data) < 4: return data
        num_entries = struct.unpack('>I', data[:4])[0]
        pos = 4
        hash_to_word = {}
        for _ in range(num_entries):
            if pos + 10 > len(data): break
            h = data[pos:pos+8]
            pos += 8
            wlen = struct.unpack('>H', data[pos:pos+2])[0]
            pos += 2
            if pos + wlen > len(data): break
            wb = data[pos:pos+wlen]
            pos += wlen
            hash_to_word[h] = wb
        out = bytearray()
        while pos < len(data):
            if pos >= len(data): break
            typ = data[pos]
            pos += 1
            if typ == 1:
                if pos + 8 > len(data): break
                h = data[pos:pos+8]
                pos += 8
                wb = hash_to_word.get(h)
                out += wb if wb else h
            elif typ == 0:
                if pos + 2 > len(data): break
                rawlen = struct.unpack('>H', data[pos:pos+2])[0]
                pos += 2
                if pos + rawlen > len(data): break
                out += data[pos:pos+rawlen]
                pos += rawlen
            else:
                break
        return bytes(out)

    # ------------------------------------------------------------------
    # Transform 25 – Dynamic Dictionary Tokenizer (text‑only, NOT bijective)
    # ------------------------------------------------------------------
    def _split_text_into_chunks(self, text: str, level: str = 'all') -> List[str]:
        if level == 'paragraph':
            return re.split(r'(\n\n)', text)
        elif level == 'line':
            return re.split(r'(\n)', text)
        elif level == 'sentence':
            return re.split(r'([.!?]+)', text)
        elif level == 'word':
            return re.split(r'(\s+|\b)', text)
        else:
            chunks = []
            paragraphs = re.split(r'(\n\n)', text)
            for i, para in enumerate(paragraphs):
                if i % 2 == 1:
                    chunks.append(para)
                    continue
                lines = re.split(r'(\n)', para)
                for j, line in enumerate(lines):
                    if j % 2 == 1:
                        chunks.append(line)
                        continue
                    sentences = re.split(r'([.!?]+)', line)
                    for k, sent in enumerate(sentences):
                        if k % 2 == 1:
                            chunks.append(sent)
                            continue
                        words = re.split(r'(\s+|\b)', sent)
                        chunks.extend(words)
            return chunks

    def _dynamic_dict_tokenize(self, data: bytes, index_bytes: int = 3) -> bytes:
        try:
            text = data.decode('utf-8')
        except:
            return b'\x00' + data
        chunks = self._split_text_into_chunks(text, 'all')
        freq = Counter(chunks)
        sorted_chunks = sorted(freq.keys(), key=lambda x: (-freq[x], -len(x), x))
        chunk_to_idx = {ch: i for i, ch in enumerate(sorted_chunks)}
        num_entries = len(sorted_chunks)
        if index_bytes == 2 and num_entries > 65535:
            index_bytes = 3
        if index_bytes == 3 and num_entries > 16777215:
            index_bytes = 8
        header = bytearray()
        header.append(index_bytes)
        header += struct.pack('>I', num_entries)
        for chunk in sorted_chunks:
            chunk_bytes = chunk.encode('utf-8')
            header += struct.pack('>I', len(chunk_bytes))
            header += chunk_bytes
        token_stream = bytearray()
        for chunk in chunks:
            idx = chunk_to_idx[chunk]
            if index_bytes == 2:
                token_stream += struct.pack('>H', idx)
            elif index_bytes == 3:
                token_stream += struct.pack('>I', idx)[1:4]
            else:
                token_stream += struct.pack('>Q', idx)
        return bytes(header) + bytes(token_stream)

    def _dynamic_dict_detokenize(self, data: bytes) -> Optional[bytes]:
        if not data: return b''
        if data[0] == 0: return data[1:]
        index_bytes = data[0]
        if index_bytes not in (2, 3, 8): return None
        pos = 1
        if pos + 4 > len(data): return None
        num_entries = struct.unpack('>I', data[pos:pos+4])[0]
        pos += 4
        dictionary = []
        for _ in range(num_entries):
            if pos + 4 > len(data): return None
            chunk_len = struct.unpack('>I', data[pos:pos+4])[0]
            pos += 4
            if pos + chunk_len > len(data): return None
            chunk = data[pos:pos+chunk_len].decode('utf-8')
            pos += chunk_len
            dictionary.append(chunk)
        tokens = []
        while pos < len(data):
            if index_bytes == 2:
                if pos + 2 > len(data): break
                idx = struct.unpack('>H', data[pos:pos+2])[0]
                pos += 2
            elif index_bytes == 3:
                if pos + 3 > len(data): break
                idx_bytes = b'\x00' + data[pos:pos+3]
                idx = struct.unpack('>I', idx_bytes)[0]
                pos += 3
            else:
                if pos + 8 > len(data): break
                idx = struct.unpack('>Q', data[pos:pos+8])[0]
                pos += 8
            if idx < len(dictionary):
                tokens.append(dictionary[idx])
            else:
                return None
        try:
            text = ''.join(tokens)
            return text.encode('utf-8')
        except:
            return None

    def transform_25(self, data: bytes) -> bytes:
        return self._dynamic_dict_tokenize(data, index_bytes=3)

    def reverse_transform_25(self, data: bytes) -> bytes:
        result = self._dynamic_dict_detokenize(data)
        return result if result is not None else b''

    # ------------------------------------------------------------------
    # Transform 26 – SHA‑256 block masking (bijective, but we exclude to be safe)
    # ------------------------------------------------------------------
    def transform_26(self, data: bytes) -> bytes:
        if not data: return b''
        secret = b"PJP_TRANSFORM26_SECRET"
        result = bytearray()
        for idx in range(0, len(data), BLOCK_SIZE):
            chunk = data[idx:idx+BLOCK_SIZE]
            block_num = idx // BLOCK_SIZE
            hasher = hashlib.sha256()
            hasher.update(secret)
            hasher.update(struct.pack(">Q", block_num))
            mask = hasher.digest()
            mask_repeated = (mask * ((len(chunk) // len(mask)) + 1))[:len(chunk)]
            xored = bytes(a ^ b for a, b in zip(chunk, mask_repeated))
            result.extend(xored)
        return bytes(result)

    def reverse_transform_26(self, data: bytes) -> bytes:
        return self.transform_26(data)

    # ------------------------------------------------------------------
    # Transform 27 – 6‑bit text compression (text‑only, NOT bijective)
    # ------------------------------------------------------------------
    def transform_27(self, data: bytes) -> bytes:
        """Encode text using 6‑bit alphabet and pack into bytes."""
        try:
            text = data.decode('utf-8')
        except UnicodeDecodeError:
            return data

        for ch in text:
            if ch not in CHAR_TO_6BIT:
                return data

        bits = []
        for ch in text:
            val = CHAR_TO_6BIT[ch]
            for i in range(5, -1, -1):
                bits.append((val >> i) & 1)

        pad = (8 - len(bits) % 8) % 8
        bits.extend([0] * pad)

        out = bytearray()
        for i in range(0, len(bits), 8):
            byte = 0
            for j in range(8):
                byte = (byte << 1) | bits[i + j]
            out.append(byte)

        length_bytes = struct.pack('<I', len(text))
        return length_bytes + bytes(out)

    def reverse_transform_27(self, data: bytes) -> bytes:
        if len(data) < 4:
            return data
        num_chars = struct.unpack('<I', data[:4])[0]
        packed = data[4:]

        bits = []
        for b in packed:
            for i in range(7, -1, -1):
                bits.append((b >> i) & 1)

        needed_bits = num_chars * 6
        if len(bits) < needed_bits:
            return data

        chars = []
        for i in range(num_chars):
            val = 0
            for j in range(6):
                val = (val << 1) | bits[i*6 + j]
            if val < 64:
                chars.append(SIXBIT_TO_CHAR[val])
            else:
                return data

        try:
            return ''.join(chars).encode('utf-8')
        except UnicodeEncodeError:
            return data

    # ------------------------------------------------------------------
    # Transforms 28, 29, 30 – per‑3‑byte subtract variants
    # ------------------------------------------------------------------
    def transform_28(self, data: bytes) -> bytes:
        if not data:
            return b''
        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        out = bytearray([pad_len])
        for i in range(0, len(padded), 3):
            chunk = padded[i:i+3]
            val = int.from_bytes(chunk, 'little')
            block_idx = i // 3
            key = (block_idx * 65537 + 12345) & 0xFFFF
            new_val = (val - key) % (1 << 24)
            out.extend(new_val.to_bytes(3, 'little'))
        return bytes(out)

    def reverse_transform_28(self, data: bytes) -> bytes:
        if not data:
            return b''
        pad_len = data[0]
        payload = data[1:]
        if len(payload) % 3 != 0:
            return data
        out = bytearray()
        for i in range(0, len(payload), 3):
            chunk = payload[i:i+3]
            val = int.from_bytes(chunk, 'little')
            block_idx = i // 3
            key = (block_idx * 65537 + 12345) & 0xFFFF
            orig_val = (val + key) % (1 << 24)
            out.extend(orig_val.to_bytes(3, 'little'))
        if pad_len > 0:
            out = out[:-pad_len]
        return bytes(out)

    def _find_best_16bit_key(self, data: bytes, quantum_boost: bool = False, time_limit: float = 60.0) -> int:
        """
        Zaden core: find the best 16‑bit key (0..65535) that minimizes the sum of
        absolute deviations from the mean after subtracting the key from each 3‑byte chunk.
        Search stops when time_limit (seconds) is exceeded; best key found so far is returned.
        """
        if len(data) < 3:
            return 0
        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        values = []
        for i in range(0, len(padded), 3):
            values.append(int.from_bytes(padded[i:i+3], 'little'))

        start_time = time.time()
        best_key = 0
        best_cost = float('inf')

        if not quantum_boost or not HAS_QISKIT:
            # Exhaustive search over 65,536 keys, but break if time limit exceeded
            for key in range(65536):
                # Check time every 1024 keys to reduce overhead
                if key % 1024 == 0:
                    if time.time() - start_time > time_limit:
                        break
                trans = [((v - key) & 0xFFFFFF) for v in values]
                mean_t = sum(trans) // len(trans)
                cost = sum(abs(t - mean_t) for t in trans)
                if cost < best_cost:
                    best_cost = cost
                    best_key = key
                    if cost == 0:
                        break
            return best_key
        else:
            # Quantum‑boosted: use a quantum circuit to generate a seed,
            # then shuffle the key order and test the first N keys.
            from qiskit import QuantumCircuit
            qc = QuantumCircuit(8)
            for i in range(8):
                qc.h(i)
                qc.rz(random.random() * 2 * math.pi, i)
            try:
                qasm = qc.qasm()
                seed = hash(qasm) & 0xFFFFFFFF
            except:
                seed = 42
            rng = random.Random(seed)
            keys = list(range(65536))
            rng.shuffle(keys)
            # Test keys until time limit is reached
            for i, key in enumerate(keys):
                if i % 1024 == 0:
                    if time.time() - start_time > time_limit:
                        break
                trans = [((v - key) & 0xFFFFFF) for v in values]
                mean_t = sum(trans) // len(trans)
                cost = sum(abs(t - mean_t) for t in trans)
                if cost < best_cost:
                    best_cost = cost
                    best_key = key
                    if cost == 0:
                        break
            return best_key

    def transform_29(self, data: bytes, quantum_boost: bool = False, time_limit: float = 60.0) -> bytes:
        if not data:
            return b''
        best_key = self._find_best_16bit_key(data, quantum_boost, time_limit)
        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        out = bytearray([pad_len])
        out.extend(best_key.to_bytes(2, 'little'))
        for i in range(0, len(padded), 3):
            chunk = padded[i:i+3]
            val = int.from_bytes(chunk, 'little')
            new_val = (val - best_key) % (1 << 24)
            out.extend(new_val.to_bytes(3, 'little'))
        return bytes(out)

    def reverse_transform_29(self, data: bytes) -> bytes:
        if not data or len(data) < 3:
            return data
        pad_len = data[0]
        if len(data) < 1 + 2:
            return data
        key = int.from_bytes(data[1:3], 'little')
        payload = data[3:]
        if len(payload) % 3 != 0:
            return data
        out = bytearray()
        for i in range(0, len(payload), 3):
            chunk = payload[i:i+3]
            val = int.from_bytes(chunk, 'little')
            orig_val = (val + key) % (1 << 24)
            out.extend(orig_val.to_bytes(3, 'little'))
        if pad_len > 0:
            out = out[:-pad_len]
        return bytes(out)

    def _find_best_24bit_key_heuristic(self, data: bytes) -> int:
        if len(data) < 3:
            return 0
        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        values = []
        for i in range(0, len(padded), 3):
            val = int.from_bytes(padded[i:i+3], 'little')
            values.append(val)
        mean = sum(values) // len(values)
        sorted_vals = sorted(values)
        median = sorted_vals[len(sorted_vals)//2]
        candidates = set()
        for base in [mean, median]:
            for offset in [0, 1, -1, 10, -10, 100, -100, 1000, -1000]:
                cand = (base + offset) % (1 << 24)
                candidates.add(cand)
        rng = random.Random(42)
        for _ in range(10):
            candidates.add(rng.randint(0, (1 << 24) - 1))
        best_key = 0
        best_cost = float('inf')
        for key in candidates:
            trans = [((v - key) & 0xFFFFFF) for v in values]
            mean_t = sum(trans) // len(trans)
            cost = sum(abs(t - mean_t) for t in trans)
            if cost < best_cost:
                best_cost = cost
                best_key = key
        return best_key

    def transform_30(self, data: bytes) -> bytes:
        if not data:
            return b''
        best_key = self._find_best_24bit_key_heuristic(data)
        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        out = bytearray([pad_len])
        out.extend(best_key.to_bytes(3, 'little'))
        for i in range(0, len(padded), 3):
            chunk = padded[i:i+3]
            val = int.from_bytes(chunk, 'little')
            new_val = (val - best_key) % (1 << 24)
            out.extend(new_val.to_bytes(3, 'little'))
        return bytes(out)

    def reverse_transform_30(self, data: bytes) -> bytes:
        if not data or len(data) < 4:
            return data
        pad_len = data[0]
        if len(data) < 1 + 3:
            return data
        key = int.from_bytes(data[1:4], 'little')
        payload = data[4:]
        if len(payload) % 3 != 0:
            return data
        out = bytearray()
        for i in range(0, len(payload), 3):
            chunk = payload[i:i+3]
            val = int.from_bytes(chunk, 'little')
            orig_val = (val + key) % (1 << 24)
            out.extend(orig_val.to_bytes(3, 'little'))
        if pad_len > 0:
            out = out[:-pad_len]
        return bytes(out)

    # ------------------------------------------------------------------
    # Transform 31 – .docx paragraph extraction with dictionary compression
    # (length‑prefixed runs for losslessness)
    # ------------------------------------------------------------------
    def _build_text_dictionary(self, text_streams: List[str], min_freq: int = 2) -> Tuple[List[str], Dict[str, int]]:
        all_tokens = []
        for text in text_streams:
            words = re.findall(r'\b[\w\-]+\b', text)
            all_tokens.extend(words)
        freq = Counter(all_tokens)
        common = [word for word, cnt in freq.items() if cnt >= min_freq]
        common.sort(key=lambda w: (-freq[w], -len(w), w))
        dictionary = common
        word_to_idx = {w: i for i, w in enumerate(dictionary)}
        return dictionary, word_to_idx

    def _encode_text_with_dict(self, text: str, dictionary: List[str], word_to_idx: Dict[str, int]) -> bytes:
        pattern = re.compile(r'(\b[\w\-]+\b)')
        parts = pattern.split(text)
        encoded = bytearray()
        for i, part in enumerate(parts):
            if i % 2 == 1:
                if part in word_to_idx:
                    idx = word_to_idx[part]
                    if len(dictionary) <= 255:
                        encoded.append(0x00)
                        encoded.append(idx)
                    elif len(dictionary) <= 65535:
                        encoded.append(0x01)
                        encoded.extend(struct.pack('>H', idx))
                    elif len(dictionary) <= 16777215:
                        encoded.append(0x02)
                        encoded.extend(struct.pack('>I', idx)[1:4])
                    else:
                        encoded.append(0x03)
                        encoded.extend(struct.pack('>Q', idx))
                else:
                    encoded.append(0x04)
                    word_bytes = part.encode('utf-8')
                    encoded.append(len(word_bytes))
                    encoded.extend(word_bytes)
            else:
                if part:
                    encoded.append(0x04)
                    raw_bytes = part.encode('utf-8')
                    encoded.append(len(raw_bytes))
                    encoded.extend(raw_bytes)
        return bytes(encoded)

    def _decode_text_with_dict(self, data: bytes, dictionary: List[str]) -> str:
        pos = 0
        out = []
        while pos < len(data):
            marker = data[pos]
            pos += 1
            if marker == 0x00:
                if pos >= len(data): break
                idx = data[pos]
                pos += 1
                if idx < len(dictionary):
                    out.append(dictionary[idx])
                else:
                    out.append(f"<ERR{idx}>")
            elif marker == 0x01:
                if pos + 1 >= len(data): break
                idx = struct.unpack('>H', data[pos:pos+2])[0]
                pos += 2
                if idx < len(dictionary):
                    out.append(dictionary[idx])
                else:
                    out.append(f"<ERR{idx}>")
            elif marker == 0x02:
                if pos + 2 >= len(data): break
                idx = struct.unpack('>I', b'\x00' + data[pos:pos+3])[0]
                pos += 3
                if idx < len(dictionary):
                    out.append(dictionary[idx])
                else:
                    out.append(f"<ERR{idx}>")
            elif marker == 0x03:
                if pos + 7 >= len(data): break
                idx = struct.unpack('>Q', data[pos:pos+8])[0]
                pos += 8
                if idx < len(dictionary):
                    out.append(dictionary[idx])
                else:
                    out.append(f"<ERR{idx}>")
            elif marker == 0x04:
                if pos >= len(data): break
                length = data[pos]
                pos += 1
                if pos + length > len(data): break
                raw = data[pos:pos+length]
                pos += length
                out.append(raw.decode('utf-8', errors='replace'))
            else:
                break
        return ''.join(out)

    def transform_31(self, data: bytes) -> bytes:
        if not data or len(data) < 4 or data[:4] != b'PK\x03\x04':
            return b'\x00' + data

        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document(io.BytesIO(data))
        except ImportError:
            # Fallback: XML only, no formatting
            try:
                with zipfile.ZipFile(io.BytesIO(data)) as zf:
                    with zf.open('word/document.xml') as f:
                        xml = f.read()
                root = ET.fromstring(xml)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                text_parts = []
                for t in root.findall('.//w:t', ns):
                    if t.text:
                        text_parts.append(t.text)
                full_text = ''.join(text_parts)
                if not full_text:
                    return b'\x00' + data
                dict_list, word_to_idx = self._build_text_dictionary([full_text])
                encoded_text = self._encode_text_with_dict(full_text, dict_list, word_to_idx)
                out = bytearray()
                out.append(0x01)
                out.append(len(dict_list))
                for word in dict_list:
                    wb = word.encode('utf-8')
                    out.extend(struct.pack('>H', len(wb)))
                    out.extend(wb)
                out.extend(encoded_text)
                return bytes(out)
            except Exception:
                return b'\x00' + data
        else:
            paragraphs_text = []
            for para in doc.paragraphs:
                para_text = ''.join(run.text for run in para.runs if run.text)
                if para_text:
                    paragraphs_text.append(para_text)
            full_text = '\n'.join(paragraphs_text)
            if not full_text:
                return b'\x00' + data

            dict_list, word_to_idx = self._build_text_dictionary([full_text])

            out = bytearray()
            out.append(0x01)
            out.append(len(dict_list))
            for word in dict_list:
                wb = word.encode('utf-8')
                out.extend(struct.pack('>H', len(wb)))
                out.extend(wb)

            for para in doc.paragraphs:
                for run in para.runs:
                    text = run.text
                    if not text:
                        continue
                    encoded_run = self._encode_text_with_dict(text, dict_list, word_to_idx)
                    size = run.font.size
                    size_val = int(size.pt) if size is not None else 12
                    style = 0
                    if run.bold: style |= 1
                    if run.italic: style |= 2
                    if run.underline: style |= 4
                    if run.font.strike: style |= 8
                    if run.font.superscript: style |= 16
                    if run.font.subscript: style |= 32
                    out.append(0x05)                # run marker
                    out.append(size_val)
                    out.append(style)
                    out.extend(struct.pack('>H', len(encoded_run)))  # length prefix
                    out.extend(encoded_run)
            return bytes(out)

    def reverse_transform_31(self, data: bytes) -> bytes:
        if not data:
            return b''
        if data[0] == 0x00:
            return data[1:]
        if data[0] != 0x01:
            return data

        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return data

        pos = 1
        if pos >= len(data):
            return data
        num_words = data[pos]
        pos += 1
        dictionary = []
        for _ in range(num_words):
            if pos + 2 > len(data):
                break
            wlen = struct.unpack('>H', data[pos:pos+2])[0]
            pos += 2
            if pos + wlen > len(data):
                break
            word = data[pos:pos+wlen].decode('utf-8')
            pos += wlen
            dictionary.append(word)

        doc = Document()
        p = doc.add_paragraph()

        while pos < len(data):
            marker = data[pos]
            pos += 1
            if marker == 0x05:
                if pos + 2 > len(data):
                    break
                size_val = data[pos]
                pos += 1
                style = data[pos]
                pos += 1
                if pos + 2 > len(data):
                    break
                run_len = struct.unpack('>H', data[pos:pos+2])[0]
                pos += 2
                if pos + run_len > len(data):
                    break
                run_data = data[pos:pos+run_len]
                pos += run_len
                decoded_text = self._decode_text_with_dict(run_data, dictionary)
                run = p.add_run(decoded_text)
                run.font.size = Pt(size_val)
                if style & 1: run.bold = True
                if style & 2: run.italic = True
                if style & 4: run.underline = True
                if style & 8: run.font.strike = True
                if style & 16: run.font.superscript = True
                if style & 32: run.font.subscript = True
            else:
                break

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # ------------------------------------------------------------------
    # Transform 32 – .docx table extraction with dictionary compression
    # (length‑prefixed runs for losslessness)
    # ------------------------------------------------------------------
    def transform_32(self, data: bytes) -> bytes:
        if not data or len(data) < 4 or data[:4] != b'PK\x03\x04':
            return b'\x00' + data

        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document(io.BytesIO(data))
        except ImportError:
            return b'\x00' + data

        tables = doc.tables
        if not tables:
            return b'\x00' + data

        all_text = []
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        full_text = '\n'.join(all_text)
        if not full_text:
            return b'\x00' + data

        dict_list, word_to_idx = self._build_text_dictionary([full_text])

        out = bytearray()
        out.append(0x02)
        out.append(len(dict_list))
        for word in dict_list:
            wb = word.encode('utf-8')
            out.extend(struct.pack('>H', len(wb)))
            out.extend(wb)

        for table in tables:
            rows = len(table.rows)
            cols = len(table.rows[0].cells) if rows > 0 else 0
            out.append(rows)
            out.append(cols)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if not run.text:
                                continue
                            encoded_run = self._encode_text_with_dict(run.text, dict_list, word_to_idx)
                            size = run.font.size
                            size_val = int(size.pt) if size is not None else 12
                            style = 0
                            if run.bold: style |= 1
                            if run.italic: style |= 2
                            if run.underline: style |= 4
                            if run.font.strike: style |= 8
                            if run.font.superscript: style |= 16
                            if run.font.subscript: style |= 32
                            out.append(0x06)
                            out.append(size_val)
                            out.append(style)
                            out.extend(struct.pack('>H', len(encoded_run)))
                            out.extend(encoded_run)
                    out.append(0x00)   # end of cell
        return bytes(out)

    def reverse_transform_32(self, data: bytes) -> bytes:
        if not data:
            return b''
        if data[0] == 0x00:
            return data[1:]
        if data[0] != 0x02:
            return data

        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return data

        pos = 1
        if pos >= len(data):
            return data
        num_words = data[pos]
        pos += 1
        dictionary = []
        for _ in range(num_words):
            if pos + 2 > len(data):
                break
            wlen = struct.unpack('>H', data[pos:pos+2])[0]
            pos += 2
            if pos + wlen > len(data):
                break
            word = data[pos:pos+wlen].decode('utf-8')
            pos += wlen
            dictionary.append(word)

        doc = Document()
        while pos < len(data):
            if pos >= len(data):
                break
            rows = data[pos]
            pos += 1
            if pos >= len(data):
                break
            cols = data[pos]
            pos += 1
            table = doc.add_table(rows=rows, cols=cols)
            for r in range(rows):
                for c in range(cols):
                    cell = table.cell(r, c)
                    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    while pos < len(data) and data[pos] != 0x00:
                        marker = data[pos]
                        pos += 1
                        if marker == 0x06:
                            if pos + 2 > len(data):
                                break
                            size_val = data[pos]
                            pos += 1
                            style = data[pos]
                            pos += 1
                            if pos + 2 > len(data):
                                break
                            run_len = struct.unpack('>H', data[pos:pos+2])[0]
                            pos += 2
                            if pos + run_len > len(data):
                                break
                            run_data = data[pos:pos+run_len]
                            pos += run_len
                            decoded_text = self._decode_text_with_dict(run_data, dictionary)
                            run = p.add_run(decoded_text)
                            run.font.size = Pt(size_val)
                            if style & 1: run.bold = True
                            if style & 2: run.italic = True
                            if style & 4: run.underline = True
                            if style & 8: run.font.strike = True
                            if style & 16: run.font.superscript = True
                            if style & 32: run.font.subscript = True
                        else:
                            break
                    # skip the 0x00 that ended the cell
                    if pos < len(data) and data[pos] == 0x00:
                        pos += 1
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # ------------------------------------------------------------------
    # Helpers: pattern, repeats, dynamic transform
    # ------------------------------------------------------------------
    def _get_pattern(self, size: int, index: int):
        random.seed(12345 + size * 100 + index)
        return [random.randint(0, 255) for _ in range(size)]

    def _calculate_repeats(self, data: bytes) -> int:
        if not data: return 1
        length = len(data)
        byte_sum = sum(data) % 256
        repeats = ((length * 13 + byte_sum * 17) % 256) + 1
        return max(1, min(256, repeats))

    def _dynamic_transform(self, n: int):
        def tf(data: bytes):
            if not data: return b''
            seed = self.get_seed(n % len(self.seed_tables), len(data))
            t = bytearray(data)
            for i in range(len(t)): t[i] ^= seed
            return bytes(t)
        return tf, tf

    # ------------------------------------------------------------------
    # Build transform maps (1..256) including 31 and 32
    # ------------------------------------------------------------------
    def _build_transform_maps(self):
        self.fwd_transforms: Dict[int, Callable] = {}
        self.rev_transforms: Dict[int, Callable] = {}

        # 1‑21
        self.fwd_transforms[1] = self.transform_00; self.rev_transforms[1] = self.reverse_transform_00
        self.fwd_transforms[2] = self.transform_01; self.rev_transforms[2] = self.reverse_transform_01
        self.fwd_transforms[3] = self.transform_02; self.rev_transforms[3] = self.reverse_transform_02
        self.fwd_transforms[4] = self.transform_03; self.rev_transforms[4] = self.reverse_transform_03
        self.fwd_transforms[5] = self.transform_04; self.rev_transforms[5] = self.reverse_transform_04
        self.fwd_transforms[6] = self.transform_05; self.rev_transforms[6] = self.reverse_transform_05
        self.fwd_transforms[7] = self.transform_06; self.rev_transforms[7] = self.reverse_transform_06
        self.fwd_transforms[8] = self.transform_07; self.rev_transforms[8] = self.reverse_transform_07
        self.fwd_transforms[9] = self.transform_08; self.rev_transforms[9] = self.reverse_transform_08
        self.fwd_transforms[10] = self.transform_09; self.rev_transforms[10] = self.reverse_transform_09
        self.fwd_transforms[11] = self.transform_10; self.rev_transforms[11] = self.reverse_transform_10
        self.fwd_transforms[12] = self.transform_11; self.rev_transforms[12] = self.reverse_transform_11
        self.fwd_transforms[13] = self.transform_12; self.rev_transforms[13] = self.reverse_transform_12
        self.fwd_transforms[14] = self.transform_13; self.rev_transforms[14] = self.reverse_transform_13
        self.fwd_transforms[15] = self.transform_15; self.rev_transforms[15] = self.reverse_transform_15
        self.fwd_transforms[16] = self.transform_16; self.rev_transforms[16] = self.reverse_transform_16
        self.fwd_transforms[17] = self.transform_17; self.rev_transforms[17] = self.reverse_transform_17
        self.fwd_transforms[18] = self.transform_18; self.rev_transforms[18] = self.reverse_transform_18
        self.fwd_transforms[19] = self.transform_19; self.rev_transforms[19] = self.reverse_transform_19
        self.fwd_transforms[20] = self.transform_20; self.rev_transforms[20] = self.reverse_transform_20
        self.fwd_transforms[21] = self.transform_21; self.rev_transforms[21] = self.reverse_transform_21

        # 22 – Base64
        self.fwd_transforms[22] = self.transform_22
        self.rev_transforms[22] = self.reverse_transform_22

        # 23‑27 (text transforms)
        self.fwd_transforms[23] = self.transform_23; self.rev_transforms[23] = self.reverse_transform_23
        self.fwd_transforms[24] = self.transform_24; self.rev_transforms[24] = self.reverse_transform_24
        self.fwd_transforms[25] = self.transform_25; self.rev_transforms[25] = self.reverse_transform_25
        self.fwd_transforms[26] = self.transform_26; self.rev_transforms[26] = self.reverse_transform_26
        self.fwd_transforms[27] = self.transform_27; self.rev_transforms[27] = self.reverse_transform_27

        # 28 – deterministic per‑3‑byte subtract
        self.fwd_transforms[28] = self.transform_28
        self.rev_transforms[28] = self.reverse_transform_28

        # 29 – global 16‑bit key subtract (Zaden core)
        self.fwd_transforms[29] = lambda d: self.transform_29(d, quantum_boost=False, time_limit=60.0)
        self.rev_transforms[29] = self.reverse_transform_29

        # 30 – global 24‑bit key subtract (heuristic)
        self.fwd_transforms[30] = self.transform_30
        self.rev_transforms[30] = self.reverse_transform_30

        # 31 – .docx paragraph with dictionary
        self.fwd_transforms[31] = self.transform_31
        self.rev_transforms[31] = self.reverse_transform_31

        # 32 – .docx table with dictionary
        self.fwd_transforms[32] = self.transform_32
        self.rev_transforms[32] = self.reverse_transform_32

        # 33‑255 dynamic
        for i in range(33, 256):
            fwd, rev = self._dynamic_transform(i)
            self.fwd_transforms[i] = fwd
            self.rev_transforms[i] = rev

        # 256 no-op
        self.fwd_transforms[256] = self.transform_256
        self.rev_transforms[256] = self.reverse_transform_256

        # Ensure all present
        for i in range(1, 257):
            if i not in self.fwd_transforms:
                raise RuntimeError(f"Transform {i} missing!")

    # ------------------------------------------------------------------
    # Build pair sequences – 2704 (52×52) using only bijective transforms
    # Exclude non‑bijective (1,14,22,23,24,25,26,27,31,32)
    # ------------------------------------------------------------------
    def _build_pair_sequences(self) -> List[Tuple[int, int]]:
        safe = []
        for i in range(1, 257):
            if i in (1, 14, 22, 23, 24, 25, 26, 27, 31, 32):
                continue
            safe.append(i)
            if len(safe) == 52:
                break
        while len(safe) < 52:
            safe.append(256)
        base = safe
        return [(t1, t2) for t1 in base for t2 in base]

    # ------------------------------------------------------------------
    # Apply and reverse sequences
    # ------------------------------------------------------------------
    def _apply_sequence(self, data: bytes, seq: Tuple[int, ...]) -> bytes:
        result = data
        for t_num in seq:
            result = self.fwd_transforms[t_num](result)
        return result

    def _reverse_sequence(self, data: bytes, seq: Tuple[int, ...]) -> bytes:
        result = data
        for t_num in reversed(seq):
            result = self.rev_transforms[t_num](result)
        return result

    # ------------------------------------------------------------------
    # Compression backends (dual mode) – tries zstd first, then paq, then raw
    # ------------------------------------------------------------------
    def _compress_backend(self, data: bytes, safe: bool = False) -> bytes:
        candidates = []
        if HAS_ZSTD:
            try:
                if safe:
                    candidates.append((b'Z', b'Z' + zstd_cctx.compress(data)))
                else:
                    candidates.append((b'Z', zstd_cctx.compress(data)))
            except:
                pass
        if paq is not None:
            try:
                if safe:
                    candidates.append((b'P', b'P' + paq.compress(data)))
                else:
                    candidates.append((b'L', paq.compress(data)))   # marker‑free for paq
            except:
                pass
        candidates.append((b'N', b'N' + data))
        if not candidates:
            return b'N' + data
        if not safe:
            _, best = min(candidates, key=lambda x: len(x[1]))
            return best
        else:
            _, best = min(candidates, key=lambda x: len(x[1]))
            return best

    def _decompress_backend(self, data: bytes, safe: bool = False) -> Optional[bytes]:
        if len(data) == 0:
            return None
        if safe:
            marker = data[0:1]
            payload = data[1:]
            if marker == b'N':
                return payload
            elif marker == b'Z' and HAS_ZSTD:
                try:
                    return zstd_dctx.decompress(payload)
                except:
                    pass
            elif marker == b'P' and paq is not None:
                try:
                    return paq.decompress(payload)
                except:
                    pass
            return None
        # marker‑free (try zstd, then paq, then raw)
        if HAS_ZSTD:
            try:
                return zstd_dctx.decompress(data)
            except:
                pass
        if paq is not None:
            try:
                return paq.decompress(data)
            except:
                pass
        if len(data) > 0 and data[0] == ord('N'):
            return data[1:]
        return None

    # ------------------------------------------------------------------
    # Variable‑length header encoding / decoding
    # ------------------------------------------------------------------
    def _encode_marker_single(self, t: int) -> bytes:
        if t <= 252:
            return bytes([t - 1])
        return bytes([254, t - 253])

    def _encode_marker_raw(self) -> bytes:
        return bytes([252])

    def _encode_marker_pair(self, t1: int, t2: int) -> bytes:
        idx = (t1 - 1) * 52 + (t2 - 1)
        return bytes([253, (idx >> 8) & 0xFF, idx & 0xFF])

    def _decode_header(self, data: bytes):
        if len(data) < 1:
            return 0, ()
        f = data[0]
        if f < 252:
            return 1, (f + 1,)
        elif f == 252:
            return 1, ()
        elif f == 253:
            if len(data) < 3:
                return 0, ()
            idx = (data[1] << 8) | data[2]
            if idx >= len(self.sequences):
                return 0, ()
            t1, t2 = self.pair_lookup[idx]
            return 3, (t1, t2)
        elif f == 254:
            if len(data) < 2:
                return 0, ()
            x = data[1]
            if x > 3:
                return 0, ()
            return 2, (253 + x,)
        else:
            return 0, ()

    # ------------------------------------------------------------------
    # Main compression with auto‑correction – flags for 28, 29, 30
    # ------------------------------------------------------------------
    def compress_with_best(self, data: bytes, safe: bool = False, ultra: bool = True,
                           include_28: bool = False, include_29: bool = False,
                           include_30: bool = False) -> bytes:
        if not data:
            backend = self._compress_backend(b'', safe)
            compressed = self._encode_marker_raw() + backend
            if not safe:
                decomp, _ = self._decompress_auto(compressed)
                if decomp != b'':
                    return self.compress_with_best(data, safe=True, ultra=ultra,
                                                   include_28=include_28, include_29=include_29,
                                                   include_30=include_30)
            return compressed

        best_total = float('inf')
        best_bytes = None

        # Build list of single transforms (1..256) – exclude 28-30 if not allowed
        single_transforms = list(range(1, 257))
        if not include_28:
            single_transforms = [t for t in single_transforms if t != 28]
        if not include_29:
            single_transforms = [t for t in single_transforms if t != 29]
        if not include_30:
            single_transforms = [t for t in single_transforms if t != 30]

        # Add quantum transforms if enabled
        if USE_QUANTUM and HAS_QISKIT:
            fast_quantum = range(257, 266)
            single_transforms.extend(fast_quantum)
            if ultra:
                single_transforms.extend(range(266, 283))

        # Filter pairs: exclude pairs containing disallowed transforms
        allowed_pairs = self.sequences
        if not include_28:
            allowed_pairs = [seq for seq in allowed_pairs if 28 not in seq]
        if not include_29:
            allowed_pairs = [seq for seq in allowed_pairs if 29 not in seq]
        if not include_30:
            allowed_pairs = [seq for seq in allowed_pairs if 30 not in seq]

        # raw
        raw_backend = self._compress_backend(data, safe)
        candidate = self._encode_marker_raw() + raw_backend
        if len(candidate) < best_total:
            best_total = len(candidate)
            best_bytes = candidate

        # singles
        for t in single_transforms:
            try:
                transformed = self.fwd_transforms[t](data)
                backend = self._compress_backend(transformed, safe)
                candidate = self._encode_marker_single(t) + backend
                if len(candidate) < best_total:
                    best_total = len(candidate)
                    best_bytes = candidate
            except:
                continue

        # pairs – only if ultra mode is on (this covers all 2704 pairs)
        if ultra:
            for t1, t2 in allowed_pairs:
                try:
                    transformed = self._apply_sequence(data, (t1, t2))
                    backend = self._compress_backend(transformed, safe)
                    candidate = self._encode_marker_pair(t1, t2) + backend
                    if len(candidate) < best_total:
                        best_total = len(candidate)
                        best_bytes = candidate
                except:
                    continue

        decomp, _ = self._decompress_auto(best_bytes)
        if decomp != data:
            if not safe:
                print("Note: marker‑free mode produced ambiguous stream, falling back to safe markers...")
                return self.compress_with_best(data, safe=True, ultra=ultra,
                                               include_28=include_28, include_29=include_29,
                                               include_30=include_30)
            else:
                raise RuntimeError(f"Safe compression failed – unexpected internal error! "
                                   f"(input len={len(data)}, output len={len(best_bytes)})")
        return best_bytes

    def _decompress_auto(self, data: bytes) -> Tuple[bytes, Optional[Tuple[int, ...]]]:
        offset, seq = self._decode_header(data)
        if offset == 0:
            return b'', None
        payload = data[offset:]
        if not payload:
            return b'', None

        first_byte = payload[0:1]
        if first_byte in (b'N', b'Z', b'P'):
            res = self._decompress_backend(payload, safe=True)
        else:
            res = self._decompress_backend(payload, safe=False)
        if res is None:
            return b'', None

        try:
            if not seq:
                result = res
            else:
                result = self._reverse_sequence(res, seq)
        except:
            return b'', None
        return result, seq

    # ------------------------------------------------------------------
    # Dictionary compression helpers (for hybrid mode)
    # ------------------------------------------------------------------
    MAGIC_DICT = b'DICT'
    MAGIC_LINE = b'LINE'

    def _tokenize_with_static_dict(self, data: bytes) -> Optional[bytes]:
        try:
            text = data.decode('utf-8')
        except:
            return None
        pattern = r'([A-Za-z0-9_]+)'
        tokens = re.split(pattern, text)
        stream = bytearray()
        for i, tok in enumerate(tokens):
            if i % 2 == 1:
                idx = self.word_to_index.get(tok)
                if idx is not None:
                    stream += b'\x01'
                    stream += struct.pack('>I', idx)
                else:
                    word_bytes = tok.encode('utf-8')
                    stream += b'\x02'
                    stream += struct.pack('>H', len(word_bytes))
                    stream += word_bytes
            else:
                if tok:
                    sep_bytes = tok.encode('utf-8')
                    stream += b'\x00'
                    stream += struct.pack('>H', len(sep_bytes))
                    stream += sep_bytes
        return bytes(stream)

    def _detokenize_static_dict(self, token_stream: bytes) -> Optional[bytes]:
        if not token_stream:
            return b''
        out = bytearray()
        pos = 0
        while pos < len(token_stream):
            if pos >= len(token_stream):
                break
            typ = token_stream[pos]
            pos += 1
            if typ == 0x01:
                if pos + 4 > len(token_stream):
                    break
                idx = struct.unpack('>I', token_stream[pos:pos+4])[0]
                pos += 4
                if idx < len(self.static_dict):
                    out += self.static_dict[idx].encode('utf-8')
                else:
                    return None
            elif typ == 0x02:
                if pos + 2 > len(token_stream):
                    break
                word_len = struct.unpack('>H', token_stream[pos:pos+2])[0]
                pos += 2
                if pos + word_len > len(token_stream):
                    break
                out += token_stream[pos:pos+word_len]
                pos += word_len
            elif typ == 0x00:
                if pos + 2 > len(token_stream):
                    break
                sep_len = struct.unpack('>H', token_stream[pos:pos+2])[0]
                pos += 2
                if pos + sep_len > len(token_stream):
                    break
                out += token_stream[pos:pos+sep_len]
                pos += sep_len
            else:
                break
        return bytes(out)

    def _compress_static_dict(self, data: bytes) -> Optional[bytes]:
        token_stream = self._tokenize_with_static_dict(data)
        if token_stream is None:
            return None
        compressed = self._compress_backend(token_stream, safe=True)
        return self.MAGIC_DICT + b'\x01' + compressed

    def _decompress_static_dict(self, compressed: bytes) -> Optional[bytes]:
        if not compressed.startswith(self.MAGIC_DICT + b'\x01'):
            return None
        payload = compressed[len(self.MAGIC_DICT) + 1:]
        token_stream = self._decompress_backend(payload, safe=True)
        if token_stream is None:
            return None
        return self._detokenize_static_dict(token_stream)

    def _compress_dynamic_dict(self, data: bytes) -> Optional[bytes]:
        try:
            token_stream = self.transform_25(data)
        except:
            return None
        compressed = self._compress_backend(token_stream, safe=True)
        return self.MAGIC_DICT + b'\x02' + compressed

    def _decompress_dynamic_dict(self, compressed: bytes) -> Optional[bytes]:
        if not compressed.startswith(self.MAGIC_DICT + b'\x02'):
            return None
        payload = compressed[len(self.MAGIC_DICT) + 1:]
        token_stream = self._decompress_backend(payload, safe=True)
        if token_stream is None:
            return None
        return self.reverse_transform_25(token_stream)

    def _tokenize_with_line_dict(self, data: bytes) -> Optional[bytes]:
        if not self.line_dict:
            return None
        try:
            text = data.decode('utf-8')
        except:
            return None

        pos = 0
        token_list = []
        while pos < len(text):
            earliest_pos = len(text) + 1
            earliest_len = 0
            earliest_idx = -1
            for idx, phrase in enumerate(self.line_dict):
                p = text.find(phrase, pos)
                if p != -1 and (p < earliest_pos or (p == earliest_pos and len(phrase) > earliest_len)):
                    earliest_pos = p
                    earliest_len = len(phrase)
                    earliest_idx = idx
            if earliest_idx != -1:
                if earliest_pos > pos:
                    token_list.append((False, text[pos:earliest_pos].encode('utf-8')))
                token_list.append((True, earliest_idx))
                pos = earliest_pos + earliest_len
            else:
                token_list.append((False, text[pos:].encode('utf-8')))
                break

        out = bytearray()
        for is_index, payload in token_list:
            if is_index:
                out += b'\x01'
                out += struct.pack('>Q', payload)
            else:
                raw_bytes = payload
                out += b'\x00'
                out += struct.pack('>H', len(raw_bytes))
                out += raw_bytes
        return bytes(out)

    def _detokenize_line_dict(self, token_stream: bytes) -> Optional[bytes]:
        if not token_stream:
            return b''
        out = bytearray()
        pos = 0
        while pos < len(token_stream):
            if pos >= len(token_stream):
                break
            typ = token_stream[pos]
            pos += 1
            if typ == 1:
                if pos + 8 > len(token_stream):
                    return None
                idx = struct.unpack('>Q', token_stream[pos:pos+8])[0]
                pos += 8
                if idx < len(self.line_dict):
                    out += self.line_dict[idx].encode('utf-8')
                else:
                    return None
            elif typ == 0:
                if pos + 2 > len(token_stream):
                    return None
                raw_len = struct.unpack('>H', token_stream[pos:pos+2])[0]
                pos += 2
                if pos + raw_len > len(token_stream):
                    return None
                out += token_stream[pos:pos+raw_len]
                pos += raw_len
            else:
                return None
        return bytes(out)

    def _compress_line_dict(self, data: bytes) -> Optional[bytes]:
        token_stream = self._tokenize_with_line_dict(data)
        if token_stream is None:
            return None
        compressed = self._compress_backend(token_stream, safe=True)
        return self.MAGIC_LINE + compressed

    def _decompress_line_dict(self, compressed: bytes) -> Optional[bytes]:
        if not compressed.startswith(self.MAGIC_LINE):
            return None
        payload = compressed[len(self.MAGIC_LINE):]
        token_stream = self._decompress_backend(payload, safe=True)
        if token_stream is None:
            return None
        return self._detokenize_line_dict(token_stream)

    # ------------------------------------------------------------------
    # Zaden Block Optimization with time‑limited search and variable‑length key coding
    # ------------------------------------------------------------------
    ZADEN_MAGIC = 0x33  # single byte header

    def _encode_key_unary(self, key: int) -> bytes:
        if key == 0:
            bits = '0'
            length = 1
        else:
            bits = bin(key)[2:]
            length = len(bits)
        prefix = '0' * (length - 1) + '1'
        encoded_str = prefix + bits
        pad = (8 - len(encoded_str) % 8) % 8
        encoded_str += '0' * pad
        out = bytearray()
        for i in range(0, len(encoded_str), 8):
            out.append(int(encoded_str[i:i+8], 2))
        return bytes(out)

    def _decode_key_unary(self, data: bytes, pos: int) -> Tuple[int, int]:
        bit_idx = pos * 8
        zeros = 0
        while True:
            byte_idx = bit_idx // 8
            bit_off = bit_idx % 8
            if byte_idx >= len(data):
                raise ValueError("End of data while decoding unary key")
            byte = data[byte_idx]
            bit = (byte >> (7 - bit_off)) & 1
            bit_idx += 1
            if bit == 1:
                break
            zeros += 1
        length = zeros + 1
        key = 0
        for _ in range(length):
            byte_idx = bit_idx // 8
            bit_off = bit_idx % 8
            if byte_idx >= len(data):
                raise ValueError("Unexpected end of data while reading key bits")
            byte = data[byte_idx]
            bit = (byte >> (7 - bit_off)) & 1
            key = (key << 1) | bit
            bit_idx += 1
        bit_idx = ((bit_idx + 7) // 8) * 8
        new_pos = bit_idx // 8
        return key, new_pos

    def _find_best_two_pass_keys(self, block: bytes, quantum_boost: bool = False, time_limit: float = 60.0) -> Tuple[int, int, bytes]:
        if len(block) < 3:
            return 0, 0, block

        time_per_pass = max(1.0, time_limit / 2)
        key1 = self._find_best_16bit_key(block, quantum_boost, time_per_pass)
        pad_len1 = (3 - len(block) % 3) % 3
        padded1 = block + b'\x00' * pad_len1
        trans1 = bytearray()
        for i in range(0, len(padded1), 3):
            v = int.from_bytes(padded1[i:i+3], 'little')
            new_v = (v - key1) & 0xFFFFFF
            trans1.extend(new_v.to_bytes(3, 'little'))
        if pad_len1:
            trans1 = trans1[:-pad_len1]
        inter = bytes(trans1)

        key2 = self._find_best_16bit_key(inter, quantum_boost, time_per_pass)
        pad_len2 = (3 - len(inter) % 3) % 3
        padded2 = inter + b'\x00' * pad_len2
        trans2 = bytearray()
        for i in range(0, len(padded2), 3):
            v = int.from_bytes(padded2[i:i+3], 'little')
            new_v = (v - key2) & 0xFFFFFF
            trans2.extend(new_v.to_bytes(3, 'little'))
        if pad_len2:
            trans2 = trans2[:-pad_len2]
        final = bytes(trans2)
        return key1, key2, final

    def _block_optimize(self, data: bytes, block_size: int = 256, quantum_boost: bool = False, time_limit: float = 60.0) -> Tuple[bytes, List[Tuple[int, int]]]:
        keys = []
        transformed_parts = []
        total_blocks = (len(data) + block_size - 1) // block_size
        for idx, i in enumerate(range(0, len(data), block_size)):
            block = data[i:i+block_size]
            print(f"Processing block {idx+1}/{total_blocks} (time limit: {time_limit:.1f}s)...")
            k1, k2, new_block = self._find_best_two_pass_keys(block, quantum_boost, time_limit)
            keys.append((k1, k2))
            transformed_parts.append(new_block)
        return b''.join(transformed_parts), keys

    # ------------------------------------------------------------------
    # Internal Zaden round‑trip test for Options 5 and 7
    # ------------------------------------------------------------------
    def _test_zaden_roundtrip(self, data: bytes) -> bool:
        try:
            block_size = 256
            quantum_boost = False
            time_limit = 5.0
            transformed_data, keys = self._block_optimize(data, block_size, quantum_boost, time_limit)
            inner_compressed = self.compress_with_best(transformed_data, safe=False, ultra=True,
                                                       include_28=True, include_29=True, include_30=True)
            magic = bytes([self.ZADEN_MAGIC])
            num_blocks = len(keys)
            header = struct.pack('<II', block_size, num_blocks)
            key_bytes = b''.join(self._encode_key_unary(k1) + self._encode_key_unary(k2) for k1, k2 in keys)
            compressed = magic + header + key_bytes + inner_compressed
            decompressed = self.decompress_block_optimized(compressed)
            return decompressed == data
        except Exception:
            return False

    # ------------------------------------------------------------------
    # Helper: compress with hybrid (Option 4 style) and return bytes
    # ------------------------------------------------------------------
    def _compress_hybrid_bytes(self, data: bytes) -> Tuple[bytes, str]:
        """Compress data using hybrid dict + Ultra + all transforms, return (best_bytes, method_name)."""
        candidates = []
        c_static = self._compress_static_dict(data)
        if c_static is not None:
            candidates.append(('Static-Word-Dict', c_static))
        c_line = self._compress_line_dict(data)
        if c_line is not None:
            candidates.append(('Line-Dict', c_line))
        c_dynamic = self._compress_dynamic_dict(data)
        if c_dynamic is not None:
            candidates.append(('Dynamic-Dict', c_dynamic))
        c_pjp = self.compress_with_best(data, safe=False, ultra=True,
                                        include_28=True, include_29=True, include_30=True)
        candidates.append(('PJP-Absolute', c_pjp))
        best_method, best_bytes = min(candidates, key=lambda x: len(x[1]))
        return best_bytes, best_method

    # ------------------------------------------------------------------
    # Algorithm 36 – 24‑bit pass search (powers of two + smart candidates)
    # ------------------------------------------------------------------
    ALGO36_MAGIC = 0x36

    def algorithm_36_compress(self, data: bytes, time_limit: float = 300.0) -> bytes:
        if len(data) == 0:
            return bytes([self.ALGO36_MAGIC, 0, 0])

        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        values = []
        for i in range(0, len(padded), 3):
            values.append(int.from_bytes(padded[i:i+3], 'little'))
        n = len(values)
        if n == 0:
            return bytes([self.ALGO36_MAGIC, pad_len, 0]) + self._compress_backend(b'', safe=False)

        candidates = set()
        for i in range(24):
            candidates.add(1 << i)
        if len(data) <= 40:
            mean = sum(values) // n
            candidates.add(mean)
            sorted_vals = sorted(values)
            median = sorted_vals[n//2]
            candidates.add(median)
            for v in values:
                candidates.add(v)
            for offset in [-10, -1, 1, 10]:
                candidates.add((mean + offset) % (1 << 24))

        best_idx = 0
        best_pass = 1
        best_metric = float('inf')
        start_time = time.time()

        for p in candidates:
            trans_vals = [((v - p) & 0xFFFFFF) for v in values]
            mean_t = sum(trans_vals) // n
            metric = sum(abs(t - mean_t) for t in trans_vals)
            if metric < best_metric:
                best_metric = metric
                best_pass = p
            if time.time() - start_time > time_limit:
                break

        best_idx = 0
        best_metric = float('inf')
        for idx, p in enumerate([1 << i for i in range(24)]):
            trans_vals = [((v - p) & 0xFFFFFF) for v in values]
            mean_t = sum(trans_vals) // n
            metric = sum(abs(t - mean_t) for t in trans_vals)
            if metric < best_metric:
                best_metric = metric
                best_idx = idx
            if time.time() - start_time > time_limit:
                break

        best_pass = 1 << best_idx
        transformed = bytearray()
        for v in values:
            new_v = (v - best_pass) & 0xFFFFFF
            transformed.extend(new_v.to_bytes(3, 'little'))

        compressed = self.compress_with_best(bytes(transformed), safe=False, ultra=True,
                                             include_28=True, include_29=True, include_30=True)
        out = bytearray([self.ALGO36_MAGIC, pad_len, best_idx])
        out.extend(compressed)
        return bytes(out)

    def algorithm_36_decompress(self, data: bytes) -> Optional[bytes]:
        if not data or data[0] != self.ALGO36_MAGIC:
            return None
        pos = 1
        if len(data) < pos + 2:
            return None
        pad_len = data[pos]; pos += 1
        idx = data[pos]; pos += 1
        if idx > 23:
            return None
        pass_val = 1 << idx
        compressed = data[pos:]

        orig_data, _ = self._decompress_auto(compressed)
        if orig_data is None:
            return None
        if len(orig_data) % 3 != 0:
            return None

        out = bytearray()
        for i in range(0, len(orig_data), 3):
            chunk = orig_data[i:i+3]
            val = int.from_bytes(chunk, 'little')
            orig_val = (val + pass_val) & 0xFFFFFF
            out.extend(orig_val.to_bytes(3, 'little'))
        if pad_len > 0:
            out = out[:-pad_len]
        return bytes(out)

    # ------------------------------------------------------------------
    # Algorithm 37 – AI Archive, 256‑byte blocks, 300s per block
    # ------------------------------------------------------------------
    ALGO37_MAGIC = 0x37

    def algorithm_37_compress(self, data: bytes, block_size: int = 256, time_limit_per_block: float = 300.0) -> bytes:
        if len(data) == 0:
            return bytes([self.ALGO37_MAGIC, 0, 0, 0, 0, 0])

        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len

        blocks = [padded[i:i+block_size] for i in range(0, len(padded), block_size)]
        num_blocks = len(blocks)

        keys = []
        transformed_blocks = []
        for idx, blk in enumerate(blocks):
            print(f"Algorithm 37: block {idx+1}/{num_blocks} ({len(blk)} bytes) ...")
            block_pad = (3 - len(blk) % 3) % 3
            blk_padded = blk + b'\x00' * block_pad
            values = [int.from_bytes(blk_padded[i:i+3], 'little') for i in range(0, len(blk_padded), 3)]

            best_key = 0
            best_metric = float('inf')
            start = time.time()
            candidates = {1 << i for i in range(24)}
            rng = random.Random(42 + idx)
            evaluated = set()

            while time.time() - start < time_limit_per_block:
                for _ in range(100):
                    candidates.add(rng.randint(0, (1 << 24) - 1))
                for c in candidates:
                    if c in evaluated:
                        continue
                    evaluated.add(c)
                    trans = [((v - c) & 0xFFFFFF) for v in values]
                    mean_t = sum(trans) // len(trans)
                    metric = sum(abs(t - mean_t) for t in trans)
                    if metric < best_metric:
                        best_metric = metric
                        best_key = c
                    if time.time() - start >= time_limit_per_block:
                        break
                if time.time() - start >= time_limit_per_block:
                    break

            keys.append(best_key)
            transformed = bytearray()
            for v in values:
                new_v = (v - best_key) & 0xFFFFFF
                transformed.extend(new_v.to_bytes(3, 'little'))
            transformed_blocks.append(bytes(transformed))

        transformed_data = b''.join(transformed_blocks)

        # PJP inner compression – ultra=True uses ALL 2704 transform pairs
        compressed_inner = self.compress_with_best(transformed_data, safe=False, ultra=True,
                                                   include_28=True, include_29=True, include_30=True)

        # PRNG XOR whitening
        seed = random.randint(0, 65535)
        rng = random.Random(seed)
        stream = bytes(rng.getrandbits(8) for _ in range(len(compressed_inner)))
        whitened_payload = bytes(a ^ b for a, b in zip(compressed_inner, stream))

        out = bytearray([self.ALGO37_MAGIC, pad_len])
        out += struct.pack('<I', num_blocks)
        for key, tr_block in zip(keys, transformed_blocks):
            out += struct.pack('<H', len(tr_block))
            out += key.to_bytes(3, 'little')
        out += struct.pack('<H', seed)
        out += whitened_payload
        return bytes(out)

    def algorithm_37_decompress(self, data: bytes) -> Optional[bytes]:
        if not data or data[0] != self.ALGO37_MAGIC:
            return None
        pos = 1
        if len(data) < pos + 5:
            return None
        pad_len = data[pos]; pos += 1
        num_blocks = struct.unpack('<I', data[pos:pos+4])[0]; pos += 4
        keys = []
        transformed_lens = []
        for _ in range(num_blocks):
            if pos + 5 > len(data):
                return None
            tr_len = struct.unpack('<H', data[pos:pos+2])[0]
            pos += 2
            key = int.from_bytes(data[pos:pos+3], 'little')
            pos += 3
            transformed_lens.append(tr_len)
            keys.append(key)

        if pos + 2 > len(data):
            return None
        seed = struct.unpack('<H', data[pos:pos+2])[0]
        pos += 2
        payload = data[pos:]

        rng = random.Random(seed)
        stream = bytes(rng.getrandbits(8) for _ in range(len(payload)))
        compressed = bytes(a ^ b for a, b in zip(payload, stream))

        inner, seq = self._decompress_auto(compressed)
        if inner is None:
            return None

        total_tr_len = sum(transformed_lens)
        if len(inner) != total_tr_len:
            return None

        out = bytearray()
        offset = 0
        for tr_len, key in zip(transformed_lens, keys):
            block_data = inner[offset:offset+tr_len]
            offset += tr_len
            if len(block_data) % 3 != 0:
                return None
            for i in range(0, len(block_data), 3):
                chunk = block_data[i:i+3]
                val = int.from_bytes(chunk, 'little')
                orig_val = (val + key) & 0xFFFFFF
                out.extend(orig_val.to_bytes(3, 'little'))
        if pad_len > 0:
            out = out[:-pad_len]
        return bytes(out)

    # ------------------------------------------------------------------
    # Option 9 – Enhanced: tries both Algorithm 37 and Absolute (hybrid+Ultra)
    #            Asks for input/output filenames, saves the smallest file.
    # ------------------------------------------------------------------
    def compress_ai_zaden(self):
        infile = input("Input file: ").strip()
        if not os.path.exists(infile):
            print(f"Error: file '{infile}' not found.")
            return
        outfile = input("Output file (default: AI_Zaden): ").strip()
        if not outfile:
            outfile = "AI_Zaden"

        try:
            with open(infile, 'rb') as f:
                data = f.read()
        except Exception as e:
            print(f"Error reading {infile}: {e}")
            return

        print("\nRunning Algorithm 37 (AI_Zaden) ...")
        t0 = time.time()
        algo37 = self.algorithm_37_compress(data, block_size=256, time_limit_per_block=300.0)
        t1 = time.time()
        print(f"Algorithm 37: {len(data)} → {len(algo37)} bytes in {t1-t0:.1f}s")

        print("\nRunning Absolute (Hybrid + all 2704 pairs) ...")
        t0 = time.time()
        abs_bytes, abs_method = self._compress_hybrid_bytes(data)
        t1 = time.time()
        print(f"Absolute ({abs_method}): {len(data)} → {len(abs_bytes)} bytes in {t1-t0:.1f}s")

        if len(algo37) < len(abs_bytes):
            best_bytes = algo37
            best_method = "Algorithm 37"
        else:
            best_bytes = abs_bytes
            best_method = f"Absolute ({abs_method})"

        with open(outfile, 'wb') as f:
            f.write(best_bytes)
        print(f"\nSaved smallest file: {len(data)} → {len(best_bytes)} bytes ({best_method}) → {outfile}")

    # ------------------------------------------------------------------
    # Standard file compression (used by options 1-4,8)
    # ------------------------------------------------------------------
    def compress_file(self, infile: str, outfile: str, ultra: bool = True, hybrid: bool = False,
                      include_28: bool = False, include_29: bool = False,
                      include_30: bool = False):
        try:
            with open(infile, 'rb') as f:
                data = f.read()
        except Exception as e:
            print(f"Error reading file: {e}")
            return

        if hybrid:
            best_bytes, method = self._compress_hybrid_bytes(data)
        else:
            candidates = []
            c_pjp = self.compress_with_best(data, safe=False, ultra=ultra,
                                            include_28=include_28, include_29=include_29,
                                            include_30=include_30)
            candidates.append(('PJP', c_pjp))
            best_method, best_bytes = min(candidates, key=lambda x: len(x[1]))
            method = best_method

        try:
            with open(outfile, 'wb') as f:
                f.write(best_bytes)
        except Exception as e:
            print(f"Error writing output file: {e}")
            return
        print(f"Compressed {len(data)} → {len(best_bytes)} bytes ({method}) → {outfile}")

    # ------------------------------------------------------------------
    # Decompression (handles standard PJP, Zaden, Algorithm 36, and Algorithm 37)
    # ------------------------------------------------------------------
    def decompress_file(self, infile: str, outfile: str):
        try:
            with open(infile, 'rb') as f:
                data = f.read()
        except Exception as e:
            print(f"Error reading file: {e}")
            return

        if len(data) > 0 and data[0] == self.ALGO37_MAGIC:
            original = self.algorithm_37_decompress(data)
            if original is not None:
                with open(outfile, 'wb') as f:
                    f.write(original)
                print(f"Decompressed (Algorithm 37) → {outfile} ({len(original)} bytes)")
                return
            else:
                print("Decompression failed for Algorithm 37 data.")
                return

        if len(data) > 0 and data[0] == self.ALGO36_MAGIC:
            original = self.algorithm_36_decompress(data)
            if original is not None:
                with open(outfile, 'wb') as f:
                    f.write(original)
                print(f"Decompressed (Algorithm 36) → {outfile} ({len(original)} bytes)")
                return
            else:
                print("Decompression failed for Algorithm 36 data.")
                return

        if len(data) > 0 and data[0] == self.ZADEN_MAGIC:
            original = self.decompress_block_optimized(data)
            if original is not None:
                with open(outfile, 'wb') as f:
                    f.write(original)
                print(f"Decompressed (Zaden) → {outfile} ({len(original)} bytes)")
                return
            else:
                print("Decompression failed for Zaden data.")
                return

        if data.startswith(self.MAGIC_LINE):
            original = self._decompress_line_dict(data)
            if original is not None:
                with open(outfile, 'wb') as f:
                    f.write(original)
                print(f"Decompressed (Line-Dict) → {outfile} ({len(original)} bytes)")
                return
        if data.startswith(self.MAGIC_DICT + b'\x01'):
            original = self._decompress_static_dict(data)
            if original is not None:
                with open(outfile, 'wb') as f:
                    f.write(original)
                print(f"Decompressed (Static-Word-Dict) → {outfile} ({len(original)} bytes)")
                return
        if data.startswith(self.MAGIC_DICT + b'\x02'):
            original = self._decompress_dynamic_dict(data)
            if original is not None:
                with open(outfile, 'wb') as f:
                    f.write(original)
                print(f"Decompressed (Dynamic-Dict) → {outfile} ({len(original)} bytes)")
                return

        original, seq = self._decompress_auto(data)
        if original == b'' and seq is None:
            print("Decompression failed – unknown format.")
            return
        with open(outfile, 'wb') as f:
            f.write(original)
        seq_str = "raw" if not seq else f"sequence {seq}"
        print(f"Decompressed ({seq_str}) → {outfile} ({len(original)} bytes)")

    def decompress_block_optimized(self, data: bytes) -> Optional[bytes]:
        if len(data) == 0 or data[0] != self.ZADEN_MAGIC:
            return None
        pos = 1
        if len(data) < pos + 8:
            return None
        block_size, num_blocks = struct.unpack('<II', data[pos:pos+8])
        pos += 8

        keys = []
        for _ in range(num_blocks):
            k1, pos = self._decode_key_unary(data, pos)
            k2, pos = self._decode_key_unary(data, pos)
            keys.append((k1, k2))

        inner_compressed = data[pos:]
        inner_data, seq = self._decompress_auto(inner_compressed)
        if inner_data is None:
            return None

        out_parts = []
        offset = 0
        for k1, k2 in keys:
            block = inner_data[offset:offset+block_size]
            pad_len2 = (3 - len(block) % 3) % 3
            if pad_len2:
                block += b'\x00' * pad_len2
            inter = bytearray()
            for i in range(0, len(block), 3):
                v = int.from_bytes(block[i:i+3], 'little')
                orig = (v + k2) & 0xFFFFFF
                inter.extend(orig.to_bytes(3, 'little'))
            if pad_len2:
                inter = inter[:-pad_len2]
            inter_block = bytes(inter)
            pad_len1 = (3 - len(inter_block) % 3) % 3
            if pad_len1:
                inter_block += b'\x00' * pad_len1
            out_block = bytearray()
            for i in range(0, len(inter_block), 3):
                v = int.from_bytes(inter_block[i:i+3], 'little')
                orig = (v + k1) & 0xFFFFFF
                out_block.extend(orig.to_bytes(3, 'little'))
            if pad_len1:
                out_block = out_block[:-pad_len1]
            out_parts.append(bytes(out_block))
            offset += block_size
        return b''.join(out_parts)

    # ------------------------------------------------------------------
    # Verify transforms (quick check on single byte)
    # ------------------------------------------------------------------
    def verify_transforms(self) -> bool:
        print("Verifying all 256+ transforms...")
        ok = True
        for t in range(1, 257):
            test = bytes([0x55])
            try:
                enc = self.fwd_transforms[t](test)
                dec = self.rev_transforms[t](enc)
                if dec == test:
                    print(f"Transform {t}: right")
                else:
                    print(f"Transform {t}: incorrect")
                    ok = False
            except Exception:
                print(f"Transform {t}: exception")
                ok = False
        if USE_QUANTUM and HAS_QISKIT:
            for t in range(257, 283):
                test = bytes([0x55])
                try:
                    enc = self.fwd_transforms[t](test)
                    dec = self.rev_transforms[t](enc)
                    if dec == test:
                        print(f"Quantum transform {t}: right")
                    else:
                        print(f"Quantum transform {t}: incorrect")
                        ok = False
                except Exception:
                    print(f"Quantum transform {t}: exception")
                    ok = False
        print("Verification complete.\n")
        return ok

    # ------------------------------------------------------------------
    # Full self‑test (exhaustive) – now includes Zaden and Algorithm 36, 37
    # ------------------------------------------------------------------
    def full_self_test(self) -> bool:
        print("=" * 60)
        print("PJP – FULL SELF‑TEST (100% lossless)")
        print("=" * 60)
        all_ok = True
        rng = random.Random(12345)

        print("Testing all single transforms on all 256 byte values...")
        for t_num in range(1, 257):
            for b in range(256):
                orig = bytes([b])
                try:
                    enc = self.fwd_transforms[t_num](orig)
                    dec = self.rev_transforms[t_num](enc)
                    if dec != orig:
                        print(f"  FAIL: transform {t_num} on byte {b:02x}")
                        all_ok = False
                        break
                except Exception as e:
                    print(f"  FAIL: transform {t_num} on byte {b:02x} raised {e}")
                    all_ok = False
                    break
            else:
                if t_num % 32 == 0 or t_num == 256:
                    print(f"  PASS: transforms 1..{t_num} OK on all bytes")
            if not all_ok:
                break
        if not all_ok:
            return False

        if USE_QUANTUM and HAS_QISKIT:
            print("Testing quantum transforms on all 256 byte values...")
            for t_num in range(257, 283):
                for b in range(256):
                    orig = bytes([b])
                    try:
                        enc = self.fwd_transforms[t_num](orig)
                        dec = self.rev_transforms[t_num](enc)
                        if dec != orig:
                            print(f"  FAIL: quantum transform {t_num} on byte {b:02x}")
                            all_ok = False
                            break
                    except Exception as e:
                        print(f"  FAIL: quantum transform {t_num} on byte {b:02x} raised {e}")
                        all_ok = False
                        break
                else:
                    if (t_num-256) % 8 == 0:
                        print(f"  PASS: quantum transforms 257..{t_num} OK on all bytes")
                if not all_ok:
                    break
            if not all_ok:
                return False

        print(f"\nTesting all {len(self.sequences)} transform pairs on all 256 byte values...")
        for idx, seq in enumerate(self.sequences):
            for b in range(256):
                orig = bytes([b])
                try:
                    enc = self._apply_sequence(orig, seq)
                    dec = self._reverse_sequence(enc, seq)
                    if dec != orig:
                        print(f"  FAIL: pair {seq} on byte {b:02x}")
                        all_ok = False
                        break
                except Exception as e:
                    print(f"  FAIL: pair {seq} on byte {b:02x} raised {e}")
                    all_ok = False
                    break
            if not all_ok:
                break
            if (idx + 1) % 256 == 0:
                print(f"  PASS: {idx + 1} pairs tested on all bytes")
        if not all_ok:
            return False

        print("\nTesting random 1000‑byte block through full compress/decompress...")
        test_data = bytes(rng.randint(0, 255) for _ in range(1000))
        for safe in [False, True]:
            compressed = self.compress_with_best(test_data, safe=safe, ultra=True,
                                                 include_28=True, include_29=True,
                                                 include_30=True)
            decompressed, _ = self._decompress_auto(compressed)
            if decompressed != test_data:
                print(f"  FAIL: random data pipeline mismatch (safe={safe})")
                return False

        print("\nTesting empty input...")
        for safe in [False, True]:
            compressed_empty = self.compress_with_best(b'', safe, include_28=True, include_29=True,
                                                       include_30=True)
            decomp_empty, _ = self._decompress_auto(compressed_empty)
            if decomp_empty != b'':
                print(f"  FAIL: empty input pipeline mismatch (safe={safe})")
                return False

        print("\nTesting static word dictionary tokenizer...")
        sample = b"The quick brown fox jumps over the lazy dog. 12345 not in dict."
        token = self._tokenize_with_static_dict(sample)
        if token is None:
            return False
        reconstructed = self._detokenize_static_dict(token)
        if reconstructed != sample:
            return False

        if self.line_dict:
            sample_line = b"This is a test. the quick brown fox jumps over the lazy dog."
            token_line = self._tokenize_with_line_dict(sample_line)
            if token_line is None:
                return False
            reconstructed_line = self._detokenize_line_dict(token_line)
            if reconstructed_line != sample_line and (reconstructed_line is None or len(reconstructed_line) != len(sample_line)):
                return False

        print("\nTesting dynamic dictionary tokenizer...")
        sample2 = b"Hello world! This is a test. Hello world again."
        encoded = self.transform_25(sample2)
        decoded = self.reverse_transform_25(encoded)
        if decoded != sample2:
            return False

        print("\nTesting 6‑bit transform...")
        sample_text = b"Hello world! How are you?\nThis is a test."
        enc27 = self.transform_27(sample_text)
        dec27 = self.reverse_transform_27(enc27)
        if dec27 != sample_text:
            all_ok = False

        print("\nTesting transforms 28‑30...")
        for i, func in enumerate([self.transform_28, self.transform_29, self.transform_30], start=28):
            test = bytes(rng.randint(0, 255) for _ in range(100))
            enc = func(test)
            if i == 28:
                dec = self.reverse_transform_28(enc)
            elif i == 29:
                dec = self.reverse_transform_29(enc)
            else:
                dec = self.reverse_transform_30(enc)
            if dec != test:
                print(f"  FAIL: transform {i} round‑trip mismatch")
                all_ok = False

        # .docx transforms
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            p = doc.add_paragraph("Hello World! ")
            r = p.add_run("This is bold.")
            r.bold = True
            r.font.size = Pt(14)
            p.add_run(" Normal text.")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0,0).text = "Cell 1,1"
            table.cell(0,1).text = "Cell 1,2"
            table.cell(1,0).text = "Cell 2,1"
            table.cell(1,1).text = "Cell 2,2"
            bio = io.BytesIO()
            doc.save(bio)
            docx_bytes = bio.getvalue()

            enc31 = self.transform_31(docx_bytes)
            dec31 = self.reverse_transform_31(enc31)
            doc31 = Document(io.BytesIO(dec31))
            if "Hello World!" not in doc31.paragraphs[0].text:
                all_ok = False

            enc32 = self.transform_32(docx_bytes)
            dec32 = self.reverse_transform_32(enc32)
            doc32 = Document(io.BytesIO(dec32))
            if len(doc32.tables) == 0 or doc32.tables[0].cell(0,0).text != "Cell 1,1":
                all_ok = False
        except ImportError:
            print("  SKIP: python-docx not installed.")

        # Zaden test
        print("\nTesting Zaden round‑trip...")
        if not self._test_zaden_roundtrip(bytes(rng.randint(0, 255) for _ in range(512))):
            all_ok = False

        # Algo 36 test
        print("Testing Algorithm 36 round‑trip...")
        test36 = bytes(rng.randint(0, 255) for _ in range(512))
        if self.algorithm_36_decompress(self.algorithm_36_compress(test36, time_limit=5.0)) != test36:
            all_ok = False

        # Algo 37 test
        print("Testing Algorithm 37 round‑trip...")
        test37 = bytes(rng.randint(0, 255) for _ in range(512))
        if self.algorithm_37_decompress(self.algorithm_37_compress(test37, block_size=256, time_limit_per_block=10.0)) != test37:
            all_ok = False

        if all_ok:
            print("\n[All tests passed – compressor is 100% lossless]")
        else:
            print("\n[FAIL] Some tests failed.")
        return all_ok

    # ------------------------------------------------------------------
    # Transform 256 – no-op
    # ------------------------------------------------------------------
    def transform_256(self, d: bytes) -> bytes:
        return d
    reverse_transform_256 = transform_256

# ------------------------------------------------------------
# Main (with persistent menu loop and strict numeric input)
# ------------------------------------------------------------
def get_menu_choice():
    while True:
        try:
            choice = int(input("> ").strip())
            if 0 <= choice <= 9:
                return choice
            else:
                print("Please enter a number from 0 to 9.")
        except ValueError:
            print("Invalid input. Please enter a number (0-9).")

def main():
    print(f"{PROGNAME} – 256 transforms + 2704 pairs + Base64 + 6‑bit text + Quantum + Transforms 28–30 + .docx transforms 31–32")
    print("Option 9: Algorithm 37 + Absolute (Hybrid) – tries both, picks smallest. Asks for I/O.")
    if paq is None and not HAS_ZSTD:
        print("Warning: No compression backend found. Dictionary streams will be stored raw.")

    c = PJPCompressor()
    c.verify_transforms()

    while True:
        print("\n" + "="*50)
        print("Menu:")
        print("1) Fast (no 28-30) – 256 singles")
        print("2) Ultra (no 28-30) – 256 singles + 2704 pairs")
        print("3) Hybrid (no 28-30) – dicts + Ultra")
        print("4) Absolute (with 28, 29, 30) – all transforms")
        print("5) Full self‑test (includes Zaden, Algo 36 & 37)")
        print("6) Decompress (extract)")
        print("7) Test 2704 pairs & extraction check")
        print("8) Fast 256 transforms test")
        print("9) AI Archive – Algorithm 37 + Absolute (user‑specified I/O)")
        print("0) Exit")
        print("="*50)

        choice = get_menu_choice()

        if choice == 0:
            print("Exiting program. Goodbye!")
            break
        elif choice == 1:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            c.compress_file(i, o, ultra=False, hybrid=False,
                            include_28=False, include_29=False, include_30=False)
        elif choice == 2:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            c.compress_file(i, o, ultra=True, hybrid=False,
                            include_28=False, include_29=False, include_30=False)
        elif choice == 3:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            c.compress_file(i, o, ultra=True, hybrid=True,
                            include_28=False, include_29=False, include_30=False)
        elif choice == 4:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            c.compress_file(i, o, ultra=True, hybrid=True,
                            include_28=True, include_29=True, include_30=True)
        elif choice == 5:
            c.full_self_test()
        elif choice == 6:
            i = input("Compressed file: ").strip()
            o = input("Output file: ").strip() or i.rsplit('.', 1)[0] + ".orig"
            c.decompress_file(i, o)
        elif choice == 7:
            c.test_2704_pairs_lossless()
        elif choice == 8:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            c.compress_file(i, o, ultra=False, hybrid=False,
                            include_28=False, include_29=False, include_30=False)
        elif choice == 9:
            c.compress_ai_zaden()

        if choice != 0:
            input("\nPress Enter to return to the menu...")

if __name__ == "__main__":
    main()
