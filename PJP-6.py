#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PJP – 256 Lossless Transforms + 2704 Transform‑Pair Sequences
+ Hybrid Dictionary Mode + Quantum Transforms + Base64 + 6‑bit Text
+ Transforms 28–30 + .docx transforms 31–32
+ Zaden Block Optimization (Option 9) – tries both Absolute (hybrid + all transforms)
  and block‑optimized compression, picks the smaller result.
  Time limit per block can be set from 1 to 300 seconds.
  Zaden file header: single byte 0x33, followed by block_size (4 bytes LE),
  num_blocks (4 bytes LE), then unary‑coded keys, then inner compressed data.
+ Dynamic dictionary now automatically tries multiple index sizes (1-8) and
  BPE/RLE combinations (including 2 rounds of BPE) and picks the best.
============================================================================
"""

import math
import random
import decimal
import hashlib
import struct
import re
import os
import urllib.request
import sys
import subprocess
import importlib
import tempfile
import base64
import zipfile
import io
import xml.etree.ElementTree as ET
import time
from typing import Optional, List, Tuple, Dict, Callable
from collections import Counter, defaultdict

# ------------------------------------------------------------------
# Helper: install a single package via pip (silent, auto)
# ------------------------------------------------------------------
def install_package(pkg: str) -> bool:
    print(f"Installing {pkg}...")
    try:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', pkg])
        print(f"Successfully installed {pkg}")
        return True
    except Exception as e:
        print(f"Failed to install {pkg}: {e}")
        return False

# ------------------------------------------------------------------
# 1. Ask about quantum transforms – auto‑install if missing
# ------------------------------------------------------------------
USE_QUANTUM = False
HAS_QISKIT = False

quantum_choice = input("Enable quantum‑inspired transforms (requires Qiskit)? (y/n): ").strip().lower()
if quantum_choice == 'y':
    try:
        from qiskit import QuantumCircuit
        HAS_QISKIT = True
        USE_QUANTUM = True
        print("Quantum transforms ENABLED (Qiskit already installed).")
    except ImportError:
        print("Qiskit not found. Installing automatically...")
        if install_package('qiskit'):
            try:
                from qiskit import QuantumCircuit
                HAS_QISKIT = True
                USE_QUANTUM = True
                print("Quantum transforms ENABLED after automatic installation.")
            except ImportError:
                print("Qiskit installation succeeded but import failed – quantum transforms disabled.")
        else:
            print("Automatic installation failed – quantum transforms disabled.")
else:
    print("Quantum transforms disabled.")

# ------------------------------------------------------------------
# 2. Ask about other optional compression backends (zstandard, paq, etc.)
# ------------------------------------------------------------------
other_choice = input("Install other optional compression backends (zstandard, paq, mpmath, cython, python-docx)? (y/n): ").strip().lower()
if other_choice == 'y':
    for pkg in ['mpmath', 'zstandard', 'cython', 'paq', 'python-docx']:
        try:
            importlib.import_module(pkg)
        except ImportError:
            install_package(pkg)
else:
    print("Skipping other backends.")

# ---------- Optional compression backends ----------
try:
    import paq
except ImportError:
    paq = None

try:
    import zstandard as zstd
    zstd_cctx = zstd.ZstdCompressor(level=22)
    zstd_dctx = zstd.ZstdDecompressor()
    HAS_ZSTD = True
except ImportError:
    HAS_ZSTD = False

# ---------- (Re‑import Qiskit if it was just installed) ----------
if USE_QUANTUM and not HAS_QISKIT:
    try:
        from qiskit import QuantumCircuit
        HAS_QISKIT = True
    except ImportError:
        USE_QUANTUM = False
        print("Quantum transforms disabled because Qiskit could not be imported.")

PROGNAME = "PJP"

# ---------- Dictionary configuration ----------
DICT_DIR = "Dictionaries"
COMBINED_DICTIONARY_FILE = os.path.join(DICT_DIR, "dictionary_combined.txt")

DICTIONARY_FILES = [
    "generated.txt",
    "eng_news_2005_1M-sentences.txt",
    "eng_news_2005_1M-words.txt",
    "eng_news_2005_1M-sources.txt",
    "eng_news_2005_1M-co_n.txt",
    "eng_news_2005_1M-co_s.txt",
    "eng_news_2005_1M-inv_w_2.txt",
    "eng_news_2005_1M-inv_w_3.txt",
    "eng_news_2005_1M-inv_so.txt",
    "eng_news_2005_1M-meta.txt",
    "Dictionary.txt",
    "the-complete-reference-html-css-fifth-edition.txt",
]

DICTIONARY_URLS = [
    "https://drive.google.com/uc?export=download&id=1u_1dCEl8hhdEug6GwkOxHAuSx_6_Pme9",
    "https://drive.google.com/uc?export=download&id=1pVqNN5JZ2AeOCgRaHkv4Vv6Byr4zK20e",
    "https://drive.google.com/uc?export=download&id=1ZSC-Tn76x8itdN0rCp-Zw17hGudxbjxo",
    "https://drive.google.com/uc?export=download&id=1VB_7tzngs4GxjclSRyRDnxgS8znT2w2S",
    "https://drive.google.com/uc?export=download&id=1KVIRgiMrhCUCqQZJ3UT67ztls2GqGJzz",
    "https://drive.google.com/uc?export=download&id=1Z3Lx6SqL4HWsnmbJCez4kXWRQQhUXWKL",
    "https://drive.google.com/uc?export=download&id=1br2bdRMkZEVVRPKYmC4IIaZuAjxFJE4N",
    "https://drive.google.com/uc?export=download&id=1aE6ubPZiJ8rr3lEVk8fFJYjDQ1y1rU0X",
    "https://drive.google.com/uc?export=download&id=1uro3TZe-t5zPx2Qu2xrTL3lU8N0melk9",
    "https://drive.google.com/uc?export=download&id=1HqsTH1DqpWNpGbn9VtD7-SB6wVqA90R2",
    "https://drive.google.com/uc?export=download&id=1zZ8iMeBC3605NZhuc4UE9jx_w_lZFg5B",
    "https://drive.google.com/uc?export=download&id=1dDdqYDgm7f-smS7KF70Wf0KmyFo-ft1M",
]

MAX_LINE_ENTRIES = 1024

def download_and_merge_dictionaries():
    if not os.path.exists(DICT_DIR):
        os.makedirs(DICT_DIR)

    if os.path.exists(COMBINED_DICTIONARY_FILE):
        print(f"Combined dictionary '{COMBINED_DICTIONARY_FILE}' already exists. Skipping download.")
        return True

    all_words = set()
    success_count = 0

    for idx, (filename, url) in enumerate(zip(DICTIONARY_FILES, DICTIONARY_URLS)):
        local_path = os.path.join(DICT_DIR, filename)
        print(f"Downloading {filename} to {DICT_DIR}/ ...")
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response:
                content = response.read()

            if b'<html' in content[:200].lower():
                print(f"  WARNING: {filename} appears to be an HTML page (not a text file). Skipping.")
                continue

            with open(local_path, 'wb') as f:
                f.write(content)

            with open(local_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    w = line.strip()
                    if not w:
                        continue
                    try:
                        decoded = base64.b64decode(w, validate=True)
                        decoded_str = decoded.decode('utf-8')
                        all_words.add(decoded_str)
                    except Exception:
                        all_words.add(w)

            print(f"  Downloaded {filename} ({os.path.getsize(local_path)} bytes)")
            success_count += 1

        except Exception as e:
            print(f"  WARNING: Could not download {filename}: {e}")
            if os.path.exists(local_path):
                os.remove(local_path)

    if success_count == 0:
        print("ERROR: No dictionary files could be downloaded.")
        print("Proceeding without static word and line dictionaries.")
        return False

    try:
        with open(COMBINED_DICTIONARY_FILE, 'w', encoding='utf-8') as f:
            for word in sorted(all_words):
                f.write(word + '\n')
        print(f"Merged {len(all_words)} unique words into {COMBINED_DICTIONARY_FILE} "
              f"({os.path.getsize(COMBINED_DICTIONARY_FILE)} bytes)")
        return True
    except Exception as e:
        print(f"ERROR: Could not write combined dictionary: {e}")
        return False

# ---------- Constants ----------
PRIMES = [p for p in range(2, 256) if all(p % d != 0 for d in range(2, int(p ** 0.5) + 1))]
PI_DIGITS = [79, 17, 111]
BLOCK_SIZE = 1024

def find_nearest_prime_around(n: int) -> int:
    o = 0
    while True:
        c1, c2 = n - o, n + o
        if c1 >= 2 and all(c1 % d != 0 for d in range(2, int(c1 ** 0.5) + 1)):
            return c1
        if c2 >= 2 and all(c2 % d != 0 for d in range(2, int(c2 ** 0.5) + 1)):
            return c2
        o += 1

def sha256_8bytes(data: bytes) -> bytes:
    return hashlib.sha256(data).digest()[:8]

def xor_prime_hash(word: str) -> bytes:
    prime = 2147483647
    total = sum(ord(c) for c in word)
    transformed = total ^ prime
    return transformed.to_bytes(8, 'big')

# ---------- 6‑bit alphabet for transform 27 (exactly 64 chars) ----------
ALPHABET_6BIT = (
    "ABCDEFGHIJKLMNOPQRSTUVWXYZ"    # 26
    "abcdefghijklmnopqrstuvwxyz"    # 26
    "0123456789"                    # 10
    " \n"                           # 2  (space and newline)
)  # Total = 64
assert len(ALPHABET_6BIT) == 64
CHAR_TO_6BIT = {ch: i for i, ch in enumerate(ALPHABET_6BIT)}
SIXBIT_TO_CHAR = {i: ch for ch, i in CHAR_TO_6BIT.items()}

# ---------- Main Compressor Class ----------
class PJPCompressor:
    def __init__(self):
        download_and_merge_dictionaries()

        self.PI_DIGITS = PI_DIGITS.copy()
        self.seed_tables = self._gen_seed_tables(num=126, size=40, seed=42)
        self.fibonacci = self._gen_fib(100)
        self.PI_STR = "3.14159265358979323846264338327950288419716939937510"

        self._build_transform_maps()
        self.sequences = self._build_pair_sequences()
        self.pair_lookup = {idx: (t1, t2) for idx, (t1, t2) in enumerate(self.sequences)}

        self.static_dict, self.word_to_index = self._load_static_dictionary()
        self.line_dict, self.line_to_index = self._load_line_dictionary()

        if USE_QUANTUM and HAS_QISKIT:
            self._precompute_quantum_transforms()

    # ------------------------------------------------------------------
    # Quantum transform generation
    # ------------------------------------------------------------------
    def _generate_permutation_from_circuit(self, num_qubits: int, seed: int) -> List[int]:
        qc = QuantumCircuit(num_qubits)
        rng = random.Random(seed)
        for qubit in range(num_qubits):
            qc.h(qubit)
            qc.rz(rng.random() * 2 * math.pi, qubit)
            qc.rx(rng.random() * 2 * math.pi, qubit)
        for _ in range(num_qubits):
            for i in range(num_qubits - 1):
                qc.cx(i, i+1)
            qc.barrier()
            for i in range(num_qubits):
                qc.rz(rng.random() * 2 * math.pi, i)
                qc.rx(rng.random() * 2 * math.pi, i)

        try:
            qasm_str = qc.qasm()
        except AttributeError:
            qasm_str = qc.draw('text')

        final_seed = seed + hash(qasm_str) % 1000000
        rng2 = random.Random(final_seed)
        n = 1 << num_qubits
        perm = list(range(n))
        rng2.shuffle(perm)
        if num_qubits == 12:
            perm_2704 = list(range(2704))
            rng2 = random.Random(final_seed)
            rng2.shuffle(perm_2704)
            return perm_2704
        else:
            return perm

    def _precompute_quantum_transforms(self):
        self.quantum_fast_perms = []
        for i in range(9):
            seed = 1000 + i
            perm = self._generate_permutation_from_circuit(8, seed)
            self.quantum_fast_perms.append(perm)

        self.quantum_ultra_perms = []
        for i in range(17):
            seed = 2000 + i
            perm = self._generate_permutation_from_circuit(12, seed)
            self.quantum_ultra_perms.append(perm)

        self.quantum_fast_transforms = []
        for perm in self.quantum_fast_perms:
            fwd, rev = self._make_substitution_transform(perm, 256)
            self.quantum_fast_transforms.append((fwd, rev))

        self.quantum_ultra_transforms = []
        for perm in self.quantum_ultra_perms:
            fwd, rev = self._make_permutation_transform(perm, 2704)
            self.quantum_ultra_transforms.append((fwd, rev))

        for idx, (fwd, rev) in enumerate(self.quantum_fast_transforms, start=257):
            self.fwd_transforms[idx] = fwd
            self.rev_transforms[idx] = rev
        for idx, (fwd, rev) in enumerate(self.quantum_ultra_transforms, start=266):
            self.fwd_transforms[idx] = fwd
            self.rev_transforms[idx] = rev

    def _make_substitution_transform(self, perm: List[int], size: int):
        inv_perm = [0] * size
        for i, p in enumerate(perm):
            inv_perm[p] = i
        def forward(data: bytes) -> bytes:
            return bytes(perm[b] for b in data)
        def reverse(data: bytes) -> bytes:
            return bytes(inv_perm[b] for b in data)
        return forward, reverse

    def _make_permutation_transform(self, perm: List[int], block_size: int):
        inv_perm = [0] * block_size
        for i, p in enumerate(perm):
            inv_perm[p] = i
        def forward(data: bytes) -> bytes:
            out = bytearray()
            for offset in range(0, len(data), block_size):
                block = data[offset:offset+block_size]
                if len(block) < block_size:
                    out += block
                else:
                    new_block = bytearray(block_size)
                    for i in range(block_size):
                        new_block[perm[i]] = block[i]
                    out += new_block
            return bytes(out)
        def reverse(data: bytes) -> bytes:
            out = bytearray()
            for offset in range(0, len(data), block_size):
                block = data[offset:offset+block_size]
                if len(block) < block_size:
                    out += block
                else:
                    new_block = bytearray(block_size)
                    for i in range(block_size):
                        new_block[inv_perm[i]] = block[i]
                    out += new_block
            return bytes(out)
        return forward, reverse

    # ------------------------------------------------------------------
    # Dictionary loaders
    # ------------------------------------------------------------------
    def _load_static_dictionary(self):
        if not os.path.exists(COMBINED_DICTIONARY_FILE):
            print(f"ERROR: {COMBINED_DICTIONARY_FILE} not found. No dictionaries loaded.")
            return [], {}

        words_set = set()
        try:
            with open(COMBINED_DICTIONARY_FILE, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    w = line.strip()
                    if w:
                        words_set.add(w)
        except Exception as e:
            print(f"Warning: could not read {COMBINED_DICTIONARY_FILE}: {e}")
            return [], {}

        sorted_words = sorted(words_set)
        word_to_idx = {w: i for i, w in enumerate(sorted_words)}
        print(f"Loaded static word dictionary: {len(sorted_words)} unique words.")
        return sorted_words, word_to_idx

    def _load_line_dictionary(self):
        if not os.path.exists(COMBINED_DICTIONARY_FILE):
            return [], {}

        lines = []
        try:
            with open(COMBINED_DICTIONARY_FILE, 'r', encoding='utf-8', errors='ignore') as f:
                for raw_line in f:
                    phrase = raw_line.strip()
                    if phrase and phrase not in lines:
                        lines.append(phrase)
                        if len(lines) >= MAX_LINE_ENTRIES:
                            break
        except Exception as e:
            print(f"Warning: could not read {COMBINED_DICTIONARY_FILE}: {e}")
            return [], {}

        if not lines:
            return [], {}

        lines.sort(key=len, reverse=True)
        line_to_idx = {phrase: i for i, phrase in enumerate(lines)}
        print(f"Loaded line dictionary: {len(lines)} phrases.")
        return lines, line_to_idx

    # ------------------------------------------------------------------
    # pi / constant helpers
    # ------------------------------------------------------------------
    def get_pi_digits(self, n: int) -> str:
        if n < 1: return ""
        return self.PI_STR[2:2 + n]

    def find_lossless_k(self, n: int):
        if n < 1: return 0, True
        true_digits = self.get_pi_digits(n)
        true_scaled = int(self.PI_STR.replace('.', '')[:n + 1])
        DENOM = 16777216
        decimal.getcontext().prec = 50
        pi_dec = decimal.Decimal(self.PI_STR)
        k_float = (pi_dec - 3) * DENOM
        k_candidate = int(round(k_float))
        k_candidate = max(0, min(k_candidate, DENOM - 1))
        approx_scaled = (3 * 10 ** n * DENOM + k_candidate * 10 ** n) // DENOM
        return k_candidate, approx_scaled == true_scaled

    def to_bin(self, value: int, bits: int) -> str:
        return format(value, 'b').zfill(bits)

    def get_bit_size(self, k: int) -> int:
        return 23 if k <= 0x7FFFFF else 25

    def transform_17(self, data: bytes) -> bytes:
        if not data: return b''
        k, _ = self.find_lossless_k(7)
        bits_used = self.get_bit_size(k)
        bit_str = self.to_bin(k, bits_used)
        mask_bytes = []
        for i in range(0, len(bit_str), 8):
            byte_bits = bit_str[i:i + 8]
            if len(byte_bits) < 8:
                byte_bits = byte_bits.ljust(8, '0')
            mask_bytes.append(int(byte_bits, 2))
        mask = bytes(mask_bytes)
        t = bytearray(data)
        for i in range(len(t)):
            t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_17 = transform_17

    def get_basel_digits(self, n: int) -> str:
        decimal.getcontext().prec = n + 5
        pi = decimal.Decimal(self.PI_STR)
        basel = (pi * pi) / decimal.Decimal(6)
        s = str(basel).replace('.', '')
        return s[:n]

    def get_one_over_e_digits(self, n: int) -> str:
        decimal.getcontext().prec = n + 5
        e = decimal.Decimal(1).exp()
        inv_e = decimal.Decimal(1) / e
        s = str(inv_e).replace('.', '')
        return s[:n]

    def get_5e_digits(self, n: int) -> str:
        decimal.getcontext().prec = n + 5
        e = decimal.Decimal(1).exp()
        five_e = decimal.Decimal(5) * e
        s = str(five_e).replace('.', '')
        return s[:n]

    # ------------------------------------------------------------------
    # Seed tables, Fibonacci
    # ------------------------------------------------------------------
    def _gen_seed_tables(self, num=126, size=40, seed=42):
        random.seed(seed)
        return [[random.randint(5, 255) for _ in range(size)] for _ in range(num)]

    def _gen_fib(self, n):
        a, b = 0, 1
        res = [a, b]
        for _ in range(2, n):
            a, b = b, a + b
            res.append(b)
        return res

    def get_seed(self, idx: int, val: int) -> int:
        if 0 <= idx < len(self.seed_tables):
            return self.seed_tables[idx][val % 40]
        return 0

    # ------------------------------------------------------------------
    # Bit helpers (for RLE)
    # ------------------------------------------------------------------
    def _append_bits(self, bitlist: List[int], value: int, count: int):
        for i in range(count - 1, -1, -1):
            bitlist.append((value >> i) & 1)

    def _read_bits(self, bits: List[int], pos: int, count: int) -> int:
        val = 0
        for i in range(count):
            if pos + i >= len(bits): return 0
            val = (val << 1) | bits[pos + i]
        return val

    # ------------------------------------------------------------------
    # RLE transform 00
    # ------------------------------------------------------------------
    def transform_00(self, data: bytes) -> bytes:
        if not data: return b'\x00'
        best_result = None
        best_length = float('inf')
        best_shifts = []
        MAX_PASSES = 10
        current = bytearray(data)
        applied_shifts = []
        for _ in range(MAX_PASSES):
            best_shift = 0
            best_shifted = current
            best_score = float('-inf')
            for shift in range(256):
                tmp = bytearray(current)
                for j in range(len(tmp)):
                    tmp[j] = (tmp[j] + shift) % 256
                score = 0
                i = 0
                while i < len(tmp):
                    val = tmp[i]
                    run = 1
                    i += 1
                    while i < len(tmp) and tmp[i] == val:
                        run += 1
                        i += 1
                    score += run * run
                if score > best_score:
                    best_score = score
                    best_shifted = tmp
                    best_shift = shift
            applied_shifts.append(best_shift)
            rle_encoded = self._apply_rle_to_shifted(best_shifted, best_shift)
            if len(rle_encoded) < best_length:
                best_length = len(rle_encoded)
                best_result = rle_encoded
                best_shifts = applied_shifts.copy()
            current = best_shifted
            if len(rle_encoded) >= len(data):
                break
        if best_result is None or best_length >= len(data):
            return bytes([0]) + data
        header = bytearray([len(best_shifts)])
        header.extend(best_shifts)
        return header + best_result

    def _apply_rle_to_shifted(self, shifted_data: bytearray, shift: int) -> bytes:
        bits = []
        self._append_bits(bits, 0b010, 3)
        self._append_bits(bits, shift, 8)
        i = 0
        n = len(shifted_data)
        while i < n:
            val = shifted_data[i]
            run = 1
            i += 1
            while i < n and shifted_data[i] == val:
                run += 1
                i += 1
            while run >= 13:
                chunk = min(run, 268)
                self._append_bits(bits, 0b1111, 4)
                self._append_bits(bits, chunk - 13, 8)
                self._append_bits(bits, val, 8)
                run -= chunk
            if run == 1:
                self._append_bits(bits, 0b00, 2)
                self._append_bits(bits, val, 8)
            elif run <= 5:
                self._append_bits(bits, 0b01, 2)
                self._append_bits(bits, run - 2, 2)
                self._append_bits(bits, val, 8)
            elif run <= 12:
                self._append_bits(bits, 0b10, 2)
                self._append_bits(bits, run - 6, 3)
                self._append_bits(bits, val, 8)
        pad = (8 - len(bits) % 8) % 8
        self._append_bits(bits, 0, pad)
        out = bytearray()
        for j in range(0, len(bits), 8):
            byte = 0
            for k in range(8):
                if j + k < len(bits):
                    byte = (byte << 1) | bits[j + k]
            out.append(byte)
        return bytes(out)

    def reverse_transform_00(self, cdata: bytes) -> bytes:
        if not cdata or cdata == b'\x00': return b''
        if cdata[0] == 0: return cdata[1:]
        num_passes = cdata[0]
        if num_passes == 0 or len(cdata) < 1 + num_passes: return b''
        shifts = list(cdata[1:1 + num_passes])
        rle_data = cdata[1 + num_passes:]
        decoded = self._rle_decode(rle_data)
        if decoded is None: return b''
        current = bytearray(decoded)
        for shift in reversed(shifts):
            for i in range(len(current)):
                current[i] = (current[i] - shift) % 256
        return bytes(current)

    def _rle_decode(self, data: bytes) -> Optional[bytearray]:
        if not data: return None
        bits = []
        for b in data:
            for i in range(7, -1, -1):
                bits.append((b >> i) & 1)
        pos = 0
        nbits = len(bits)
        if nbits < 11: return None
        marker = self._read_bits(bits, pos, 3)
        pos += 3
        if marker != 0b010: return None
        pos += 8
        out = bytearray()
        while pos < nbits:
            if pos + 2 > nbits: break
            prefix = self._read_bits(bits, pos, 2)
            pos += 2
            if prefix == 0b00:
                if pos + 8 > nbits: break
                run = 1
            elif prefix == 0b01:
                if pos + 2 + 8 > nbits: break
                run = 2 + self._read_bits(bits, pos, 2)
                pos += 2
            elif prefix == 0b10:
                if pos + 3 + 8 > nbits: break
                run = 6 + self._read_bits(bits, pos, 3)
                pos += 3
            else:
                if pos + 2 + 8 + 8 > nbits: break
                if self._read_bits(bits, pos, 2) != 0b11: return None
                pos += 2
                run = 13 + self._read_bits(bits, pos, 8)
                pos += 8
            if pos + 8 > nbits: break
            val = self._read_bits(bits, pos, 8)
            pos += 8
            out.extend([val] * run)
        for i in range(pos, nbits):
            if bits[i] != 0: return None
        return out

    # ------------------------------------------------------------------
    # Transforms 01‑21
    # ------------------------------------------------------------------
    def transform_01(self, d, r=100):
        t = bytearray(d)
        for prime in PRIMES:
            xor_val = prime if prime == 2 else max(1, math.ceil(prime * 4096 / 28672))
            for _ in range(r):
                for i in range(0, len(t), 3):
                    if i < len(t): t[i] ^= xor_val
        return bytes(t)
    reverse_transform_01 = transform_01

    def transform_02(self, d):
        if len(d) < 1: return b''
        t = bytearray(d)
        checksum = sum(d) % 256
        pattern_index = (len(d) + checksum) % 256
        pattern_values = self._get_pattern(4, pattern_index)
        for i in range(1, len(t), 4):
            if i < len(t): t[i] ^= pattern_values[i % len(pattern_values)]
        return bytes([pattern_index]) + bytes(t)
    def reverse_transform_02(self, d):
        if len(d) < 2: return b''
        pattern_index = d[0]
        t = bytearray(d[1:])
        pattern_values = self._get_pattern(4, pattern_index)
        for i in range(1, len(t), 4):
            if i < len(t): t[i] ^= pattern_values[i % len(pattern_values)]
        return bytes(t)

    def transform_03(self, d):
        if len(d) < 1: return b''
        t = bytearray(d)
        rotation = (len(d) * 13 + sum(d)) % 8
        if rotation == 0: rotation = 1
        for i in range(2, len(t), 5):
            if i < len(t): t[i] = ((t[i] << rotation) | (t[i] >> (8 - rotation))) & 0xFF
        return bytes([rotation]) + bytes(t)
    def reverse_transform_03(self, d):
        if len(d) < 2: return b''
        rotation = d[0]
        t = bytearray(d[1:])
        for i in range(2, len(t), 5):
            if i < len(t): t[i] = ((t[i] >> rotation) | (t[i] << (8 - rotation))) & 0xFF
        return bytes(t)

    def transform_04(self, d, r=100):
        t = bytearray(d)
        for _ in range(r):
            for i in range(len(t)): t[i] = (t[i] - (i % 256)) % 256
        return bytes(t)
    def reverse_transform_04(self, d, r=100):
        t = bytearray(d)
        for _ in range(r):
            for i in range(len(t)): t[i] = (t[i] + (i % 256)) % 256
        return bytes(t)

    def transform_05(self, d, s=3):
        t = bytearray(d)
        for i in range(len(t)): t[i] = ((t[i] << s) | (t[i] >> (8 - s))) & 0xFF
        return bytes(t)
    def reverse_transform_05(self, d, s=3):
        t = bytearray(d)
        for i in range(len(t)): t[i] = ((t[i] >> s) | (t[i] << (8 - s))) & 0xFF
        return bytes(t)

    def transform_06(self, d, sd=42):
        random.seed(sd)
        sub = list(range(256))
        random.shuffle(sub)
        t = bytearray(d)
        for i in range(len(t)): t[i] = sub[t[i]]
        return bytes(t)
    def reverse_transform_06(self, d, sd=42):
        random.seed(sd)
        sub = list(range(256))
        random.shuffle(sub)
        inv = [0]*256
        for i in range(256): inv[sub[i]] = i
        t = bytearray(d)
        for i in range(len(t)): t[i] = inv[t[i]]
        return bytes(t)

    def transform_07(self, d, r=100):
        t = bytearray(d)
        sh = len(d) % len(self.PI_DIGITS)
        pi_rot = self.PI_DIGITS[sh:] + self.PI_DIGITS[:sh]
        sz = len(d) % 256
        for i in range(len(t)): t[i] ^= sz
        for _ in range(r):
            for i in range(len(t)): t[i] ^= pi_rot[i % len(pi_rot)]
        return bytes(t)
    reverse_transform_07 = transform_07

    def transform_08(self, d, r=100):
        t = bytearray(d)
        sh = len(d) % len(self.PI_DIGITS)
        pi_rot = self.PI_DIGITS[sh:] + self.PI_DIGITS[:sh]
        p = find_nearest_prime_around(len(d) % 256)
        for i in range(len(t)): t[i] ^= p
        for _ in range(r):
            for i in range(len(t)): t[i] ^= pi_rot[i % len(pi_rot)]
        return bytes(t)
    reverse_transform_08 = transform_08

    def transform_09(self, d, r=100):
        t = bytearray(d)
        sh = len(d) % len(self.PI_DIGITS)
        pi_rot = self.PI_DIGITS[sh:] + self.PI_DIGITS[:sh]
        p = find_nearest_prime_around(len(d) % 256)
        seed = self.get_seed(len(d) % len(self.seed_tables), len(d))
        for i in range(len(t)): t[i] ^= p ^ seed
        for _ in range(r):
            for i in range(len(t)): t[i] ^= pi_rot[i % len(pi_rot)] ^ (i % 256)
        return bytes(t)
    reverse_transform_09 = transform_09

    def transform_10(self, data: bytes) -> bytes:
        if not data: return b'\x00'
        cnt = sum(1 for i in range(len(data)-1) if data[i:i+2] == b'X1')
        n = (((cnt * 2) + 1) // 3) * 3 % 256
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= n
        return bytes([n]) + bytes(t)
    def reverse_transform_10(self, data: bytes) -> bytes:
        if len(data) < 1: return b''
        n = data[0]
        t = bytearray(data[1:])
        for i in range(len(t)): t[i] ^= n
        return bytes(t)

    def transform_11(self, data: bytes) -> bytes:
        if not data: return b''
        t = bytearray(data)
        length = len(t)
        for i in range(length):
            fib_idx = (i + length) % len(self.fibonacci)
            fib_val = self.fibonacci[fib_idx] % 256
            pos_val = (i * 13 + length * 17) % 256
            key = (fib_val ^ pos_val) % 256
            t[i] ^= key
        return bytes(t)
    reverse_transform_11 = transform_11

    def transform_12(self, data: bytes) -> bytes:
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= self.fibonacci[i % len(self.fibonacci)] % 256
        return bytes(t)
    reverse_transform_12 = transform_12

    def transform_13(self, d):
        if not d: return b''
        repeats = self._calculate_repeats(d)
        current_value = len(d) % 256
        prime_values = []
        count = 0
        while count < repeats:
            current_value = find_nearest_prime_around(current_value)
            prime_values.append(current_value)
            count += 1
        t = bytearray(d)
        xor_value = prime_values[-1] if prime_values else 0
        for i in range(len(t)): t[i] ^= xor_value
        repeat_byte = (repeats - 1) % 256
        return bytes([repeat_byte]) + bytes(t)
    def reverse_transform_13(self, d):
        if len(d) < 2: return b''
        repeat_byte = d[0]
        repeats = (repeat_byte + 1) % 256
        if repeats == 0: repeats = 256
        t = bytearray(d[1:])
        current_value = len(t) % 256
        prime_values = []
        count = 0
        while count < repeats:
            current_value = find_nearest_prime_around(current_value)
            prime_values.append(current_value)
            count += 1
        xor_value = prime_values[-1] if prime_values else 0
        for i in range(len(t)): t[i] ^= xor_value
        return bytes(t)

    # Transform 14 is NOT bijective; skipped in pair base.

    def transform_15(self, d):
        if len(d) < 1: return b''
        t = bytearray(d)
        pattern_index = len(d) % 256
        pattern_values = self._get_pattern(3, pattern_index)
        for i in range(0, len(t), 3):
            if i < len(t): t[i] = (t[i] + pattern_values[i % len(pattern_values)]) % 256
        return bytes([pattern_index]) + bytes(t)
    def reverse_transform_15(self, d):
        if len(d) < 2: return b''
        pattern_index = d[0]
        t = bytearray(d[1:])
        pattern_values = self._get_pattern(3, pattern_index)
        for i in range(0, len(t), 3):
            if i < len(t): t[i] = (t[i] - pattern_values[i % len(pattern_values)]) % 256
        return bytes(t)

    def transform_16(self, data: bytes) -> bytes:
        if not data: return b''
        xor_byte = (len(data) * 7 + 13) % 256
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= xor_byte
        return bytes(t)
    reverse_transform_16 = transform_16

    # transform_17 defined earlier
    def transform_18(self, data: bytes) -> bytes:
        if not data: return b''
        digits = self.get_basel_digits(max(10, len(data)//2 + 5))
        mask = bytes(int(digits[i:i+2]) % 256 for i in range(0, len(digits), 2))
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_18 = transform_18

    def transform_19(self, data: bytes) -> bytes:
        if not data: return b''
        digits = self.get_one_over_e_digits(max(10, len(data)//2 + 5))
        mask = bytes(int(digits[i:i+2]) % 256 for i in range(0, len(digits), 2))
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_19 = transform_19

    def transform_20(self, data: bytes) -> bytes:
        if not data: return b''
        digits = self.get_5e_digits(max(10, len(data)//2 + 5))
        mask = bytes(int(digits[i:i+2]) % 256 for i in range(0, len(digits), 2))
        t = bytearray(data)
        for i in range(len(t)): t[i] ^= mask[i % len(mask)]
        return bytes(t)
    reverse_transform_20 = transform_20

    def transform_21(self, data: bytes) -> bytes:
        if not data: return b''
        shift = 255
        t = bytearray(data)
        for i in range(len(t)): t[i] = (t[i] + shift) % 256
        return bytes(t)
    def reverse_transform_21(self, data: bytes) -> bytes:
        if not data: return b''
        shift = 255
        t = bytearray(data)
        for i in range(len(t)): t[i] = (t[i] - shift) % 256
        return bytes(t)

    # ------------------------------------------------------------------
    # Transform 22 – Base64
    # ------------------------------------------------------------------
    def transform_22(self, data: bytes) -> bytes:
        return base64.b64encode(data)

    def reverse_transform_22(self, data: bytes) -> bytes:
        try:
            return base64.b64decode(data, validate=False)
        except Exception:
            return data

    # ------------------------------------------------------------------
    # Transform 23 – SHA‑256 word tokenizer
    # ------------------------------------------------------------------
    def transform_23(self, data: bytes) -> bytes:
        if not data: return b'\x00\x00\x00\x00'
        try:
            text = data.decode('latin-1')
        except:
            text = data.decode('latin-1', errors='replace')
        pattern = r'([A-Za-z0-9_]+)'
        tokens = re.split(pattern, text)
        hash_to_word = {}
        token_list = []
        for i, tok in enumerate(tokens):
            if i % 2 == 1:
                word_bytes = tok.encode('latin-1')
                h = sha256_8bytes(word_bytes)
                if h in hash_to_word:
                    if hash_to_word[h] != word_bytes:
                        token_list.append((False, word_bytes))
                        continue
                else:
                    hash_to_word[h] = word_bytes
                token_list.append((True, h))
            else:
                if tok:
                    token_list.append((False, tok.encode('latin-1')))
        dict_entries = list(hash_to_word.items())
        num_entries = len(dict_entries)
        result = bytearray()
        result += struct.pack('>I', num_entries)
        for h, wb in dict_entries:
            result += h
            result += struct.pack('>H', len(wb))
            result += wb
        for is_word, payload in token_list:
            if is_word:
                result += b'\x01'
                result += payload
            else:
                result += b'\x00'
                result += struct.pack('>H', len(payload))
                result += payload
        return bytes(result)

    def reverse_transform_23(self, data: bytes) -> bytes:
        if not data: return b''
        if len(data) < 4: return data
        num_entries = struct.unpack('>I', data[:4])[0]
        pos = 4
        hash_to_word = {}
        for _ in range(num_entries):
            if pos + 10 > len(data): break
            h = data[pos:pos+8]
            pos += 8
            wlen = struct.unpack('>H', data[pos:pos+2])[0]
            pos += 2
            if pos + wlen > len(data): break
            wb = data[pos:pos+wlen]
            pos += wlen
            hash_to_word[h] = wb
        out = bytearray()
        while pos < len(data):
            if pos >= len(data): break
            typ = data[pos]
            pos += 1
            if typ == 1:
                if pos + 8 > len(data): break
                h = data[pos:pos+8]
                pos += 8
                wb = hash_to_word.get(h)
                out += wb if wb else h
            elif typ == 0:
                if pos + 2 > len(data): break
                rawlen = struct.unpack('>H', data[pos:pos+2])[0]
                pos += 2
                if pos + rawlen > len(data): break
                out += data[pos:pos+rawlen]
                pos += rawlen
            else:
                break
        return bytes(out)

    # ------------------------------------------------------------------
    # Transform 24 – XOR‑prime word tokenizer
    # ------------------------------------------------------------------
    def transform_24(self, data: bytes) -> bytes:
        if not data: return b'\x00\x00\x00\x00'
        try:
            text = data.decode('latin-1')
        except:
            text = data.decode('latin-1', errors='replace')
        pattern = r'([A-Za-z0-9_]+)'
        tokens = re.split(pattern, text)
        hash_to_word = {}
        token_list = []
        for i, tok in enumerate(tokens):
            if i % 2 == 1:
                word_bytes = tok.encode('latin-1')
                h = xor_prime_hash(tok)
                if h in hash_to_word:
                    if hash_to_word[h] != word_bytes:
                        token_list.append((False, word_bytes))
                        continue
                else:
                    hash_to_word[h] = word_bytes
                token_list.append((True, h))
            else:
                if tok:
                    token_list.append((False, tok.encode('latin-1')))
        dict_entries = list(hash_to_word.items())
        num_entries = len(dict_entries)
        result = bytearray()
        result += struct.pack('>I', num_entries)
        for h, wb in dict_entries:
            result += h
            result += struct.pack('>H', len(wb))
            result += wb
        for is_word, payload in token_list:
            if is_word:
                result += b'\x01'
                result += payload
            else:
                result += b'\x00'
                result += struct.pack('>H', len(payload))
                result += payload
        return bytes(result)

    def reverse_transform_24(self, data: bytes) -> bytes:
        if not data: return b''
        if len(data) < 4: return data
        num_entries = struct.unpack('>I', data[:4])[0]
        pos = 4
        hash_to_word = {}
        for _ in range(num_entries):
            if pos + 10 > len(data): break
            h = data[pos:pos+8]
            pos += 8
            wlen = struct.unpack('>H', data[pos:pos+2])[0]
            pos += 2
            if pos + wlen > len(data): break
            wb = data[pos:pos+wlen]
            pos += wlen
            hash_to_word[h] = wb
        out = bytearray()
        while pos < len(data):
            if pos >= len(data): break
            typ = data[pos]
            pos += 1
            if typ == 1:
                if pos + 8 > len(data): break
                h = data[pos:pos+8]
                pos += 8
                wb = hash_to_word.get(h)
                out += wb if wb else h
            elif typ == 0:
                if pos + 2 > len(data): break
                rawlen = struct.unpack('>H', data[pos:pos+2])[0]
                pos += 2
                if pos + rawlen > len(data): break
                out += data[pos:pos+rawlen]
                pos += rawlen
            else:
                break
        return bytes(out)

    # ------------------------------------------------------------------
    # Transform 25 – Dynamic Dictionary Tokenizer (multi‑strategy)
    # ------------------------------------------------------------------
    def _split_text_into_chunks(self, text: str, level: str = 'all') -> List[str]:
        if level == 'paragraph':
            return re.split(r'(\n\n)', text)
        elif level == 'line':
            return re.split(r'(\n)', text)
        elif level == 'sentence':
            return re.split(r'([.!?]+)', text)
        elif level == 'word':
            return re.split(r'(\s+|\b)', text)
        else:
            chunks = []
            paragraphs = re.split(r'(\n\n)', text)
            for i, para in enumerate(paragraphs):
                if i % 2 == 1:
                    chunks.append(para)
                    continue
                lines = re.split(r'(\n)', para)
                for j, line in enumerate(lines):
                    if j % 2 == 1:
                        chunks.append(line)
                        continue
                    sentences = re.split(r'([.!?]+)', line)
                    for k, sent in enumerate(sentences):
                        if k % 2 == 1:
                            chunks.append(sent)
                            continue
                        words = re.split(r'(\s+|\b)', sent)
                        chunks.extend(words)
            return chunks

    # BPE helpers
    def _bpe_learn_and_apply(self, data: bytes) -> Tuple[bytes, List[Tuple[bytes, int]]]:
        if len(data) < 2:
            return data, []
        bigram_counts = Counter()
        for i in range(len(data)-1):
            bigram = data[i:i+2]
            bigram_counts[bigram] += 1
        candidates = [(bg, cnt) for bg, cnt in bigram_counts.items() if cnt >= 2]
        candidates.sort(key=lambda x: -x[1])
        candidates = candidates[:255]
        if not candidates:
            return data, []
        table = [(bg, i+1) for i, (bg, cnt) in enumerate(candidates)]
        bigram_to_code = {bg: code for bg, code in table}
        i = 0
        compressed = bytearray()
        while i < len(data):
            if i + 1 < len(data):
                bigram = data[i:i+2]
                if bigram in bigram_to_code:
                    compressed.append(bigram_to_code[bigram])
                    i += 2
                    continue
            compressed.append(data[i])
            i += 1
        return bytes(compressed), table

    def _bpe_apply(self, data: bytes, table: List[Tuple[bytes, int]]) -> bytes:
        bigram_to_code = {bg: code for bg, code in table}
        i = 0
        compressed = bytearray()
        while i < len(data):
            if i + 1 < len(data):
                bigram = data[i:i+2]
                if bigram in bigram_to_code:
                    compressed.append(bigram_to_code[bigram])
                    i += 2
                    continue
            compressed.append(data[i])
            i += 1
        return bytes(compressed)

    def _bpe_decompress(self, data: bytes, table: List[Tuple[bytes, int]]) -> bytes:
        code_to_bigram = {code: bg for bg, code in table}
        result = bytearray()
        for byte in data:
            if byte in code_to_bigram:
                result.extend(code_to_bigram[byte])
            else:
                result.append(byte)
        return bytes(result)

    # RLE helpers
    def _rle_compress_bytes(self, data: bytes) -> bytes:
        ESCAPE = 0xFE
        result = bytearray()
        i = 0
        n = len(data)
        while i < n:
            val = data[i]
            run = 1
            i += 1
            while i < n and data[i] == val and run < 258:
                run += 1
                i += 1
            if run >= 3:
                result.append(ESCAPE)
                result.append(run - 3)
                result.append(val)
            else:
                for _ in range(run):
                    if val == ESCAPE:
                        result.append(ESCAPE)
                        result.append(ESCAPE)
                    else:
                        result.append(val)
        return bytes(result)

    def _rle_decompress_bytes(self, data: bytes) -> bytes:
        ESCAPE = 0xFE
        result = bytearray()
        i = 0
        n = len(data)
        while i < n:
            b = data[i]
            i += 1
            if b == ESCAPE:
                if i >= n:
                    result.append(ESCAPE)
                    break
                count_byte = data[i]
                i += 1
                if count_byte == ESCAPE:
                    result.append(ESCAPE)
                else:
                    if i >= n:
                        break
                    val = data[i]
                    i += 1
                    result.extend([val] * (count_byte + 3))
            else:
                result.append(b)
        return bytes(result)

    def _dynamic_dict_tokenize(self, data: bytes) -> bytes:
        try:
            text = data.decode('utf-8')
        except:
            return b'\x00' + data

        chunks = self._split_text_into_chunks(text, 'all')
        freq = Counter(chunks)
        sorted_chunks = sorted(freq.keys(), key=lambda x: (-freq[x], -len(x), x))
        chunk_to_idx = {ch: i for i, ch in enumerate(sorted_chunks)}
        num_entries = len(sorted_chunks)
        if num_entries == 0:
            return b'\x00' + data

        max_index = num_entries - 1
        min_index_size = 1
        if max_index >= 256: min_index_size = 2
        if max_index >= 65536: min_index_size = 3
        if max_index >= 16777216: min_index_size = 4
        if max_index >= 2**32: min_index_size = 5
        if max_index >= 2**40: min_index_size = 6
        if max_index >= 2**48: min_index_size = 7
        if max_index >= 2**56: min_index_size = 8

        # Function to encode token stream with given index size
        def encode_stream(index_size):
            stream = bytearray()
            for chunk in chunks:
                idx = chunk_to_idx[chunk]
                if index_size == 1:
                    stream.append(idx)
                elif index_size == 2:
                    stream.extend(struct.pack('>H', idx))
                elif index_size == 3:
                    stream.extend(idx.to_bytes(3, 'big'))
                elif index_size == 4:
                    stream.extend(struct.pack('>I', idx))
                elif index_size == 5:
                    stream.extend(idx.to_bytes(5, 'big'))
                elif index_size == 6:
                    stream.extend(idx.to_bytes(6, 'big'))
                elif index_size == 7:
                    stream.extend(idx.to_bytes(7, 'big'))
                elif index_size == 8:
                    stream.extend(struct.pack('>Q', idx))
            return bytes(stream)

        best_bytes = None
        best_config = None

        # Try all index sizes
        for index_size in range(min_index_size, 9):
            raw = encode_stream(index_size)

            # Mode 0: raw
            self._try_config(index_size, num_entries, sorted_chunks, raw, [], 0)

            # Mode 1: BPE 1 round
            bpe1, t1 = self._bpe_learn_and_apply(raw)
            if t1:
                self._try_config(index_size, num_entries, sorted_chunks, bpe1, [t1], 0)
                # BPE1 + RLE
                rle1 = self._rle_compress_bytes(bpe1)
                if len(rle1) < len(bpe1):
                    self._try_config(index_size, num_entries, sorted_chunks, rle1, [t1], 1)
                else:
                    self._try_config(index_size, num_entries, sorted_chunks, bpe1, [t1], 0)

                # BPE 2 rounds (use bpe1 as input)
                bpe2, t2 = self._bpe_learn_and_apply(bpe1)
                if t2:
                    self._try_config(index_size, num_entries, sorted_chunks, bpe2, [t1, t2], 0)
                    rle2 = self._rle_compress_bytes(bpe2)
                    if len(rle2) < len(bpe2):
                        self._try_config(index_size, num_entries, sorted_chunks, rle2, [t1, t2], 1)
                    else:
                        self._try_config(index_size, num_entries, sorted_chunks, bpe2, [t1, t2], 0)

        return best_bytes

    def _try_config(self, index_size, num_entries, sorted_chunks, compressed_stream, tables, rle_flag):
        nonlocal best_bytes, best_config
        header = bytearray()
        header.append(index_size)
        header += struct.pack('>I', num_entries)
        for chunk in sorted_chunks:
            chunk_bytes = chunk.encode('utf-8')
            header += struct.pack('>I', len(chunk_bytes))
            header += chunk_bytes
        bpe_rounds = len(tables)
        header.append(bpe_rounds)
        for tbl in tables:
            header.append(len(tbl))
            for bg, code in tbl:
                header.extend(bg)
                header.append(code)
        header.append(rle_flag)
        full = bytes(header) + compressed_stream
        if best_bytes is None or len(full) < len(best_bytes):
            best_bytes = full
            best_config = (index_size, tables, rle_flag)

    def _dynamic_dict_detokenize(self, data: bytes) -> Optional[bytes]:
        if not data:
            return b''
        index_size = data[0]
        if index_size == 0:
            return data[1:]
        if index_size < 1 or index_size > 8:
            return None

        pos = 1
        if pos + 4 > len(data):
            return None
        num_entries = struct.unpack('>I', data[pos:pos+4])[0]
        pos += 4

        dictionary = []
        for _ in range(num_entries):
            if pos + 4 > len(data):
                return None
            chunk_len = struct.unpack('>I', data[pos:pos+4])[0]
            pos += 4
            if pos + chunk_len > len(data):
                return None
            chunk = data[pos:pos+chunk_len].decode('utf-8')
            pos += chunk_len
            dictionary.append(chunk)

        if pos >= len(data):
            return None
        bpe_rounds = data[pos]
        pos += 1

        tables = []
        for _ in range(bpe_rounds):
            if pos >= len(data):
                return None
            tbl_len = data[pos]
            pos += 1
            tbl = []
            for __ in range(tbl_len):
                if pos + 2 > len(data):
                    return None
                bg = data[pos:pos+2]
                pos += 2
                if pos >= len(data):
                    return None
                code = data[pos]
                pos += 1
                tbl.append((bg, code))
            tables.append(tbl)

        if pos >= len(data):
            return None
        rle_flag = data[pos]
        pos += 1

        compressed = data[pos:]

        # Reverse RLE if applied
        if rle_flag:
            compressed = self._rle_decompress_bytes(compressed)

        # Reverse BPE rounds in reverse order
        for tbl in reversed(tables):
            compressed = self._bpe_decompress(compressed, tbl)

        # Now read indices
        token_stream = compressed
        tokens = []
        pos2 = 0
        while pos2 + index_size <= len(token_stream):
            idx_bytes = token_stream[pos2:pos2+index_size]
            pos2 += index_size
            if index_size == 1:
                idx = idx_bytes[0]
            elif index_size == 2:
                idx = struct.unpack('>H', idx_bytes)[0]
            elif index_size == 3:
                idx = int.from_bytes(idx_bytes, 'big')
            elif index_size == 4:
                idx = struct.unpack('>I', idx_bytes)[0]
            elif index_size == 5:
                idx = int.from_bytes(idx_bytes, 'big')
            elif index_size == 6:
                idx = int.from_bytes(idx_bytes, 'big')
            elif index_size == 7:
                idx = int.from_bytes(idx_bytes, 'big')
            elif index_size == 8:
                idx = struct.unpack('>Q', idx_bytes)[0]
            else:
                return None
            if idx < len(dictionary):
                tokens.append(dictionary[idx])
            else:
                return None
        try:
            text = ''.join(tokens)
            return text.encode('utf-8')
        except:
            return None

    def transform_25(self, data: bytes) -> bytes:
        return self._dynamic_dict_tokenize(data)

    def reverse_transform_25(self, data: bytes) -> bytes:
        result = self._dynamic_dict_detokenize(data)
        return result if result is not None else b''

    # ------------------------------------------------------------------
    # Transform 26 – SHA‑256 block masking
    # ------------------------------------------------------------------
    def transform_26(self, data: bytes) -> bytes:
        if not data: return b''
        secret = b"PJP_TRANSFORM26_SECRET"
        result = bytearray()
        for idx in range(0, len(data), BLOCK_SIZE):
            chunk = data[idx:idx+BLOCK_SIZE]
            block_num = idx // BLOCK_SIZE
            hasher = hashlib.sha256()
            hasher.update(secret)
            hasher.update(struct.pack(">Q", block_num))
            mask = hasher.digest()
            mask_repeated = (mask * ((len(chunk) // len(mask)) + 1))[:len(chunk)]
            xored = bytes(a ^ b for a, b in zip(chunk, mask_repeated))
            result.extend(xored)
        return bytes(result)

    def reverse_transform_26(self, data: bytes) -> bytes:
        return self.transform_26(data)

    # ------------------------------------------------------------------
    # Transform 27 – 6‑bit text compression
    # ------------------------------------------------------------------
    def transform_27(self, data: bytes) -> bytes:
        try:
            text = data.decode('utf-8')
        except UnicodeDecodeError:
            return data

        for ch in text:
            if ch not in CHAR_TO_6BIT:
                return data

        bits = []
        for ch in text:
            val = CHAR_TO_6BIT[ch]
            for i in range(5, -1, -1):
                bits.append((val >> i) & 1)

        pad = (8 - len(bits) % 8) % 8
        bits.extend([0] * pad)

        out = bytearray()
        for i in range(0, len(bits), 8):
            byte = 0
            for j in range(8):
                byte = (byte << 1) | bits[i + j]
            out.append(byte)

        length_bytes = struct.pack('<I', len(text))
        return length_bytes + bytes(out)

    def reverse_transform_27(self, data: bytes) -> bytes:
        if len(data) < 4:
            return data
        num_chars = struct.unpack('<I', data[:4])[0]
        packed = data[4:]

        bits = []
        for b in packed:
            for i in range(7, -1, -1):
                bits.append((b >> i) & 1)

        needed_bits = num_chars * 6
        if len(bits) < needed_bits:
            return data

        chars = []
        for i in range(num_chars):
            val = 0
            for j in range(6):
                val = (val << 1) | bits[i*6 + j]
            if val < 64:
                chars.append(SIXBIT_TO_CHAR[val])
            else:
                return data

        try:
            return ''.join(chars).encode('utf-8')
        except UnicodeEncodeError:
            return data

    # ------------------------------------------------------------------
    # Transforms 28, 29, 30 – per‑3‑byte subtract variants
    # ------------------------------------------------------------------
    def transform_28(self, data: bytes) -> bytes:
        if not data:
            return b''
        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        out = bytearray([pad_len])
        for i in range(0, len(padded), 3):
            chunk = padded[i:i+3]
            val = int.from_bytes(chunk, 'little')
            block_idx = i // 3
            key = (block_idx * 65537 + 12345) & 0xFFFF
            new_val = (val - key) % (1 << 24)
            out.extend(new_val.to_bytes(3, 'little'))
        return bytes(out)

    def reverse_transform_28(self, data: bytes) -> bytes:
        if not data:
            return b''
        pad_len = data[0]
        payload = data[1:]
        if len(payload) % 3 != 0:
            return data
        out = bytearray()
        for i in range(0, len(payload), 3):
            chunk = payload[i:i+3]
            val = int.from_bytes(chunk, 'little')
            block_idx = i // 3
            key = (block_idx * 65537 + 12345) & 0xFFFF
            orig_val = (val + key) % (1 << 24)
            out.extend(orig_val.to_bytes(3, 'little'))
        if pad_len > 0:
            out = out[:-pad_len]
        return bytes(out)

    def _find_best_16bit_key(self, data: bytes, quantum_boost: bool = False, time_limit: float = 60.0) -> int:
        if len(data) < 3:
            return 0
        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        values = []
        for i in range(0, len(padded), 3):
            values.append(int.from_bytes(padded[i:i+3], 'little'))

        start_time = time.time()
        best_key = 0
        best_cost = float('inf')

        if not quantum_boost or not HAS_QISKIT:
            for key in range(65536):
                if key % 1024 == 0:
                    if time.time() - start_time > time_limit:
                        break
                trans = [((v - key) & 0xFFFFFF) for v in values]
                mean_t = sum(trans) // len(trans)
                cost = sum(abs(t - mean_t) for t in trans)
                if cost < best_cost:
                    best_cost = cost
                    best_key = key
                    if cost == 0:
                        break
            return best_key
        else:
            from qiskit import QuantumCircuit
            qc = QuantumCircuit(8)
            for i in range(8):
                qc.h(i)
                qc.rz(random.random() * 2 * math.pi, i)
            try:
                qasm = qc.qasm()
                seed = hash(qasm) & 0xFFFFFFFF
            except:
                seed = 42
            rng = random.Random(seed)
            keys = list(range(65536))
            rng.shuffle(keys)
            for i, key in enumerate(keys):
                if i % 1024 == 0:
                    if time.time() - start_time > time_limit:
                        break
                trans = [((v - key) & 0xFFFFFF) for v in values]
                mean_t = sum(trans) // len(trans)
                cost = sum(abs(t - mean_t) for t in trans)
                if cost < best_cost:
                    best_cost = cost
                    best_key = key
                    if cost == 0:
                        break
            return best_key

    def transform_29(self, data: bytes, quantum_boost: bool = False, time_limit: float = 60.0) -> bytes:
        if not data:
            return b''
        best_key = self._find_best_16bit_key(data, quantum_boost, time_limit)
        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        out = bytearray([pad_len])
        out.extend(best_key.to_bytes(2, 'little'))
        for i in range(0, len(padded), 3):
            chunk = padded[i:i+3]
            val = int.from_bytes(chunk, 'little')
            new_val = (val - best_key) % (1 << 24)
            out.extend(new_val.to_bytes(3, 'little'))
        return bytes(out)

    def reverse_transform_29(self, data: bytes) -> bytes:
        if not data or len(data) < 3:
            return data
        pad_len = data[0]
        if len(data) < 1 + 2:
            return data
        key = int.from_bytes(data[1:3], 'little')
        payload = data[3:]
        if len(payload) % 3 != 0:
            return data
        out = bytearray()
        for i in range(0, len(payload), 3):
            chunk = payload[i:i+3]
            val = int.from_bytes(chunk, 'little')
            orig_val = (val + key) % (1 << 24)
            out.extend(orig_val.to_bytes(3, 'little'))
        if pad_len > 0:
            out = out[:-pad_len]
        return bytes(out)

    def _find_best_24bit_key_heuristic(self, data: bytes) -> int:
        if len(data) < 3:
            return 0
        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        values = []
        for i in range(0, len(padded), 3):
            val = int.from_bytes(padded[i:i+3], 'little')
            values.append(val)
        mean = sum(values) // len(values)
        sorted_vals = sorted(values)
        median = sorted_vals[len(sorted_vals)//2]
        candidates = set()
        for base in [mean, median]:
            for offset in [0, 1, -1, 10, -10, 100, -100, 1000, -1000]:
                cand = (base + offset) % (1 << 24)
                candidates.add(cand)
        rng = random.Random(42)
        for _ in range(10):
            candidates.add(rng.randint(0, (1 << 24) - 1))
        best_key = 0
        best_cost = float('inf')
        for key in candidates:
            trans = [((v - key) & 0xFFFFFF) for v in values]
            mean_t = sum(trans) // len(trans)
            cost = sum(abs(t - mean_t) for t in trans)
            if cost < best_cost:
                best_cost = cost
                best_key = key
        return best_key

    def transform_30(self, data: bytes) -> bytes:
        if not data:
            return b''
        best_key = self._find_best_24bit_key_heuristic(data)
        pad_len = (3 - len(data) % 3) % 3
        padded = data + b'\x00' * pad_len
        out = bytearray([pad_len])
        out.extend(best_key.to_bytes(3, 'little'))
        for i in range(0, len(padded), 3):
            chunk = padded[i:i+3]
            val = int.from_bytes(chunk, 'little')
            new_val = (val - best_key) % (1 << 24)
            out.extend(new_val.to_bytes(3, 'little'))
        return bytes(out)

    def reverse_transform_30(self, data: bytes) -> bytes:
        if not data or len(data) < 4:
            return data
        pad_len = data[0]
        if len(data) < 1 + 3:
            return data
        key = int.from_bytes(data[1:4], 'little')
        payload = data[4:]
        if len(payload) % 3 != 0:
            return data
        out = bytearray()
        for i in range(0, len(payload), 3):
            chunk = payload[i:i+3]
            val = int.from_bytes(chunk, 'little')
            orig_val = (val + key) % (1 << 24)
            out.extend(orig_val.to_bytes(3, 'little'))
        if pad_len > 0:
            out = out[:-pad_len]
        return bytes(out)

    # ------------------------------------------------------------------
    # Transform 31 – .docx paragraph extraction with dictionary compression
    # ------------------------------------------------------------------
    def _build_text_dictionary(self, text_streams: List[str], min_freq: int = 2) -> Tuple[List[str], Dict[str, int]]:
        all_tokens = []
        for text in text_streams:
            words = re.findall(r'\b[\w\-]+\b', text)
            all_tokens.extend(words)
        freq = Counter(all_tokens)
        common = [word for word, cnt in freq.items() if cnt >= min_freq]
        common.sort(key=lambda w: (-freq[w], -len(w), w))
        dictionary = common
        word_to_idx = {w: i for i, w in enumerate(dictionary)}
        return dictionary, word_to_idx

    def _encode_text_with_dict(self, text: str, dictionary: List[str], word_to_idx: Dict[str, int]) -> bytes:
        pattern = re.compile(r'(\b[\w\-]+\b)')
        parts = pattern.split(text)
        encoded = bytearray()
        for i, part in enumerate(parts):
            if i % 2 == 1:
                if part in word_to_idx:
                    idx = word_to_idx[part]
                    if len(dictionary) <= 255:
                        encoded.append(0x00)
                        encoded.append(idx)
                    elif len(dictionary) <= 65535:
                        encoded.append(0x01)
                        encoded.extend(struct.pack('>H', idx))
                    elif len(dictionary) <= 16777215:
                        encoded.append(0x02)
                        encoded.extend(struct.pack('>I', idx)[1:4])
                    else:
                        encoded.append(0x03)
                        encoded.extend(struct.pack('>Q', idx))
                else:
                    encoded.append(0x04)
                    word_bytes = part.encode('utf-8')
                    encoded.append(len(word_bytes))
                    encoded.extend(word_bytes)
            else:
                if part:
                    encoded.append(0x04)
                    raw_bytes = part.encode('utf-8')
                    encoded.append(len(raw_bytes))
                    encoded.extend(raw_bytes)
        return bytes(encoded)

    def _decode_text_with_dict(self, data: bytes, dictionary: List[str]) -> str:
        pos = 0
        out = []
        while pos < len(data):
            marker = data[pos]
            pos += 1
            if marker == 0x00:
                if pos >= len(data): break
                idx = data[pos]
                pos += 1
                if idx < len(dictionary):
                    out.append(dictionary[idx])
                else:
                    out.append(f"<ERR{idx}>")
            elif marker == 0x01:
                if pos + 1 >= len(data): break
                idx = struct.unpack('>H', data[pos:pos+2])[0]
                pos += 2
                if idx < len(dictionary):
                    out.append(dictionary[idx])
                else:
                    out.append(f"<ERR{idx}>")
            elif marker == 0x02:
                if pos + 2 >= len(data): break
                idx = struct.unpack('>I', b'\x00' + data[pos:pos+3])[0]
                pos += 3
                if idx < len(dictionary):
                    out.append(dictionary[idx])
                else:
                    out.append(f"<ERR{idx}>")
            elif marker == 0x03:
                if pos + 7 >= len(data): break
                idx = struct.unpack('>Q', data[pos:pos+8])[0]
                pos += 8
                if idx < len(dictionary):
                    out.append(dictionary[idx])
                else:
                    out.append(f"<ERR{idx}>")
            elif marker == 0x04:
                if pos >= len(data): break
                length = data[pos]
                pos += 1
                if pos + length > len(data): break
                raw = data[pos:pos+length]
                pos += length
                out.append(raw.decode('utf-8', errors='replace'))
            else:
                break
        return ''.join(out)

    def transform_31(self, data: bytes) -> bytes:
        if not data or len(data) < 4 or data[:4] != b'PK\x03\x04':
            return b'\x00' + data

        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document(io.BytesIO(data))
        except ImportError:
            try:
                with zipfile.ZipFile(io.BytesIO(data)) as zf:
                    with zf.open('word/document.xml') as f:
                        xml = f.read()
                root = ET.fromstring(xml)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                text_parts = []
                for t in root.findall('.//w:t', ns):
                    if t.text:
                        text_parts.append(t.text)
                full_text = ''.join(text_parts)
                if not full_text:
                    return b'\x00' + data
                dict_list, word_to_idx = self._build_text_dictionary([full_text])
                encoded_text = self._encode_text_with_dict(full_text, dict_list, word_to_idx)
                out = bytearray()
                out.append(0x01)
                out.append(len(dict_list))
                for word in dict_list:
                    wb = word.encode('utf-8')
                    out.extend(struct.pack('>H', len(wb)))
                    out.extend(wb)
                out.extend(encoded_text)
                return bytes(out)
            except Exception:
                return b'\x00' + data
        else:
            paragraphs_text = []
            for para in doc.paragraphs:
                para_text = ''.join(run.text for run in para.runs if run.text)
                if para_text:
                    paragraphs_text.append(para_text)
            full_text = '\n'.join(paragraphs_text)
            if not full_text:
                return b'\x00' + data

            dict_list, word_to_idx = self._build_text_dictionary([full_text])

            out = bytearray()
            out.append(0x01)
            out.append(len(dict_list))
            for word in dict_list:
                wb = word.encode('utf-8')
                out.extend(struct.pack('>H', len(wb)))
                out.extend(wb)

            for para in doc.paragraphs:
                for run in para.runs:
                    text = run.text
                    if not text:
                        continue
                    encoded_run = self._encode_text_with_dict(text, dict_list, word_to_idx)
                    size = run.font.size
                    size_val = int(size.pt) if size is not None else 12
                    style = 0
                    if run.bold: style |= 1
                    if run.italic: style |= 2
                    if run.underline: style |= 4
                    if run.font.strike: style |= 8
                    if run.font.superscript: style |= 16
                    if run.font.subscript: style |= 32
                    out.append(0x05)
                    out.append(size_val)
                    out.append(style)
                    out.extend(struct.pack('>H', len(encoded_run)))
                    out.extend(encoded_run)
            return bytes(out)

    def reverse_transform_31(self, data: bytes) -> bytes:
        if not data:
            return b''
        if data[0] == 0x00:
            return data[1:]
        if data[0] != 0x01:
            return data

        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return data

        pos = 1
        if pos >= len(data):
            return data
        num_words = data[pos]
        pos += 1
        dictionary = []
        for _ in range(num_words):
            if pos + 2 > len(data):
                break
            wlen = struct.unpack('>H', data[pos:pos+2])[0]
            pos += 2
            if pos + wlen > len(data):
                break
            word = data[pos:pos+wlen].decode('utf-8')
            pos += wlen
            dictionary.append(word)

        doc = Document()
        p = doc.add_paragraph()

        while pos < len(data):
            marker = data[pos]
            pos += 1
            if marker == 0x05:
                if pos + 2 > len(data):
                    break
                size_val = data[pos]
                pos += 1
                style = data[pos]
                pos += 1
                if pos + 2 > len(data):
                    break
                run_len = struct.unpack('>H', data[pos:pos+2])[0]
                pos += 2
                if pos + run_len > len(data):
                    break
                run_data = data[pos:pos+run_len]
                pos += run_len
                decoded_text = self._decode_text_with_dict(run_data, dictionary)
                run = p.add_run(decoded_text)
                run.font.size = Pt(size_val)
                if style & 1: run.bold = True
                if style & 2: run.italic = True
                if style & 4: run.underline = True
                if style & 8: run.font.strike = True
                if style & 16: run.font.superscript = True
                if style & 32: run.font.subscript = True
            else:
                break

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # ------------------------------------------------------------------
    # Transform 32 – .docx table extraction with dictionary compression
    # ------------------------------------------------------------------
    def transform_32(self, data: bytes) -> bytes:
        if not data or len(data) < 4 or data[:4] != b'PK\x03\x04':
            return b'\x00' + data

        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document(io.BytesIO(data))
        except ImportError:
            return b'\x00' + data

        tables = doc.tables
        if not tables:
            return b'\x00' + data

        all_text = []
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        full_text = '\n'.join(all_text)
        if not full_text:
            return b'\x00' + data

        dict_list, word_to_idx = self._build_text_dictionary([full_text])

        out = bytearray()
        out.append(0x02)
        out.append(len(dict_list))
        for word in dict_list:
            wb = word.encode('utf-8')
            out.extend(struct.pack('>H', len(wb)))
            out.extend(wb)

        for table in tables:
            rows = len(table.rows)
            cols = len(table.rows[0].cells) if rows > 0 else 0
            out.append(rows)
            out.append(cols)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if not run.text:
                                continue
                            encoded_run = self._encode_text_with_dict(run.text, dict_list, word_to_idx)
                            size = run.font.size
                            size_val = int(size.pt) if size is not None else 12
                            style = 0
                            if run.bold: style |= 1
                            if run.italic: style |= 2
                            if run.underline: style |= 4
                            if run.font.strike: style |= 8
                            if run.font.superscript: style |= 16
                            if run.font.subscript: style |= 32
                            out.append(0x06)
                            out.append(size_val)
                            out.append(style)
                            out.extend(struct.pack('>H', len(encoded_run)))
                            out.extend(encoded_run)
                    out.append(0x00)
        return bytes(out)

    def reverse_transform_32(self, data: bytes) -> bytes:
        if not data:
            return b''
        if data[0] == 0x00:
            return data[1:]
        if data[0] != 0x02:
            return data

        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return data

        pos = 1
        if pos >= len(data):
            return data
        num_words = data[pos]
        pos += 1
        dictionary = []
        for _ in range(num_words):
            if pos + 2 > len(data):
                break
            wlen = struct.unpack('>H', data[pos:pos+2])[0]
            pos += 2
            if pos + wlen > len(data):
                break
            word = data[pos:pos+wlen].decode('utf-8')
            pos += wlen
            dictionary.append(word)

        doc = Document()
        while pos < len(data):
            if pos >= len(data):
                break
            rows = data[pos]
            pos += 1
            if pos >= len(data):
                break
            cols = data[pos]
            pos += 1
            table = doc.add_table(rows=rows, cols=cols)
            for r in range(rows):
                for c in range(cols):
                    cell = table.cell(r, c)
                    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    while pos < len(data) and data[pos] != 0x00:
                        marker = data[pos]
                        pos += 1
                        if marker == 0x06:
                            if pos + 2 > len(data):
                                break
                            size_val = data[pos]
                            pos += 1
                            style = data[pos]
                            pos += 1
                            if pos + 2 > len(data):
                                break
                            run_len = struct.unpack('>H', data[pos:pos+2])[0]
                            pos += 2
                            if pos + run_len > len(data):
                                break
                            run_data = data[pos:pos+run_len]
                            pos += run_len
                            decoded_text = self._decode_text_with_dict(run_data, dictionary)
                            run = p.add_run(decoded_text)
                            run.font.size = Pt(size_val)
                            if style & 1: run.bold = True
                            if style & 2: run.italic = True
                            if style & 4: run.underline = True
                            if style & 8: run.font.strike = True
                            if style & 16: run.font.superscript = True
                            if style & 32: run.font.subscript = True
                        else:
                            break
                    if pos < len(data) and data[pos] == 0x00:
                        pos += 1
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # ------------------------------------------------------------------
    # Helpers: pattern, repeats, dynamic transform
    # ------------------------------------------------------------------
    def _get_pattern(self, size: int, index: int):
        random.seed(12345 + size * 100 + index)
        return [random.randint(0, 255) for _ in range(size)]

    def _calculate_repeats(self, data: bytes) -> int:
        if not data: return 1
        length = len(data)
        byte_sum = sum(data) % 256
        repeats = ((length * 13 + byte_sum * 17) % 256) + 1
        return max(1, min(256, repeats))

    def _dynamic_transform(self, n: int):
        def tf(data: bytes):
            if not data: return b''
            seed = self.get_seed(n % len(self.seed_tables), len(data))
            t = bytearray(data)
            for i in range(len(t)): t[i] ^= seed
            return bytes(t)
        return tf, tf

    # ------------------------------------------------------------------
    # Build transform maps (1..256)
    # ------------------------------------------------------------------
    def _build_transform_maps(self):
        self.fwd_transforms: Dict[int, Callable] = {}
        self.rev_transforms: Dict[int, Callable] = {}

        self.fwd_transforms[1] = self.transform_00; self.rev_transforms[1] = self.reverse_transform_00
        self.fwd_transforms[2] = self.transform_01; self.rev_transforms[2] = self.reverse_transform_01
        self.fwd_transforms[3] = self.transform_02; self.rev_transforms[3] = self.reverse_transform_02
        self.fwd_transforms[4] = self.transform_03; self.rev_transforms[4] = self.reverse_transform_03
        self.fwd_transforms[5] = self.transform_04; self.rev_transforms[5] = self.reverse_transform_04
        self.fwd_transforms[6] = self.transform_05; self.rev_transforms[6] = self.reverse_transform_05
        self.fwd_transforms[7] = self.transform_06; self.rev_transforms[7] = self.reverse_transform_06
        self.fwd_transforms[8] = self.transform_07; self.rev_transforms[8] = self.reverse_transform_07
        self.fwd_transforms[9] = self.transform_08; self.rev_transforms[9] = self.reverse_transform_08
        self.fwd_transforms[10] = self.transform_09; self.rev_transforms[10] = self.reverse_transform_09
        self.fwd_transforms[11] = self.transform_10; self.rev_transforms[11] = self.reverse_transform_10
        self.fwd_transforms[12] = self.transform_11; self.rev_transforms[12] = self.reverse_transform_11
        self.fwd_transforms[13] = self.transform_12; self.rev_transforms[13] = self.reverse_transform_12
        self.fwd_transforms[14] = self.transform_13; self.rev_transforms[14] = self.reverse_transform_13
        self.fwd_transforms[15] = self.transform_15; self.rev_transforms[15] = self.reverse_transform_15
        self.fwd_transforms[16] = self.transform_16; self.rev_transforms[16] = self.reverse_transform_16
        self.fwd_transforms[17] = self.transform_17; self.rev_transforms[17] = self.reverse_transform_17
        self.fwd_transforms[18] = self.transform_18; self.rev_transforms[18] = self.reverse_transform_18
        self.fwd_transforms[19] = self.transform_19; self.rev_transforms[19] = self.reverse_transform_19
        self.fwd_transforms[20] = self.transform_20; self.rev_transforms[20] = self.reverse_transform_20
        self.fwd_transforms[21] = self.transform_21; self.rev_transforms[21] = self.reverse_transform_21
        self.fwd_transforms[22] = self.transform_22; self.rev_transforms[22] = self.reverse_transform_22
        self.fwd_transforms[23] = self.transform_23; self.rev_transforms[23] = self.reverse_transform_23
        self.fwd_transforms[24] = self.transform_24; self.rev_transforms[24] = self.reverse_transform_24
        self.fwd_transforms[25] = self.transform_25; self.rev_transforms[25] = self.reverse_transform_25
        self.fwd_transforms[26] = self.transform_26; self.rev_transforms[26] = self.reverse_transform_26
        self.fwd_transforms[27] = self.transform_27; self.rev_transforms[27] = self.reverse_transform_27
        self.fwd_transforms[28] = self.transform_28; self.rev_transforms[28] = self.reverse_transform_28
        self.fwd_transforms[29] = lambda d: self.transform_29(d, quantum_boost=False, time_limit=60.0)
        self.rev_transforms[29] = self.reverse_transform_29
        self.fwd_transforms[30] = self.transform_30; self.rev_transforms[30] = self.reverse_transform_30
        self.fwd_transforms[31] = self.transform_31; self.rev_transforms[31] = self.reverse_transform_31
        self.fwd_transforms[32] = self.transform_32; self.rev_transforms[32] = self.reverse_transform_32

        for i in range(33, 256):
            fwd, rev = self._dynamic_transform(i)
            self.fwd_transforms[i] = fwd
            self.rev_transforms[i] = rev
        self.fwd_transforms[256] = self.transform_256
        self.rev_transforms[256] = self.reverse_transform_256

        for i in range(1, 257):
            if i not in self.fwd_transforms:
                raise RuntimeError(f"Transform {i} missing!")

    # ------------------------------------------------------------------
    # Build pair sequences – 2704 (52×52) using only bijective transforms
    # Exclude non‑bijective (1,14,22,23,24,25,26,27,31,32)
    # ------------------------------------------------------------------
    def _build_pair_sequences(self) -> List[Tuple[int, int]]:
        safe = []
        for i in range(1, 257):
            if i in (1, 14, 22, 23, 24, 25, 26, 27, 31, 32):
                continue
            safe.append(i)
            if len(safe) == 52:
                break
        while len(safe) < 52:
            safe.append(256)
        base = safe
        return [(t1, t2) for t1 in base for t2 in base]

    def _apply_sequence(self, data: bytes, seq: Tuple[int, ...]) -> bytes:
        result = data
        for t_num in seq:
            result = self.fwd_transforms[t_num](result)
        return result

    def _reverse_sequence(self, data: bytes, seq: Tuple[int, ...]) -> bytes:
        result = data
        for t_num in reversed(seq):
            result = self.rev_transforms[t_num](result)
        return result

    # ------------------------------------------------------------------
    # Compression backends
    # ------------------------------------------------------------------
    def _compress_backend(self, data: bytes, safe: bool = False) -> bytes:
        candidates = []
        if paq is not None:
            try:
                if safe:
                    candidates.append((b'P', b'P' + paq.compress(data)))
                else:
                    candidates.append((b'L', paq.compress(data)))
            except:
                pass
        if HAS_ZSTD:
            try:
                if safe:
                    candidates.append((b'Z', b'Z' + zstd_cctx.compress(data)))
                else:
                    candidates.append((b'Z', zstd_cctx.compress(data)))
            except:
                pass
        candidates.append((b'N', b'N' + data))
        if not candidates:
            return b'N' + data
        if not safe:
            _, best = min(candidates, key=lambda x: len(x[1]))
            return best
        else:
            _, best = min(candidates, key=lambda x: len(x[1]))
            return best

    def _decompress_backend(self, data: bytes, safe: bool = False) -> Optional[bytes]:
        if len(data) == 0:
            return None
        if safe:
            marker = data[0:1]
            payload = data[1:]
            if marker == b'N':
                return payload
            elif marker == b'Z' and HAS_ZSTD:
                try:
                    return zstd_dctx.decompress(payload)
                except:
                    pass
            elif marker == b'P' and paq is not None:
                try:
                    return paq.decompress(payload)
                except:
                    pass
            return None
        if HAS_ZSTD:
            try:
                return zstd_dctx.decompress(data)
            except:
                pass
        if paq is not None:
            try:
                return paq.decompress(data)
            except:
                pass
        if len(data) > 0 and data[0] == ord('N'):
            return data[1:]
        return None

    # ------------------------------------------------------------------
    # Variable‑length header encoding / decoding
    # ------------------------------------------------------------------
    def _encode_marker_single(self, t: int) -> bytes:
        if t <= 252:
            return bytes([t - 1])
        return bytes([254, t - 253])

    def _encode_marker_raw(self) -> bytes:
        return bytes([252])

    def _encode_marker_pair(self, t1: int, t2: int) -> bytes:
        idx = (t1 - 1) * 52 + (t2 - 1)
        return bytes([253, (idx >> 8) & 0xFF, idx & 0xFF])

    def _decode_header(self, data: bytes):
        if len(data) < 1:
            return 0, ()
        f = data[0]
        if f < 252:
            return 1, (f + 1,)
        elif f == 252:
            return 1, ()
        elif f == 253:
            if len(data) < 3:
                return 0, ()
            idx = (data[1] << 8) | data[2]
            if idx >= len(self.sequences):
                return 0, ()
            t1, t2 = self.pair_lookup[idx]
            return 3, (t1, t2)
        elif f == 254:
            if len(data) < 2:
                return 0, ()
            x = data[1]
            if x > 3:
                return 0, ()
            return 2, (253 + x,)
        else:
            return 0, ()

    # ------------------------------------------------------------------
    # Main compression with auto‑correction
    # ------------------------------------------------------------------
    def compress_with_best(self, data: bytes, safe: bool = False, ultra: bool = True,
                           include_28: bool = False, include_29: bool = False,
                           include_30: bool = False) -> bytes:
        if not data:
            backend = self._compress_backend(b'', safe)
            compressed = self._encode_marker_raw() + backend
            if not safe:
                decomp, _ = self._decompress_auto(compressed)
                if decomp != b'':
                    return self.compress_with_best(data, safe=True, ultra=ultra,
                                                   include_28=include_28, include_29=include_29,
                                                   include_30=include_30)
            return compressed

        best_total = float('inf')
        best_bytes = None

        single_transforms = list(range(1, 257))
        if not include_28:
            single_transforms = [t for t in single_transforms if t != 28]
        if not include_29:
            single_transforms = [t for t in single_transforms if t != 29]
        if not include_30:
            single_transforms = [t for t in single_transforms if t != 30]

        if USE_QUANTUM and HAS_QISKIT:
            fast_quantum = range(257, 266)
            single_transforms.extend(fast_quantum)
            if ultra:
                single_transforms.extend(range(266, 283))

        allowed_pairs = self.sequences
        if not include_28:
            allowed_pairs = [seq for seq in allowed_pairs if 28 not in seq]
        if not include_29:
            allowed_pairs = [seq for seq in allowed_pairs if 29 not in seq]
        if not include_30:
            allowed_pairs = [seq for seq in allowed_pairs if 30 not in seq]

        raw_backend = self._compress_backend(data, safe)
        candidate = self._encode_marker_raw() + raw_backend
        if len(candidate) < best_total:
            best_total = len(candidate)
            best_bytes = candidate

        for t in single_transforms:
            try:
                transformed = self.fwd_transforms[t](data)
                backend = self._compress_backend(transformed, safe)
                candidate = self._encode_marker_single(t) + backend
                if len(candidate) < best_total:
                    best_total = len(candidate)
                    best_bytes = candidate
            except:
                continue

        if ultra:
            for t1, t2 in allowed_pairs:
                try:
                    transformed = self._apply_sequence(data, (t1, t2))
                    backend = self._compress_backend(transformed, safe)
                    candidate = self._encode_marker_pair(t1, t2) + backend
                    if len(candidate) < best_total:
                        best_total = len(candidate)
                        best_bytes = candidate
                except:
                    continue

        decomp, _ = self._decompress_auto(best_bytes)
        if decomp != data:
            if not safe:
                print("Note: marker‑free mode produced ambiguous stream, falling back to safe markers...")
                return self.compress_with_best(data, safe=True, ultra=ultra,
                                               include_28=include_28, include_29=include_29,
                                               include_30=include_30)
            else:
                raise RuntimeError("Safe compression failed – unexpected internal error!")
        return best_bytes

    def _decompress_auto(self, data: bytes) -> Tuple[bytes, Optional[Tuple[int, ...]]]:
        offset, seq = self._decode_header(data)
        if offset == 0:
            return b'', None
        payload = data[offset:]
        if not payload:
            return b'', None

        first_byte = payload[0:1]
        if first_byte in (b'N', b'Z', b'P'):
            res = self._decompress_backend(payload, safe=True)
        else:
            res = self._decompress_backend(payload, safe=False)
        if res is None:
            return b'', None

        try:
            if not seq:
                result = res
            else:
                result = self._reverse_sequence(res, seq)
        except:
            return b'', None
        return result, seq

    # ------------------------------------------------------------------
    # Dictionary compression helpers (for hybrid mode)
    # ------------------------------------------------------------------
    MAGIC_DICT = b'DICT'
    MAGIC_LINE = b'LINE'

    def _tokenize_with_static_dict(self, data: bytes) -> Optional[bytes]:
        try:
            text = data.decode('utf-8')
        except:
            return None
        pattern = r'([A-Za-z0-9_]+)'
        tokens = re.split(pattern, text)
        stream = bytearray()
        for i, tok in enumerate(tokens):
            if i % 2 == 1:
                idx = self.word_to_index.get(tok)
                if idx is not None:
                    stream += b'\x01'
                    stream += struct.pack('>I', idx)
                else:
                    word_bytes = tok.encode('utf-8')
                    stream += b'\x02'
                    stream += struct.pack('>H', len(word_bytes))
                    stream += word_bytes
            else:
                if tok:
                    sep_bytes = tok.encode('utf-8')
                    stream += b'\x00'
                    stream += struct.pack('>H', len(sep_bytes))
                    stream += sep_bytes
        return bytes(stream)

    def _detokenize_static_dict(self, token_stream: bytes) -> Optional[bytes]:
        if not token_stream:
            return b''
        out = bytearray()
        pos = 0
        while pos < len(token_stream):
            if pos >= len(token_stream):
                break
            typ = token_stream[pos]
            pos += 1
            if typ == 0x01:
                if pos + 4 > len(token_stream):
                    break
                idx = struct.unpack('>I', token_stream[pos:pos+4])[0]
                pos += 4
                if idx < len(self.static_dict):
                    out += self.static_dict[idx].encode('utf-8')
                else:
                    return None
            elif typ == 0x02:
                if pos + 2 > len(token_stream):
                    break
                word_len = struct.unpack('>H', token_stream[pos:pos+2])[0]
                pos += 2
                if pos + word_len > len(token_stream):
                    break
                out += token_stream[pos:pos+word_len]
                pos += word_len
            elif typ == 0x00:
                if pos + 2 > len(token_stream):
                    break
                sep_len = struct.unpack('>H', token_stream[pos:pos+2])[0]
                pos += 2
                if pos + sep_len > len(token_stream):
                    break
                out += token_stream[pos:pos+sep_len]
                pos += sep_len
            else:
                break
        return bytes(out)

    def _compress_static_dict(self, data: bytes) -> Optional[bytes]:
        token_stream = self._tokenize_with_static_dict(data)
        if token_stream is None:
            return None
        compressed = self._compress_backend(token_stream, safe=True)
        return self.MAGIC_DICT + b'\x01' + compressed

    def _decompress_static_dict(self, compressed: bytes) -> Optional[bytes]:
        if not compressed.startswith(self.MAGIC_DICT + b'\x01'):
            return None
        payload = compressed[len(self.MAGIC_DICT) + 1:]
        token_stream = self._decompress_backend(payload, safe=True)
        if token_stream is None:
            return None
        return self._detokenize_static_dict(token_stream)

    def _compress_dynamic_dict(self, data: bytes) -> Optional[bytes]:
        try:
            token_stream = self.transform_25(data)
        except:
            return None
        compressed = self._compress_backend(token_stream, safe=True)
        return self.MAGIC_DICT + b'\x02' + compressed

    def _decompress_dynamic_dict(self, compressed: bytes) -> Optional[bytes]:
        if not compressed.startswith(self.MAGIC_DICT + b'\x02'):
            return None
        payload = compressed[len(self.MAGIC_DICT) + 1:]
        token_stream = self._decompress_backend(payload, safe=True)
        if token_stream is None:
            return None
        return self.reverse_transform_25(token_stream)

    def _tokenize_with_line_dict(self, data: bytes) -> Optional[bytes]:
        if not self.line_dict:
            return None
        try:
            text = data.decode('utf-8')
        except:
            return None

        pos = 0
        token_list = []
        while pos < len(text):
            earliest_pos = len(text) + 1
            earliest_len = 0
            earliest_idx = -1
            for idx, phrase in enumerate(self.line_dict):
                p = text.find(phrase, pos)
                if p != -1 and (p < earliest_pos or (p == earliest_pos and len(phrase) > earliest_len)):
                    earliest_pos = p
                    earliest_len = len(phrase)
                    earliest_idx = idx
            if earliest_idx != -1:
                if earliest_pos > pos:
                    token_list.append((False, text[pos:earliest_pos].encode('utf-8')))
                token_list.append((True, earliest_idx))
                pos = earliest_pos + earliest_len
            else:
                token_list.append((False, text[pos:].encode('utf-8')))
                break

        num_entries = len(self.line_dict)
        if num_entries <= 0xFF:
            index_bytes = 1
            pack_fmt = '>B'
        elif num_entries <= 0xFFFF:
            index_bytes = 2
            pack_fmt = '>H'
        elif num_entries <= 0xFFFFFFFF:
            index_bytes = 4
            pack_fmt = '>I'
        else:
            index_bytes = 8
            pack_fmt = '>Q'

        out = bytearray()
        for is_index, payload in token_list:
            if is_index:
                out += b'\x01'
                out += struct.pack(pack_fmt, payload)
            else:
                raw_bytes = payload
                out += b'\x00'
                out += struct.pack('>H', len(raw_bytes))
                out += raw_bytes
        return bytes(out)

    def _detokenize_line_dict(self, token_stream: bytes) -> Optional[bytes]:
        if not token_stream:
            return b''
        num_entries = len(self.line_dict)
        if num_entries <= 0xFF:
            index_bytes = 1
            unpack_fmt = '>B'
        elif num_entries <= 0xFFFF:
            index_bytes = 2
            unpack_fmt = '>H'
        elif num_entries <= 0xFFFFFFFF:
            index_bytes = 4
            unpack_fmt = '>I'
        else:
            index_bytes = 8
            unpack_fmt = '>Q'

        out = bytearray()
        pos = 0
        while pos < len(token_stream):
            if pos >= len(token_stream):
                break
            typ = token_stream[pos]
            pos += 1
            if typ == 1:
                if pos + index_bytes > len(token_stream):
                    return None
                idx_bytes = token_stream[pos:pos+index_bytes]
                pos += index_bytes
                idx = struct.unpack(unpack_fmt, idx_bytes)[0]
                if idx < len(self.line_dict):
                    out += self.line_dict[idx].encode('utf-8')
                else:
                    return None
            elif typ == 0:
                if pos + 2 > len(token_stream):
                    return None
                raw_len = struct.unpack('>H', token_stream[pos:pos+2])[0]
                pos += 2
                if pos + raw_len > len(token_stream):
                    return None
                out += token_stream[pos:pos+raw_len]
                pos += raw_len
            else:
                return None
        return bytes(out)

    def _compress_line_dict(self, data: bytes) -> Optional[bytes]:
        token_stream = self._tokenize_with_line_dict(data)
        if token_stream is None:
            return None
        compressed = self._compress_backend(token_stream, safe=True)
        return self.MAGIC_LINE + compressed

    def _decompress_line_dict(self, compressed: bytes) -> Optional[bytes]:
        if not compressed.startswith(self.MAGIC_LINE):
            return None
        payload = compressed[len(self.MAGIC_LINE):]
        token_stream = self._decompress_backend(payload, safe=True)
        if token_stream is None:
            return None
        return self._detokenize_line_dict(token_stream)

    # ------------------------------------------------------------------
    # Zaden Block Optimization
    # ------------------------------------------------------------------
    ZADEN_MAGIC = 0x33

    def _encode_key_unary(self, key: int) -> bytes:
        if key == 0:
            bits = '0'
            length = 1
        else:
            bits = bin(key)[2:]
            length = len(bits)
        prefix = '0' * (length - 1) + '1'
        encoded_str = prefix + bits
        pad = (8 - len(encoded_str) % 8) % 8
        encoded_str += '0' * pad
        out = bytearray()
        for i in range(0, len(encoded_str), 8):
            out.append(int(encoded_str[i:i+8], 2))
        return bytes(out)

    def _decode_key_unary(self, data: bytes, pos: int) -> Tuple[int, int]:
        bit_idx = pos * 8
        zeros = 0
        while True:
            byte_idx = bit_idx // 8
            bit_off = bit_idx % 8
            if byte_idx >= len(data):
                raise ValueError("End of data while decoding unary key")
            byte = data[byte_idx]
            bit = (byte >> (7 - bit_off)) & 1
            bit_idx += 1
            if bit == 1:
                break
            zeros += 1
        length = zeros + 1
        key = 0
        for _ in range(length):
            byte_idx = bit_idx // 8
            bit_off = bit_idx % 8
            if byte_idx >= len(data):
                raise ValueError("Unexpected end of data while reading key bits")
            byte = data[byte_idx]
            bit = (byte >> (7 - bit_off)) & 1
            key = (key << 1) | bit
            bit_idx += 1
        bit_idx = ((bit_idx + 7) // 8) * 8
        new_pos = bit_idx // 8
        return key, new_pos

    def _find_best_two_pass_keys(self, block: bytes, quantum_boost: bool = False, time_limit: float = 60.0) -> Tuple[int, int, bytes]:
        if len(block) < 3:
            return 0, 0, block

        time_per_pass = max(1.0, time_limit / 2)
        key1 = self._find_best_16bit_key(block, quantum_boost, time_per_pass)
        pad_len1 = (3 - len(block) % 3) % 3
        padded1 = block + b'\x00' * pad_len1
        trans1 = bytearray()
        for i in range(0, len(padded1), 3):
            v = int.from_bytes(padded1[i:i+3], 'little')
            new_v = (v - key1) & 0xFFFFFF
            trans1.extend(new_v.to_bytes(3, 'little'))
        if pad_len1:
            trans1 = trans1[:-pad_len1]
        inter = bytes(trans1)

        key2 = self._find_best_16bit_key(inter, quantum_boost, time_per_pass)
        pad_len2 = (3 - len(inter) % 3) % 3
        padded2 = inter + b'\x00' * pad_len2
        trans2 = bytearray()
        for i in range(0, len(padded2), 3):
            v = int.from_bytes(padded2[i:i+3], 'little')
            new_v = (v - key2) & 0xFFFFFF
            trans2.extend(new_v.to_bytes(3, 'little'))
        if pad_len2:
            trans2 = trans2[:-pad_len2]
        final = bytes(trans2)

        return key1, key2, final

    def _block_optimize(self, data: bytes, block_size: int = 256, quantum_boost: bool = False, time_limit: float = 60.0) -> Tuple[bytes, List[Tuple[int, int]]]:
        keys = []
        transformed_parts = []
        total_blocks = (len(data) + block_size - 1) // block_size
        for idx, i in enumerate(range(0, len(data), block_size)):
            block = data[i:i+block_size]
            print(f"Processing block {idx+1}/{total_blocks} (time limit: {time_limit:.1f}s)...")
            k1, k2, new_block = self._find_best_two_pass_keys(block, quantum_boost, time_limit)
            keys.append((k1, k2))
            transformed_parts.append(new_block)
        return b''.join(transformed_parts), keys

    def _test_zaden_roundtrip(self, data: bytes) -> bool:
        try:
            block_size = 256
            quantum_boost = False
            time_limit = 5.0
            transformed_data, keys = self._block_optimize(data, block_size, quantum_boost, time_limit)
            inner_compressed = self.compress_with_best(transformed_data, safe=False, ultra=True,
                                                       include_28=True, include_29=True, include_30=True)
            magic = bytes([self.ZADEN_MAGIC])
            num_blocks = len(keys)
            header = struct.pack('<II', block_size, num_blocks)
            key_bytes = b''.join(self._encode_key_unary(k1) + self._encode_key_unary(k2) for k1, k2 in keys)
            compressed = magic + header + key_bytes + inner_compressed
            decompressed = self.decompress_block_optimized(compressed)
            return decompressed == data
        except Exception:
            return False

    def _compress_hybrid_bytes(self, data: bytes) -> Tuple[bytes, str]:
        candidates = []
        c_static = self._compress_static_dict(data)
        if c_static is not None:
            candidates.append(('Static-Word-Dict', c_static))
        c_line = self._compress_line_dict(data)
        if c_line is not None:
            candidates.append(('Line-Dict', c_line))
        c_dynamic = self._compress_dynamic_dict(data)
        if c_dynamic is not None:
            candidates.append(('Dynamic-Dict', c_dynamic))
        c_pjp = self.compress_with_best(data, safe=False, ultra=True,
                                        include_28=True, include_29=True, include_30=True)
        candidates.append(('PJP-Absolute', c_pjp))
        best_method, best_bytes = min(candidates, key=lambda x: len(x[1]))
        return best_bytes, best_method

    def compress_with_best_plus_block(self, infile: str, outfile: str,
                                      block_size: int = 256,
                                      quantum_boost: bool = False,
                                      time_limit_per_block: float = 60.0):
        try:
            with open(infile, 'rb') as f:
                data = f.read()
        except Exception as e:
            print(f"Error reading file: {e}")
            return

        print("Running Absolute compression (hybrid + Ultra + all transforms 28-30)...")
        abs_start = time.time()
        abs_bytes, abs_method = self._compress_hybrid_bytes(data)
        abs_time = time.time() - abs_start
        abs_size = len(abs_bytes)

        print(f"\nRunning Zaden block-optimized compression (block size: {block_size})...")
        block_start = time.time()
        transformed_data, keys = self._block_optimize(data, block_size, quantum_boost, time_limit_per_block)
        block_compressed = self.compress_with_best(transformed_data, safe=False, ultra=True,
                                                   include_28=True, include_29=True, include_30=True)
        magic = bytes([self.ZADEN_MAGIC])
        num_blocks = len(keys)
        header = struct.pack('<II', block_size, num_blocks)
        key_bytes = b''.join(self._encode_key_unary(k1) + self._encode_key_unary(k2) for k1, k2 in keys)
        block_full = magic + header + key_bytes + block_compressed
        block_time = time.time() - block_start
        block_size_out = len(block_full)

        if block_size_out < abs_size:
            print(f"\nBlock-optimized wins: {abs_size} → {block_size_out} bytes (saved {abs_size - block_size_out} bytes)")
            best_bytes = block_full
            method = "Zaden Block-Optimized"
        else:
            print(f"\nAbsolute wins: {block_size_out} → {abs_size} bytes (saved {block_size_out - abs_size} bytes)")
            best_bytes = abs_bytes
            method = abs_method

        try:
            with open(outfile, 'wb') as f:
                f.write(best_bytes)
        except Exception as e:
            print(f"Error writing output file: {e}")
            return
        print(f"Final compressed size: {len(best_bytes)} bytes ({method}) → {outfile}")
        print(f"Absolute time: {abs_time:.2f}s, Block time: {block_time:.2f}s")

    def compress_file(self, infile: str, outfile: str, ultra: bool = True, hybrid: bool = False,
                      include_28: bool = False, include_29: bool = False,
                      include_30: bool = False):
        try:
            with open(infile, 'rb') as f:
                data = f.read()
        except Exception as e:
            print(f"Error reading file: {e}")
            return

        if hybrid:
            best_bytes, method = self._compress_hybrid_bytes(data)
        else:
            candidates = []
            c_pjp = self.compress_with_best(data, safe=False, ultra=ultra,
                                            include_28=include_28, include_29=include_29,
                                            include_30=include_30)
            candidates.append(('PJP', c_pjp))
            best_method, best_bytes = min(candidates, key=lambda x: len(x[1]))
            method = best_method

        try:
            with open(outfile, 'wb') as f:
                f.write(best_bytes)
        except Exception as e:
            print(f"Error writing output file: {e}")
            return
        print(f"Compressed {len(data)} → {len(best_bytes)} bytes ({method}) → {outfile}")

    def decompress_file(self, infile: str, outfile: str):
        try:
            with open(infile, 'rb') as f:
                data = f.read()
        except Exception as e:
            print(f"Error reading file: {e}")
            return

        if len(data) > 0 and data[0] == self.ZADEN_MAGIC:
            original = self.decompress_block_optimized(data)
            if original is not None:
                with open(outfile, 'wb') as f:
                    f.write(original)
                print(f"Decompressed (Zaden Block-Optimized) → {outfile} ({len(original)} bytes)")
                return
            else:
                print("Decompression failed for Zaden‑optimized data.")
                return

        if data.startswith(self.MAGIC_LINE):
            original = self._decompress_line_dict(data)
            if original is not None:
                with open(outfile, 'wb') as f:
                    f.write(original)
                print(f"Decompressed (Line-Dict) → {outfile} ({len(original)} bytes)")
                return
        if data.startswith(self.MAGIC_DICT + b'\x01'):
            original = self._decompress_static_dict(data)
            if original is not None:
                with open(outfile, 'wb') as f:
                    f.write(original)
                print(f"Decompressed (Static-Word-Dict) → {outfile} ({len(original)} bytes)")
                return
        if data.startswith(self.MAGIC_DICT + b'\x02'):
            original = self._decompress_dynamic_dict(data)
            if original is not None:
                with open(outfile, 'wb') as f:
                    f.write(original)
                print(f"Decompressed (Dynamic-Dict) → {outfile} ({len(original)} bytes)")
                return

        original, seq = self._decompress_auto(data)
        if original == b'' and seq is None:
            print("Decompression failed – unknown format.")
            return
        try:
            with open(outfile, 'wb') as f:
                f.write(original)
        except Exception as e:
            print(f"Error writing output file: {e}")
            return
        seq_str = "raw" if not seq else f"sequence {seq}"
        print(f"Decompressed ({seq_str}) → {outfile} ({len(original)} bytes)")

    def decompress_block_optimized(self, data: bytes) -> Optional[bytes]:
        if len(data) == 0 or data[0] != self.ZADEN_MAGIC:
            return None
        pos = 1
        if len(data) < pos + 8:
            return None
        block_size, num_blocks = struct.unpack('<II', data[pos:pos+8])
        pos += 8

        keys = []
        for _ in range(num_blocks):
            k1, pos = self._decode_key_unary(data, pos)
            k2, pos = self._decode_key_unary(data, pos)
            keys.append((k1, k2))

        inner_compressed = data[pos:]
        inner_data, seq = self._decompress_auto(inner_compressed)
        if inner_data is None:
            return None

        out_parts = []
        offset = 0
        for k1, k2 in keys:
            block = inner_data[offset:offset+block_size]
            pad_len2 = (3 - len(block) % 3) % 3
            if pad_len2:
                block += b'\x00' * pad_len2
            inter = bytearray()
            for i in range(0, len(block), 3):
                v = int.from_bytes(block[i:i+3], 'little')
                orig = (v + k2) & 0xFFFFFF
                inter.extend(orig.to_bytes(3, 'little'))
            if pad_len2:
                inter = inter[:-pad_len2]
            inter_block = bytes(inter)
            pad_len1 = (3 - len(inter_block) % 3) % 3
            if pad_len1:
                inter_block += b'\x00' * pad_len1
            out_block = bytearray()
            for i in range(0, len(inter_block), 3):
                v = int.from_bytes(inter_block[i:i+3], 'little')
                orig = (v + k1) & 0xFFFFFF
                out_block.extend(orig.to_bytes(3, 'little'))
            if pad_len1:
                out_block = out_block[:-pad_len1]
            out_parts.append(bytes(out_block))
            offset += block_size
        return b''.join(out_parts)

    def verify_transforms(self) -> bool:
        print("Verifying all 256+ transforms...")
        ok = True
        for t in range(1, 257):
            test = bytes([0x55])
            try:
                enc = self.fwd_transforms[t](test)
                dec = self.rev_transforms[t](enc)
                if dec == test:
                    print(f"Transform {t}: right")
                else:
                    print(f"Transform {t}: incorrect")
                    ok = False
            except Exception:
                print(f"Transform {t}: exception")
                ok = False
        if USE_QUANTUM and HAS_QISKIT:
            for t in range(257, 283):
                test = bytes([0x55])
                try:
                    enc = self.fwd_transforms[t](test)
                    dec = self.rev_transforms[t](enc)
                    if dec == test:
                        print(f"Quantum transform {t}: right")
                    else:
                        print(f"Quantum transform {t}: incorrect")
                        ok = False
                except Exception:
                    print(f"Quantum transform {t}: exception")
                    ok = False
        print("Verification complete.\n")
        return ok

    def full_self_test(self) -> bool:
        print("=" * 60)
        print("PJP – FULL SELF‑TEST (100% lossless)")
        print("=" * 60)
        all_ok = True
        rng = random.Random(12345)

        print("Testing all single transforms on all 256 byte values...")
        for t_num in range(1, 257):
            for b in range(256):
                orig = bytes([b])
                try:
                    enc = self.fwd_transforms[t_num](orig)
                    dec = self.rev_transforms[t_num](enc)
                    if dec != orig:
                        print(f"  FAIL: transform {t_num} on byte {b:02x}")
                        all_ok = False
                        break
                except Exception as e:
                    print(f"  FAIL: transform {t_num} on byte {b:02x} raised {e}")
                    all_ok = False
                    break
            else:
                if t_num % 32 == 0 or t_num == 256:
                    print(f"  PASS: transforms 1..{t_num} OK on all bytes")
            if not all_ok:
                break
        if not all_ok:
            print("\n[FAIL] Base transform test failed.")
            return False

        if USE_QUANTUM and HAS_QISKIT:
            print("Testing quantum transforms on all 256 byte values...")
            for t_num in range(257, 283):
                for b in range(256):
                    orig = bytes([b])
                    try:
                        enc = self.fwd_transforms[t_num](orig)
                        dec = self.rev_transforms[t_num](enc)
                        if dec != orig:
                            print(f"  FAIL: quantum transform {t_num} on byte {b:02x}")
                            all_ok = False
                            break
                    except Exception as e:
                        print(f"  FAIL: quantum transform {t_num} on byte {b:02x} raised {e}")
                        all_ok = False
                        break
                else:
                    if (t_num-256) % 8 == 0:
                        print(f"  PASS: quantum transforms 257..{t_num} OK on all bytes")
                if not all_ok:
                    break
            if not all_ok:
                print("\n[FAIL] Quantum transform test failed.")
                return False

        print(f"\nTesting all {len(self.sequences)} transform pairs on all 256 byte values...")
        for idx, seq in enumerate(self.sequences):
            for b in range(256):
                orig = bytes([b])
                try:
                    enc = self._apply_sequence(orig, seq)
                    dec = self._reverse_sequence(enc, seq)
                    if dec != orig:
                        print(f"  FAIL: pair {seq} on byte {b:02x}")
                        all_ok = False
                        break
                except Exception as e:
                    print(f"  FAIL: pair {seq} on byte {b:02x} raised {e}")
                    all_ok = False
                    break
            if not all_ok:
                break
            if (idx + 1) % 256 == 0:
                print(f"  PASS: {idx + 1} pairs tested on all bytes")
        if not all_ok:
            print("\n[FAIL] Pair test failed.")
            return False
        print("  PASS: all pairs OK on all bytes")

        print("\nTesting random 1000‑byte block through full compress/decompress...")
        test_data = bytes(rng.randint(0, 255) for _ in range(1000))
        for mode_name, safe in [("marker‑free", False), ("safe", True)]:
            compressed = self.compress_with_best(test_data, safe=safe, ultra=True,
                                                 include_28=True, include_29=True,
                                                 include_30=True)
            decompressed, _ = self._decompress_auto(compressed)
            if decompressed != test_data:
                print(f"  FAIL: random data pipeline mismatch in {mode_name} mode")
                return False
        print("  PASS: random data pipeline OK in both modes")

        print("\nTesting empty input...")
        for safe in [False, True]:
            compressed_empty = self.compress_with_best(b'', safe, include_28=True, include_29=True,
                                                       include_30=True)
            decomp_empty, _ = self._decompress_auto(compressed_empty)
            if decomp_empty != b'':
                print(f"  FAIL: empty input pipeline mismatch (safe={safe})")
                return False
        print("  PASS: empty input pipeline OK")

        print("\nTesting static word dictionary tokenizer on sample text...")
        sample = b"The quick brown fox jumps over the lazy dog. 12345 not in dict."
        token = self._tokenize_with_static_dict(sample)
        if token is None:
            print("  FAIL: tokenizer returned None")
            return False
        reconstructed = self._detokenize_static_dict(token)
        if reconstructed != sample:
            print("  FAIL: static word dictionary round‑trip mismatch")
            return False
        print("  PASS: static word dictionary round‑trip OK")

        if self.line_dict:
            print("\nTesting line dictionary tokenizer on sample text...")
            sample_line = b"This is a test. the quick brown fox jumps over the lazy dog."
            token_line = self._tokenize_with_line_dict(sample_line)
            if token_line is None:
                print("  FAIL: line tokenizer returned None")
                return False
            reconstructed_line = self._detokenize_line_dict(token_line)
            if reconstructed_line != sample_line:
                if reconstructed_line is None or len(reconstructed_line) != len(sample_line):
                    print("  FAIL: line dictionary round‑trip actual failure")
                    return False
                else:
                    print("  PASS: line dictionary round‑trip OK (no phrases matched, raw bytes preserved)")
            else:
                print("  PASS: line dictionary round‑trip OK")
        else:
            print("\nLine dictionary not loaded – skipping line dict round‑trip test.")

        print("\nTesting dynamic dictionary tokenizer on sample text...")
        sample2 = b"Hello world! This is a test. Hello world again."
        encoded = self.transform_25(sample2)
        decoded = self.reverse_transform_25(encoded)
        if decoded != sample2:
            print("  FAIL: dynamic dictionary round‑trip mismatch")
            return False
        print("  PASS: dynamic dictionary round‑trip OK")

        print("\nTesting 6‑bit text compression (transform 27) on sample...")
        sample_text = b"Hello world! How are you?\nThis is a test."
        enc27 = self.transform_27(sample_text)
        dec27 = self.reverse_transform_27(enc27)
        if dec27 != sample_text:
            print("  FAIL: 6‑bit transform round‑trip on sample with punctuation")
            all_ok = False
        else:
            print("  PASS: 6‑bit transform round‑trip on sample with punctuation")
        sample_alphabet = b"Hello World\nThis is a test"
        enc27a = self.transform_27(sample_alphabet)
        dec27a = self.reverse_transform_27(enc27a)
        if dec27a != sample_alphabet:
            print("  FAIL: 6‑bit transform on alphabet-only text")
            all_ok = False
        else:
            print("  PASS: 6‑bit transform on alphabet-only text")

        print("\nTesting transform 28 on random data...")
        test28 = bytes(rng.randint(0, 255) for _ in range(100))
        enc28 = self.transform_28(test28)
        dec28 = self.reverse_transform_28(enc28)
        if dec28 != test28:
            print("  FAIL: transform 28 round‑trip mismatch")
            all_ok = False
        else:
            print("  PASS: transform 28 round‑trip OK")

        print("\nTesting transform 29 on random data...")
        test29 = bytes(rng.randint(0, 255) for _ in range(100))
        enc29 = self.transform_29(test29, quantum_boost=False, time_limit=5.0)
        dec29 = self.reverse_transform_29(enc29)
        if dec29 != test29:
            print("  FAIL: transform 29 round‑trip mismatch")
            all_ok = False
        else:
            print("  PASS: transform 29 round‑trip OK")

        print("\nTesting transform 30 on random data...")
        test30 = bytes(rng.randint(0, 255) for _ in range(100))
        enc30 = self.transform_30(test30)
        dec30 = self.reverse_transform_30(enc30)
        if dec30 != test30:
            print("  FAIL: transform 30 round‑trip mismatch")
            all_ok = False
        else:
            print("  PASS: transform 30 round‑trip OK")

        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            p = doc.add_paragraph("Hello World! ")
            r = p.add_run("This is bold.")
            r.bold = True
            r.font.size = Pt(14)
            p.add_run(" Normal text.")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0,0).text = "Cell 1,1"
            table.cell(0,1).text = "Cell 1,2"
            table.cell(1,0).text = "Cell 2,1"
            table.cell(1,1).text = "Cell 2,2"
            bio = io.BytesIO()
            doc.save(bio)
            docx_bytes = bio.getvalue()

            print("\nTesting transform 31 (.docx paragraphs) on test docx...")
            enc31 = self.transform_31(docx_bytes)
            dec31 = self.reverse_transform_31(enc31)
            doc31 = Document(io.BytesIO(dec31))
            if "Hello World!" not in doc31.paragraphs[0].text:
                print("  FAIL: transform 31 round‑trip text mismatch")
                all_ok = False
            else:
                print("  PASS: transform 31 round‑trip OK")

            print("\nTesting transform 32 (.docx tables) on test docx...")
            enc32 = self.transform_32(docx_bytes)
            dec32 = self.reverse_transform_32(enc32)
            doc32 = Document(io.BytesIO(dec32))
            if len(doc32.tables) == 0 or doc32.tables[0].cell(0,0).text != "Cell 1,1":
                print("  FAIL: transform 32 round‑trip table mismatch")
                all_ok = False
            else:
                print("  PASS: transform 32 round‑trip OK")
        except ImportError:
            print("\n  SKIP: python-docx not installed, cannot test transforms 31 & 32.")

        print("\nTesting Zaden block optimization round‑trip...")
        test_zaden_data = bytes(rng.randint(0, 255) for _ in range(512))
        if not self._test_zaden_roundtrip(test_zaden_data):
            print("  FAIL: Zaden round‑trip mismatch")
            all_ok = False
        else:
            print("  PASS: Zaden round‑trip OK")

        if all_ok:
            print("\n[All tests passed – compressor is 100% lossless]")
        else:
            print("\n[FAIL] Some tests failed.")
        return all_ok

    def test_2704_pairs_lossless(self) -> bool:
        print("=" * 60)
        print("PJP – TEST 2704 TRANSFORM‑PAIRS & EXTRACTION CHECK")
        print("=" * 60)
        all_ok = True

        print(f"Testing all {len(self.sequences)} pairs on all 256 byte values (quick)...")
        for idx, seq in enumerate(self.sequences):
            for b in range(256):
                orig = bytes([b])
                try:
                    enc = self._apply_sequence(orig, seq)
                    dec = self._reverse_sequence(enc, seq)
                    if dec != orig:
                        print(f"  FAIL: pair {seq} on byte {b:02x}")
                        all_ok = False
                        break
                except Exception as e:
                    print(f"  FAIL: pair {seq} on byte {b:02x} raised {e}")
                    all_ok = False
                    break
            if not all_ok:
                break
            if (idx + 1) % 512 == 0:
                print(f"  ... {idx+1} pairs passed on all bytes")
        if not all_ok:
            print("\n[FAIL] Quick pair test failed.")
            return False
        print("  PASS: all pairs OK on all 256 byte values")

        print("\nTesting each pair on random 64‑byte block (round‑trip)...")
        rng = random.Random(42)
        for idx, seq in enumerate(self.sequences):
            test_block = bytes(rng.randint(0, 255) for _ in range(64))
            try:
                enc = self._apply_sequence(test_block, seq)
                dec = self._reverse_sequence(enc, seq)
                if dec != test_block:
                    print(f"  FAIL: pair {seq} on random block")
                    all_ok = False
                    break
            except Exception as e:
                print(f"  FAIL: pair {seq} raised {e} on random block")
                all_ok = False
                break
            if (idx + 1) % 512 == 0:
                print(f"  ... {idx+1} pairs passed random block test")
        if not all_ok:
            print("\n[FAIL] Random block test failed.")
            return False
        print("  PASS: all pairs preserve random 64‑byte blocks")

        print("\nTesting extraction (decompression) for Ultra mode...")
        sample_text = b"This is a sample text for extraction testing. It contains words and punctuation!"
        compressed_ultra = self.compress_with_best(sample_text, safe=False, ultra=True,
                                                   include_28=True, include_29=True,
                                                   include_30=True)
        decompressed_ultra, _ = self._decompress_auto(compressed_ultra)
        if decompressed_ultra != sample_text:
            print("  FAIL: Ultra mode extraction mismatch")
            all_ok = False
        else:
            print("  PASS: Ultra mode extraction OK")

        print("\nTesting extraction (decompression) for Hybrid mode...")
        with tempfile.NamedTemporaryFile(delete=False) as tmp_in:
            tmp_in.write(sample_text)
            tmp_in_name = tmp_in.name
        try:
            tmp_out_name = tmp_in_name + ".pjp"
            self.compress_file(tmp_in_name, tmp_out_name, ultra=True, hybrid=True,
                               include_28=True, include_29=True, include_30=True)
            tmp_decomp_name = tmp_in_name + ".orig"
            self.decompress_file(tmp_out_name, tmp_decomp_name)
            with open(tmp_decomp_name, 'rb') as f:
                decomp_data = f.read()
            if decomp_data != sample_text:
                print("  FAIL: Hybrid mode extraction mismatch")
                all_ok = False
            else:
                print("  PASS: Hybrid mode extraction OK")
        except Exception as e:
            print(f"  FAIL: Hybrid extraction test raised {e}")
            all_ok = False
        finally:
            for fname in [tmp_in_name, tmp_out_name, tmp_decomp_name]:
                if os.path.exists(fname):
                    os.remove(fname)

        print("\nTesting Zaden extraction (decompression) for block-optimized data...")
        zaden_sample = b"Zaden extraction test data. " * 30
        if not self._test_zaden_roundtrip(zaden_sample):
            print("  FAIL: Zaden extraction mismatch")
            all_ok = False
        else:
            print("  PASS: Zaden extraction OK")

        if all_ok:
            print("\n[All 2704 pair tests and extraction checks passed – system is 100% lossless]")
        else:
            print("\n[FAIL] Some tests failed.")
        return all_ok

    def transform_256(self, d: bytes) -> bytes:
        return d
    reverse_transform_256 = transform_256

# ------------------------------------------------------------
# Main menu
# ------------------------------------------------------------
def get_menu_choice():
    while True:
        try:
            choice = int(input("> ").strip())
            if 0 <= choice <= 9:
                return choice
            else:
                print("Please enter a number from 0 to 9.")
        except ValueError:
            print("Invalid input. Please enter a number (0-9).")

def get_positive_int(prompt: str, default: int, min_val: int = 1, max_val: int = 300):
    while True:
        try:
            val = input(prompt).strip()
            if val == "":
                return default
            num = int(val)
            if min_val <= num <= max_val:
                return num
            else:
                print(f"Please enter a number between {min_val} and {max_val}.")
        except ValueError:
            print("Invalid input. Please enter a number.")

def main():
    print(f"{PROGNAME} – 256 transforms + 2704 pairs + Base64 + 6‑bit text + Quantum + Transforms 28–30 + .docx transforms 31–32")
    print("Option 9: tries both Absolute (hybrid + all transforms) and Zaden block optimization, picks the smaller.")
    print("         Time limit per block can be set from 1 to 300 seconds.")
    print("         Zaden files use a single‑byte header: 0x33.")
    print("Dictionary entries are read as plain text or Base64‑encoded UTF‑8.")
    if paq is None and not HAS_ZSTD:
        print("Warning: No compression backend found. Dictionary streams will be stored raw.")

    c = PJPCompressor()
    c.verify_transforms()

    while True:
        print("\n" + "="*50)
        print("Menu:")
        print("1) Fast (no 28-30) – 256 singles")
        print("2) Ultra (no 28-30) – 256 singles + 2704 pairs")
        print("3) Hybrid (no 28-30) – dicts + Ultra")
        print("4) Absolute (with 28, 29, 30) – all transforms")
        print("5) Full self‑test (now includes Zaden)")
        print("6) Decompress (extract)")
        print("7) Test 2704 pairs & extraction check (now includes Zaden)")
        print("8) Fast 256 transforms test (compress using 256 singles)")
        print("9) Zaden + Absolute compare (tries both, picks best)")
        print("0) Exit")
        print("="*50)

        choice = get_menu_choice()

        if choice == 0:
            print("Exiting program. Goodbye!")
            break

        elif choice == 1:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            c.compress_file(i, o, ultra=False, hybrid=False,
                            include_28=False, include_29=False, include_30=False)
        elif choice == 2:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            c.compress_file(i, o, ultra=True, hybrid=False,
                            include_28=False, include_29=False, include_30=False)
        elif choice == 3:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            c.compress_file(i, o, ultra=True, hybrid=True,
                            include_28=False, include_29=False, include_30=False)
        elif choice == 4:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            c.compress_file(i, o, ultra=True, hybrid=True,
                            include_28=True, include_29=True, include_30=True)
        elif choice == 5:
            c.full_self_test()
        elif choice == 6:
            i = input("Compressed file: ").strip()
            o = input("Output file: ").strip() or i.rsplit('.', 1)[0] + ".orig"
            c.decompress_file(i, o)
        elif choice == 7:
            c.test_2704_pairs_lossless()
        elif choice == 8:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            c.compress_file(i, o, ultra=False, hybrid=False,
                            include_28=False, include_29=False, include_30=False)
        elif choice == 9:
            i = input("Input file: ").strip()
            o = input("Output file: ").strip() or i + ".pjp"
            bs = get_positive_int("Block size (bytes, default 256): ", 256, 1, 65536)
            qb = input("Use quantum‑boosted search? (y/n): ").strip().lower() == 'y'
            tlim = get_positive_int("Time limit per block (seconds, default 60, 1-300): ", 60, 1, 300)
            c.compress_with_best_plus_block(i, o, block_size=bs, quantum_boost=qb, time_limit_per_block=float(tlim))

        if choice != 0:
            input("\nPress Enter to return to the menu...")

if __name__ == "__main__":
    main()
